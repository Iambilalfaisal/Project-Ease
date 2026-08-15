import dataclasses
import io
import json
import logging
import mimetypes
import os
import re
import time
from collections.abc import AsyncGenerator, Awaitable, Callable
from pathlib import Path
from typing import Any, cast

from azure.cognitiveservices.speech import (
    ResultReason,
    SpeechConfig,
    SpeechSynthesisOutputFormat,
    SpeechSynthesisResult,
    SpeechSynthesizer,
)
from azure.core.credentials import AzureKeyCredential
from azure.identity.aio import (
    AzureDeveloperCliCredential,
    ManagedIdentityCredential,
    get_bearer_token_provider,
)
from azure.monitor.opentelemetry import configure_azure_monitor
from azure.search.documents.aio import SearchClient
from azure.search.documents.indexes.aio import SearchIndexClient
from azure.search.documents.knowledgebases.aio import KnowledgeBaseRetrievalClient
from opentelemetry.instrumentation.aiohttp_client import AioHttpClientInstrumentor
from opentelemetry.instrumentation.asgi import OpenTelemetryMiddleware
from opentelemetry.instrumentation.httpx import (
    HTTPXClientInstrumentor,
)
from opentelemetry.instrumentation.openai import OpenAIInstrumentor
from quart import (
    Blueprint,
    Quart,
    abort,
    current_app,
    jsonify,
    make_response,
    request,
    send_file,
    send_from_directory,
)
from quart_cors import cors

from approaches.approach import Approach, DataPoints
from approaches.chatreadretrieveread import ChatReadRetrieveReadApproach
from approaches.promptmanager import PromptManager
from chat_history.cosmosdb import chat_history_cosmosdb_bp
from config import (
    CONFIG_AGENTIC_KNOWLEDGEBASE_ENABLED,
    CONFIG_AUTH_CLIENT,
    CONFIG_CHAT_APPROACH,
    CONFIG_CHAT_HISTORY_BROWSER_ENABLED,
    CONFIG_CHAT_HISTORY_COSMOS_ENABLED,
    CONFIG_CREDENTIAL,
    CONFIG_DEFAULT_REASONING_EFFORT,
    CONFIG_DEFAULT_RETRIEVAL_REASONING_EFFORT,
    CONFIG_GLOBAL_BLOB_MANAGER,
    CONFIG_INGESTER,
    CONFIG_KNOWLEDGEBASE_CLIENT,
    CONFIG_KNOWLEDGEBASE_CLIENT_WITH_SHAREPOINT,
    CONFIG_KNOWLEDGEBASE_CLIENT_WITH_WEB,
    CONFIG_KNOWLEDGEBASE_CLIENT_WITH_WEB_AND_SHAREPOINT,
    CONFIG_LANGUAGE_PICKER_ENABLED,
    CONFIG_MULTIMODAL_ENABLED,
    CONFIG_OPENAI_CLIENT,
    CONFIG_QUERY_REWRITING_ENABLED,
    CONFIG_RAG_SEARCH_IMAGE_EMBEDDINGS,
    CONFIG_RAG_SEARCH_TEXT_EMBEDDINGS,
    CONFIG_RAG_SEND_IMAGE_SOURCES,
    CONFIG_RAG_SEND_TEXT_SOURCES,
    CONFIG_REASONING_EFFORT_ENABLED,
    CONFIG_REASONING_EFFORT_OPTIONS,
    CONFIG_SEARCH_CLIENT,
    CONFIG_SEMANTIC_RANKER_DEPLOYED,
    CONFIG_SHAREPOINT_SOURCE_ENABLED,
    CONFIG_SPEECH_INPUT_ENABLED,
    CONFIG_SPEECH_OUTPUT_AZURE_ENABLED,
    CONFIG_SPEECH_OUTPUT_BROWSER_ENABLED,
    CONFIG_SPEECH_SERVICE_ID,
    CONFIG_SPEECH_SERVICE_LOCATION,
    CONFIG_SPEECH_SERVICE_TOKEN,
    CONFIG_SPEECH_SERVICE_VOICE,
    CONFIG_STREAMING_ENABLED,
    CONFIG_USER_BLOB_MANAGER,
    CONFIG_USER_UPLOAD_ENABLED,
    CONFIG_VECTOR_SEARCH_ENABLED,
    CONFIG_WEB_SOURCE_ENABLED,
)
from core.authentication import AuthenticationHelper
from core.sessionhelper import create_session_id
from decorators import authenticated, authenticated_path
from error import error_dict, error_response
from prepdocs import (
    OpenAIHost,
    setup_embeddings_service,
    setup_file_processors,
    setup_image_embeddings_service,
    setup_openai_client,
    setup_search_info,
)
from prepdocslib.blobmanager import AdlsBlobManager, BlobManager
from prepdocslib.embeddings import ImageEmbeddings
from prepdocslib.filestrategy import UploadUserFileStrategy
from prepdocslib.listfilestrategy import File

bp = Blueprint("routes", __name__, static_folder="static")
# Fix Windows registry issue with mimetypes
mimetypes.add_type("application/javascript", ".js")
mimetypes.add_type("text/css", ".css")


@bp.route("/")
async def index():
    return await bp.send_static_file("index.html")


# Dedicated MSAL popup-redirect page. Per MSAL best practice this must be a
# minimal page that only loads the redirect-bridge script — no routing, no
# other application code — so we serve a separate redirect.html static asset
# (not index.html). msal-browser 5.x uses a BroadcastChannel handshake and
# the bridge script inside redirect.html posts the auth response back to the
# opener window and closes the popup.
# See https://github.com/AzureAD/microsoft-authentication-library-for-js/blob/dev/lib/msal-browser/docs/login-user.md#redirecturi-considerations
@bp.route("/redirect")
async def redirect():
    return await bp.send_static_file("redirect.html")


@bp.route("/favicon.ico")
async def favicon():
    return await bp.send_static_file("favicon.ico")


@bp.route("/assets/<path:path>")
async def assets(path):
    return await send_from_directory(Path(__file__).resolve().parent / "static" / "assets", path)


@bp.route("/content/<path>")
@authenticated_path
async def content_file(path: str, auth_claims: dict[str, Any]):
    """
    Serve content files from blob storage from within the app to keep the example self-contained.
    *** NOTE *** if you are using app services authentication, this route will return unauthorized to all users that are not logged in
    if AZURE_ENFORCE_ACCESS_CONTROL is not set or false, logged in users can access all files regardless of access control
    if AZURE_ENFORCE_ACCESS_CONTROL is set to true, logged in users can only access files they have access to
    This is also slow and memory hungry.
    """
    # Remove page number from path, filename-1.txt -> filename.txt
    # This shouldn't typically be necessary as browsers don't send hash fragments to servers
    if path.find("#page=") > 0:
        path_parts = path.rsplit("#page=", 1)
        path = path_parts[0]
    current_app.logger.info("Opening file %s", path)
    blob_manager: BlobManager = current_app.config[CONFIG_GLOBAL_BLOB_MANAGER]

    # Get bytes and properties from the blob manager
    result = await blob_manager.download_blob(path)

    if result is None:
        current_app.logger.info("Path not found in general Blob container: %s", path)
        if current_app.config[CONFIG_USER_UPLOAD_ENABLED]:
            user_oid = auth_claims["oid"]
            user_blob_manager: AdlsBlobManager = current_app.config[CONFIG_USER_BLOB_MANAGER]
            result = await user_blob_manager.download_blob(path, user_oid=user_oid)
            if result is None:
                current_app.logger.exception("Path not found in DataLake: %s", path)

    if not result:
        abort(404)

    content, properties = result

    if not properties or "content_settings" not in properties:
        abort(404)

    mime_type = properties["content_settings"]["content_type"]
    if mime_type == "application/octet-stream":
        mime_type = mimetypes.guess_type(path)[0] or "application/octet-stream"

    # Create a BytesIO object from the bytes
    blob_file = io.BytesIO(content)
    return await send_file(blob_file, mimetype=mime_type, as_attachment=False, attachment_filename=path)


def _safe_asdict(o: Any) -> Any:
    """Like dataclasses.asdict() but without deepcopy, and skips non-serializable objects."""
    if dataclasses.is_dataclass(o) and not isinstance(o, type):
        return {f.name: _safe_asdict(getattr(o, f.name)) for f in dataclasses.fields(o)}
    elif isinstance(o, list):
        return [_safe_asdict(v) for v in o]
    elif isinstance(o, dict):
        return {k: _safe_asdict(v) for k, v in o.items()}
    elif isinstance(o, tuple):
        return tuple(_safe_asdict(v) for v in o)
    elif isinstance(o, (str, int, float, bool)) or o is None:
        return o
    else:
        # Skip non-primitive objects (e.g. AsyncOpenAI client, thread locks)
        return None


class JSONEncoder(json.JSONEncoder):
    def default(self, o):
        if dataclasses.is_dataclass(o) and not isinstance(o, type):
            as_dict = _safe_asdict(o)
            if isinstance(o, DataPoints):
                return {k: v for k, v in as_dict.items() if v is not None}
            data_points_payload = as_dict.get("data_points") if isinstance(as_dict, dict) else None
            if isinstance(data_points_payload, dict) and data_points_payload.get("citation_activity_details") is None:
                data_points_payload.pop("citation_activity_details")
            return as_dict
        return super().default(o)


async def format_as_ndjson(r: AsyncGenerator[dict, None]) -> AsyncGenerator[str, None]:
    try:
        async for event in r:
            yield json.dumps(event, ensure_ascii=False, cls=JSONEncoder) + "\n"
    except Exception as error:
        logging.exception("Exception while generating response stream: %s", error)
        yield json.dumps(error_dict(error))


def _inject_employee_scope(context: dict) -> None:
    """PROJECT EASE: If the caller is an employee, restrict search to their permitted category docs."""
    pe = _get_session()
    if pe and pe.get("role") == "employee":
        user_id = pe.get("user_id") or ""
        org_id  = pe.get("org") or ""
        cat_ids = get_permitted_categories(user_id)
        docs    = get_docs_for_categories(org_id, cat_ids) if cat_ids else []
        context.setdefault("overrides", {})["permitted_sourcefiles"] = [d["filename"] for d in docs]


@bp.route("/chat", methods=["POST"])
@authenticated
async def chat(auth_claims: dict[str, Any]):
    if not request.is_json:
        return jsonify({"error": "request must be json"}), 415
    request_json = await request.get_json()
    context = request_json.get("context", {})
    context["auth_claims"] = auth_claims
    _inject_employee_scope(context)  # PROJECT EASE: employee category scoping
    try:
        approach: Approach = cast(Approach, current_app.config[CONFIG_CHAT_APPROACH])

        # If session state is provided, persists the session state,
        # else creates a new session_id depending on the chat history options enabled.
        session_state = request_json.get("session_state")
        if session_state is None:
            session_state = create_session_id(
                current_app.config[CONFIG_CHAT_HISTORY_COSMOS_ENABLED],
                current_app.config[CONFIG_CHAT_HISTORY_BROWSER_ENABLED],
            )
        result = await approach.run(
            request_json["messages"],
            context=context,
            session_state=session_state,
        )
        # Log the search — extract last user message as query
        try:
            messages = request_json.get("messages", [])
            query = next((m["content"] for m in reversed(messages) if m.get("role") == "user"), None)
            _audit(_get_session(), "search", details={"query": query})
        except Exception:
            pass
        return jsonify(result)
    except Exception as error:
        return error_response(error, "/chat")


@bp.route("/chat/stream", methods=["POST"])
@authenticated
async def chat_stream(auth_claims: dict[str, Any]):
    if not request.is_json:
        return jsonify({"error": "request must be json"}), 415
    request_json = await request.get_json()
    context = request_json.get("context", {})
    context["auth_claims"] = auth_claims
    _inject_employee_scope(context)  # PROJECT EASE: employee category scoping
    try:
        approach: Approach = cast(Approach, current_app.config[CONFIG_CHAT_APPROACH])

        # If session state is provided, persists the session state,
        # else creates a new session_id depending on the chat history options enabled.
        session_state = request_json.get("session_state")
        if session_state is None:
            session_state = create_session_id(
                current_app.config[CONFIG_CHAT_HISTORY_COSMOS_ENABLED],
                current_app.config[CONFIG_CHAT_HISTORY_BROWSER_ENABLED],
            )
        result = await approach.run_stream(
            request_json["messages"],
            context=context,
            session_state=session_state,
        )
        response = await make_response(format_as_ndjson(result))
        response.timeout = None  # type: ignore
        response.mimetype = "application/json-lines"
        return response
    except Exception as error:
        return error_response(error, "/chat")


# Send MSAL.js settings to the client UI
@bp.route("/auth_setup", methods=["GET"])
def auth_setup():
    auth_helper = current_app.config.get(CONFIG_AUTH_CLIENT)
    if auth_helper is None:
        # Local-dev / auth-only mode (no Azure credentials configured) — mirrors
        # the frontend's own DEFAULT_AUTH_SETUP fallback so this never 500s.
        return jsonify({
            "useLogin": False,
            "requireAccessControl": False,
            "enableUnauthenticatedAccess": True,
            "msalConfig": {
                "auth": {"clientId": "", "authority": "", "redirectUri": "/redirect",
                         "postLogoutRedirectUri": "/", "navigateToLoginRequestUrl": False},
                "cache": {"cacheLocation": "sessionStorage", "storeAuthStateInCookie": False},
            },
            "loginRequest": {"scopes": []},
            "tokenRequest": {"scopes": []},
        })
    return jsonify(auth_helper.get_auth_setup_for_client())


# ─── PROJECT EASE: DB-backed auth ────────────────────────────────────────────
import secrets as _secrets
from db import (
    SYSTEM,
    init_db, get_user_by_email, check_password as _check_pw,
    get_org, get_categories, create_category, delete_category,
    get_documents, create_document, update_document_status, delete_document,
    get_doc_counts, get_users_for_org, create_user, delete_user,
    get_permitted_categories, set_permissions, get_docs_for_categories,
    get_user_by_id, get_user_by_whatsapp, update_user_whatsapp,
    search_matters_by_keyword, get_filenames_for_matter,
    create_reset_token, use_reset_token,
    # admin
    get_all_orgs, create_org, update_org, delete_org,
    get_org_details, get_platform_stats,
    # registration — Task #41
    register_org, get_pending_registrations, approve_registration, update_org_profile,
    # Matter/Client Management — Task #31
    get_clients, create_client, update_client, delete_client, get_client_with_matters,
    get_matter_teams, create_matter_team, update_matter_team, delete_matter_team,
    add_matter_team_member, remove_matter_team_member,
    get_custom_courts, add_custom_court, delete_custom_court,
    get_matters, create_matter, update_matter, delete_matter,
    get_matter_with_docs, link_document_to_matter, unlink_document_from_matter,
    # Audit Log — Task #30
    log_event, get_audit_logs, count_audit_logs,
    # Upgrade flow — Task #45
    PLAN_CONFIG, create_upgrade_request, get_upgrade_requests, resolve_upgrade_request,
    # Court Calendar — Task #32
    get_hearings, create_hearing, get_hearing, update_hearing, delete_hearing,
    get_deadlines, create_deadline, get_deadline, update_deadline, delete_deadline,
    get_hearings_needing_reminder, mark_hearing_reminder_sent,
    get_deadlines_needing_reminder, mark_deadline_reminder_sent,
    # Fee tracking & Invoices — Task #44
    get_fees, create_fee, get_fee, update_fee, delete_fee,
    get_invoices, create_invoice, get_invoice_with_fees, update_invoice,
    # Case Law — Task #33
    create_case_law_doc, get_case_law_doc, list_case_law_docs,
    set_case_law_doc_status, delete_case_law_doc,
    # Templates & Drafting — Task #38
    list_templates, create_template, get_template, update_template, delete_template,
    TEMPLATE_TYPES,
    # Client Portal — Task #39
    create_client_token, get_client_token_by_id, get_client_token_by_value,
    list_client_tokens, revoke_client_token,
    # Court Orders — Task #130
    get_court_orders, create_court_order, update_court_order, delete_court_order,
    get_matter_client_contact, find_client_by_phone_suffix, get_client_matters_status,
    get_clients_with_hearings_in_range,
    # Adverse Parties — Task #131
    get_adverse_parties, create_adverse_party, update_adverse_party, delete_adverse_party,
    # Limitation Tracker — Task #132
    LIMITATION_PERIODS, compute_limitation_date, get_matters_with_approaching_limitation,
    # Time Tracking — Task #133
    get_time_entries, create_time_entry, update_time_entry, delete_time_entry, bill_time_entries,
    # Cause List — Task #137
    parse_cause_list_text, store_cause_list, get_cause_list_entries,
    link_cause_list_entry, delete_cause_list_entry, get_today_cause_list_matches,
    get_orgs_with_owner_wa, get_org_owner_contact,
    # Matter Notes — Task #138
    NOTE_TYPES, get_matter_notes, create_matter_note, update_matter_note, delete_matter_note,
    # Matter Priority — Task #139
    MATTER_PRIORITIES,
    # Document Requests — Task #140
    DOC_REQUEST_STATUSES, get_document_requests, create_document_request,
    update_document_request, delete_document_request,
    # Witnesses — Task #141
    WITNESS_TYPES, STATEMENT_STATUSES, get_witnesses, create_witness,
    update_witness, delete_witness,
    # Matter Deadlines — Task #142
    DEADLINE_PRIORITIES, get_matter_deadlines, create_matter_deadline,
    update_matter_deadline, delete_matter_deadline,
    # Matter Expenses — Task #143
    EXPENSE_CATEGORIES, get_matter_expenses, create_matter_expense,
    update_matter_expense, delete_matter_expense,
    # Matter Correspondence — Task #144
    CORR_DIRECTIONS, CORR_TYPES, get_matter_correspondence, create_matter_correspondence,
    update_matter_correspondence, delete_matter_correspondence,
    # Bail & Interim Relief — Task #145
    RELIEF_TYPES, RELIEF_STATUSES, get_matter_relief, create_matter_relief,
    update_matter_relief, delete_matter_relief,
    # Matter Outcome — Task #146
    OUTCOME_TYPES, get_matter_outcome, upsert_matter_outcome,
    # Matter Charges — Task #147
    PLEA_OPTIONS, get_matter_charges, create_matter_charge,
    update_matter_charge, delete_matter_charge,
    # Matter FIR — Task #148
    get_matter_fir, create_matter_fir, update_matter_fir, delete_matter_fir,
    # Matter Challan — Task #149
    CHALLAN_TYPES, CHALLAN_STATUSES,
    get_matter_challan, create_matter_challan, update_matter_challan, delete_matter_challan,
    # Court Fee Calculator — Task #152
    COURT_FEE_TYPES, compute_court_fee,
    get_court_fee_payments, create_court_fee_payment,
    update_court_fee_payment, delete_court_fee_payment,
    # Associate Fees — Task #153
    get_associate_fees, get_associate_fees_summary,
    create_associate_fee, update_associate_fee, delete_associate_fee,
    # Trust Ledger — Task #154
    TRUST_TXN_TYPES,
    get_trust_ledger, create_trust_entry, update_trust_entry, delete_trust_entry,
    # Cheque Tracker — Task #155
    CHEQUE_TYPES, CHEQUE_STATUSES,
    get_matter_cheques, create_matter_cheque, update_matter_cheque, delete_matter_cheque,
    # WHT Invoice — Task #157
    compute_wht,
    # Intelligence Notes — Task #158
    get_opposing_counsel, create_opposing_counsel, update_opposing_counsel, delete_opposing_counsel,
    get_judge_notes, create_judge_note, update_judge_note, delete_judge_note,
    get_judge_note, get_judge_track_record,
    # Feature Flags — Task #162
    FEATURE_KEYS, FEATURE_LABELS, get_org_flags, set_org_flags, get_all_org_flags, is_feature_enabled,
    # Legal Notices — Task #165
    NOTICE_TYPES, NOTICE_VIA, NOTICE_STATUSES,
    get_legal_notices, create_legal_notice, update_legal_notice, delete_legal_notice,
    # Court Transfers — Task #170
    get_court_transfers, create_court_transfer, update_court_transfer, delete_court_transfer,
    # Bail Bonds — Task #167
    BAIL_TYPES, BAIL_STATUSES,
    get_bail_bonds, create_bail_bond, update_bail_bond, delete_bail_bond,
    get_org_bail_stages, add_org_bail_stage, update_org_bail_stage,
    get_bail_stage_completions, set_bail_stage_completion,
    # Staff Attendance & Salary — Task #171
    STAFF_ROLES, ATT_STATUSES, SALARY_PAY_MODES,
    get_staff, create_staff, update_staff, delete_staff,
    get_attendance, upsert_attendance,
    get_salary_payments, create_salary_payment, delete_salary_payment,
    # Outstanding Dues — Task #169
    get_outstanding_invoices,
)
# Note: Conflict check (Task #150) reuses get_clients + get_matters already imported above.

# Initialise DB (creates tables + seeds dev data) at import time
init_db()

_sessions: dict[str, dict] = {}  # token → session dict  (in-memory; fine for MVP)


def _get_session(req=None) -> dict | None:
    r = req or request
    token = r.headers.get("Authorization", "").removeprefix("Bearer ").strip()
    return _sessions.get(token)


def _get_ip() -> str | None:
    """Best-effort caller IP from X-Forwarded-For or remote_addr."""
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.remote_addr or None


def _audit(session: dict | None, event_type: str, **kwargs):
    """Convenience wrapper — pulls org/user/actor fields from session."""
    log_event(
        event_type  = event_type,
        org_id      = (session or {}).get("org"),
        user_id     = (session or {}).get("user_id"),
        actor_name  = (session or {}).get("name"),
        actor_role  = (session or {}).get("role"),
        ip_address  = _get_ip(),
        **kwargs,
    )


@bp.route("/auth/login", methods=["POST"])
async def auth_login():
    """Check email + password against DB and return a session token."""
    data = await request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not email or not password:
        return jsonify({"error": "Email and password are required"}), 400

    # Platform admin — hardcoded (no org)
    if email == "admin@projectease.com" and password == "admin123":
        token = _secrets.token_hex(32)
        session_data = {"email": email, "name": "Platform Admin", "role": "platform_admin", "org": None}
        _sessions[token] = session_data
        log_event("login_success", actor_name=email, actor_role="platform_admin", ip_address=_get_ip(),
                  details={"email": email})
        return jsonify({"token": token, "user": session_data})

    user = get_user_by_email(email)
    if not user or not _check_pw(password, user["password_hash"]):
        log_event("login_fail", ip_address=_get_ip(), details={"email": email})
        return jsonify({"error": "Invalid email or password"}), 401

    # Self-registered orgs start life as pending_payment (Task #41) — block login
    # until an admin approves the registration. Whitelisted on "active" (rather
    # than blacklisting "pending_payment") so any future non-active status is
    # blocked by default too, not silently let through.
    org = get_org(user["org_id"])
    if org and org.get("status") not in (None, "active"):
        log_event("login_fail", org_id=user["org_id"], user_id=user["user_id"],
                  ip_address=_get_ip(), details={"email": email, "reason": "org_not_active"})
        return jsonify({"error": "Your organization's registration is still pending approval. "
                                  "You'll be notified by email once it's been reviewed."}), 403

    token = _secrets.token_hex(32)
    session_data = {
        "user_id": user["user_id"],
        "email":   user["email"],
        "name":    user["name"],
        "role":    user["role"],
        "org":     user["org_id"],
    }
    _sessions[token] = session_data
    log_event("login_success",
              org_id=user["org_id"], user_id=user["user_id"],
              actor_name=user["name"], actor_role=user["role"],
              ip_address=_get_ip(), details={"email": email})
    return jsonify({
        "token": token,
        "user": {
            **session_data,
            "must_change_password": bool(user.get("must_change_password", 0)),
        },
    })


@bp.route("/auth/me", methods=["GET"])
async def auth_me():
    """Return the current user for a valid session token."""
    token = request.headers.get("Authorization", "").removeprefix("Bearer ").strip()
    session = _sessions.get(token)
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify(session)


@bp.route("/auth/logout", methods=["POST"])
async def auth_logout():
    """Invalidate a session token."""
    token = request.headers.get("Authorization", "").removeprefix("Bearer ").strip()
    session = _sessions.pop(token, None)
    _audit(session, "logout")
    return jsonify({"success": True})


# ─── PROJECT EASE: Document Upload ───────────────────────────────────────────
# Accepts any supported file type, tags with org_id, and indexes into AI Search.
# Supported: PDF (text + scanned), DOCX, PPTX, XLSX, PNG, JPG, TXT, CSV, MD
@bp.route("/upload", methods=["POST"])
async def upload_document():
    """Upload a document and index it into Azure AI Search for the caller's org."""
    # Auth check
    token = request.headers.get("Authorization", "").removeprefix("Bearer ").strip()
    session = _sessions.get(token)
    if not session:
        return jsonify({"error": "Unauthorized"}), 401

    org_id: str | None = session.get("org")
    user_id: str | None = session.get("user_id")

    # ── Plan limit check ──────────────────────────────────────────────────────
    if org_id:
        org = get_org(org_id)
        if org:
            plan_key = org.get("plan", "trial")
            limits   = PLAN_CONFIG.get(plan_key, PLAN_CONFIG["trial"])
            counts   = get_doc_counts(org_id)
            if counts["total_docs"] >= limits["max_docs"]:
                return jsonify({
                    "error": f"Document limit reached ({limits['max_docs']} docs on {limits['label']} plan). "
                             "Please upgrade your plan to upload more documents.",
                    "limit_reached": "docs",
                }), 403
    # ─────────────────────────────────────────────────────────────────────────

    files = await request.files
    form  = await request.form
    if "file" not in files:
        return jsonify({"error": "No file provided"}), 400

    uploaded_file = files["file"]
    filename: str = uploaded_file.filename or "upload"
    category_id: str | None = form.get("category_id") or None

    # Allowed extensions
    ALLOWED_EXTENSIONS = {
        ".pdf", ".docx", ".pptx", ".xlsx",
        ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
        ".txt", ".csv", ".md", ".json", ".html",
    }
    import os as _os
    ext = _os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type '{ext}' is not supported."}), 400

    try:
        import tempfile, asyncio as _asyncio
        from dotenv import load_dotenv as _ld
        _ld(dotenv_path=_os.path.join(_os.path.dirname(__file__), ".env"), override=True)

        from azure.core.credentials import AzureKeyCredential as _AKC
        from azure.core.credentials_async import AsyncTokenCredential
        from azure.identity.aio import AzureDeveloperCliCredential as _AzCred
        from prepdocslib.servicesetup import (
            OpenAIHost, setup_search_info, setup_blob_manager,
            setup_embeddings_service, setup_openai_client,
            build_file_processors, setup_figure_processor,
        )
        from prepdocslib.listfilestrategy import LocalListFileStrategy
        from prepdocslib.filestrategy import FileStrategy
        from prepdocslib.strategy import DocumentAction

        # ── credentials ──
        azure_credential = _AzCred(process_timeout=60)
        OPENAI_HOST = OpenAIHost(_os.environ["OPENAI_HOST"])

        search_key = _os.getenv("AZURE_SEARCH_KEY")
        search_cred = _AKC(search_key) if search_key else azure_credential

        search_info = setup_search_info(
            search_service=_os.environ["AZURE_SEARCH_SERVICE"],
            index_name=_os.environ["AZURE_SEARCH_INDEX"],
            azure_credential=azure_credential,
            search_key=search_key,
        )

        blob_manager = setup_blob_manager(
            azure_credential=azure_credential,
            storage_account=_os.environ["AZURE_STORAGE_ACCOUNT"],
            storage_container=_os.environ["AZURE_STORAGE_CONTAINER"],
            storage_resource_group=_os.environ.get("AZURE_STORAGE_RESOURCE_GROUP"),
            subscription_id=_os.environ.get("AZURE_SUBSCRIPTION_ID"),
            storage_key=_os.getenv("AZURE_STORAGE_KEY"),
        )

        openai_client, azure_openai_endpoint = setup_openai_client(
            openai_host=OPENAI_HOST,
            azure_credential=azure_credential,
            azure_openai_service=_os.getenv("AZURE_OPENAI_SERVICE"),
            azure_openai_api_key=_os.getenv("AZURE_OPENAI_API_KEY_OVERRIDE"),
        )

        doc_int_service = _os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE")
        doc_int_key = _os.getenv("AZURE_DOCUMENTINTELLIGENCE_KEY")

        file_processors, figure_processor = build_file_processors(
            azure_credential=azure_credential,
            document_intelligence_service=doc_int_service,
            document_intelligence_key=doc_int_key if doc_int_key and doc_int_key != "YOUR_DOC_INTELLIGENCE_KEY_HERE" else None,
            use_local_pdf_parser=not doc_int_service,
            use_local_html_parser=not doc_int_service,
        ), None

        # Unpack tuple from build_file_processors (returns dict, not tuple)
        if isinstance(file_processors, tuple):
            file_processors, figure_processor = file_processors

        embeddings_service = setup_embeddings_service(
            OPENAI_HOST,
            openai_client,
            emb_model_name=_os.environ["AZURE_OPENAI_EMB_MODEL_NAME"],
            emb_model_dimensions=int(_os.getenv("AZURE_OPENAI_EMB_DIMENSIONS", "1536")),
            azure_openai_deployment=_os.getenv("AZURE_OPENAI_EMB_DEPLOYMENT"),
            azure_openai_endpoint=azure_openai_endpoint,
        )

        # ── save upload to a temp file ──
        file_bytes = uploaded_file.read()
        size_bytes = len(file_bytes)
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        # Create DB record immediately (status=processing)
        doc_record = create_document(
            org_id=org_id or "",
            filename=filename,
            actor=user_id or SYSTEM,
            size_bytes=size_bytes,
            uploaded_by=user_id,
            category_id=category_id,
        )

        # ── Tag document with org_id via category field ──
        list_strategy = LocalListFileStrategy(path_pattern=tmp_path)

        strategy = FileStrategy(
            search_info=search_info,
            list_file_strategy=list_strategy,
            blob_manager=blob_manager,
            file_processors=file_processors,
            document_action=DocumentAction.Add,
            embeddings=embeddings_service,
            image_embeddings=None,
            search_field_name_embedding=_os.getenv("AZURE_SEARCH_FIELD_NAME_EMBEDDING", "embedding"),
            use_acls=False,
            category=org_id,       # ← org_id stored as category for tenant filtering
            figure_processor=figure_processor,
        )

        # ── Delete the md5 cache so the file always gets processed ──
        md5_path = tmp_path + ".md5"
        if _os.path.exists(md5_path):
            _os.remove(md5_path)

        loop = _asyncio.get_event_loop()
        await strategy.setup()
        await strategy.run()

        # cleanup
        _os.remove(tmp_path)
        if _os.path.exists(md5_path):
            _os.remove(md5_path)

        try:
            await blob_manager.close_clients()
            await openai_client.close()
            await azure_credential.close()
        except Exception:
            pass

        update_document_status(doc_record["doc_id"], "ready", actor=user_id or SYSTEM)
        _audit(_get_session(), "doc_upload",
               resource_type="document", resource_id=doc_record["doc_id"],
               resource_name=filename,
               details={"size_bytes": doc_record.get("size_bytes"), "category_id": category_id})
        return jsonify({"success": True, "filename": filename, "org": org_id, "doc": doc_record})

    except Exception as e:
        current_app.logger.exception("Upload failed: %s", e)
        return jsonify({"error": str(e)}), 500


# ─── PROJECT EASE: Documents API ─────────────────────────────────────────────

@bp.route("/me", methods=["GET"])
async def get_my_profile():
    """Return current user's profile and their permitted categories."""
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    org_id  = session.get("org") or ""
    user_id = session.get("user_id") or ""
    org     = get_org(org_id) or {}
    cat_ids = set(get_permitted_categories(user_id))
    all_cats = get_categories(org_id)
    permitted = [c for c in all_cats if c["category_id"] in cat_ids]
    return jsonify({
        "user_id":              user_id,
        "name":                 session.get("name", ""),
        "email":                session.get("email", ""),
        "role":                 session.get("role", ""),
        "org_name":             org.get("name", ""),
        "permitted_categories": permitted,
    })


@bp.route("/documents", methods=["GET"])
async def list_documents():
    """Return documents for the caller's org. Employees see only permitted-category docs."""
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""

    if session.get("role") == "employee":
        user_id = session.get("user_id") or ""
        cat_ids = get_permitted_categories(user_id)
        docs    = get_docs_for_categories(org_id, cat_ids) if cat_ids else []
        total_bytes = sum((d.get("size_bytes") or 0) for d in docs)
        return jsonify({"documents": docs, "usage": {"total_docs": len(docs), "total_bytes": total_bytes}})

    docs   = get_documents(org_id)
    counts = get_doc_counts(org_id)
    return jsonify({"documents": docs, "usage": counts})


@bp.route("/documents/<doc_id>", methods=["DELETE"])
async def remove_document(doc_id: str):
    """Delete a document from the index, blob storage, and DB."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org")

    doc = delete_document(doc_id, org_id or "", actor=session.get("user_id") or SYSTEM)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    _audit(session, "doc_delete",
           resource_type="document", resource_id=doc_id,
           resource_name=doc.get("filename"))

    # Remove from Azure AI Search index
    try:
        import os as _os
        from dotenv import load_dotenv as _ld
        _ld(dotenv_path=_os.path.join(_os.path.dirname(__file__), ".env"), override=True)

        from azure.core.credentials import AzureKeyCredential as _AKC
        from azure.identity.aio import AzureDeveloperCliCredential as _AzCred
        from prepdocslib.servicesetup import setup_search_info, setup_blob_manager
        from prepdocslib.listfilestrategy import LocalListFileStrategy
        from prepdocslib.filestrategy import FileStrategy
        from prepdocslib.strategy import DocumentAction

        azure_credential = _AzCred(process_timeout=60)
        search_key = _os.getenv("AZURE_SEARCH_KEY")

        search_info = setup_search_info(
            search_service=_os.environ["AZURE_SEARCH_SERVICE"],
            index_name=_os.environ["AZURE_SEARCH_INDEX"],
            azure_credential=azure_credential,
            search_key=search_key,
        )
        blob_manager = setup_blob_manager(
            azure_credential=azure_credential,
            storage_account=_os.environ["AZURE_STORAGE_ACCOUNT"],
            storage_container=_os.environ["AZURE_STORAGE_CONTAINER"],
            storage_resource_group=_os.environ.get("AZURE_STORAGE_RESOURCE_GROUP"),
            subscription_id=_os.environ.get("AZURE_SUBSCRIPTION_ID"),
            storage_key=_os.getenv("AZURE_STORAGE_KEY"),
        )

        # FileStrategy with Remove action uses the filename to delete chunks from the index
        import tempfile, os as _os2
        # Create a dummy temp file so LocalListFileStrategy can resolve the path
        with tempfile.NamedTemporaryFile(delete=False, suffix=_os.path.splitext(doc["filename"])[1]) as tmp:
            tmp_path = tmp.name

        list_strategy = LocalListFileStrategy(path_pattern=tmp_path)
        # Rename temp file to match original filename for correct index key lookup
        target_path = _os2.path.join(_os2.path.dirname(tmp_path), doc["filename"])
        _os2.rename(tmp_path, target_path)

        strategy = FileStrategy(
            search_info=search_info,
            list_file_strategy=LocalListFileStrategy(path_pattern=target_path),
            blob_manager=blob_manager,
            file_processors={},
            document_action=DocumentAction.Remove,
            embeddings=None,
            image_embeddings=None,
            use_acls=False,
            category=org_id,
        )
        await strategy.setup()
        await strategy.run()
        _os2.remove(target_path)

        try:
            await blob_manager.close_clients()
            await azure_credential.close()
        except Exception:
            pass
    except Exception as e:
        current_app.logger.warning("Index removal failed (doc still deleted from DB): %s", e)

    return jsonify({"success": True, "doc_id": doc_id})


# ─── PROJECT EASE: Categories API ────────────────────────────────────────────

@bp.route("/categories", methods=["GET"])
async def list_categories():
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"categories": get_categories(session.get("org") or "")})


@bp.route("/categories", methods=["POST"])
async def add_category():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Category name is required"}), 400
    try:
        cat = create_category(session.get("org") or "", name, actor=session.get("user_id") or SYSTEM)
        return jsonify(cat), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@bp.route("/categories/<category_id>", methods=["DELETE"])
async def remove_category(category_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_category(category_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── PROJECT EASE: Org self-service API ──────────────────────────────────────

@bp.route("/org", methods=["GET"])
async def get_own_org():
    """Return the caller's own org record."""
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    org = get_org(session.get("org") or "")
    if not org:
        return jsonify({"error": "Org not found"}), 404
    counts = get_doc_counts(session.get("org") or "")
    members = get_users_for_org(session.get("org") or "")
    return jsonify({**org, **counts, "user_count": len(members)})


@bp.route("/org", methods=["PUT"])
async def update_own_org():
    """Let an org owner update their org name / industry."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    updated = update_org(
        session.get("org") or "",
        actor=actor,
        name=data.get("name"),
        industry=data.get("industry"),
    )
    _audit(session, "org_update", details={"name": data.get("name"), "industry": data.get("industry")})
    return jsonify(updated)


# ─── PROJECT EASE: Self-service Registration ─────────────────────────────────

@bp.route("/register", methods=["POST"])
async def public_register():
    """Public endpoint — no auth. Creates org (pending_payment) + owner user."""
    import asyncio as _asyncio
    from email_helper import send_registration_pending as _send_pending

    data        = await request.get_json(silent=True) or {}
    firm_name   = (data.get("firm_name")   or "").strip()
    owner_name  = (data.get("owner_name")  or "").strip()
    owner_email = (data.get("owner_email") or "").strip().lower()
    password    = data.get("password", "")
    city        = (data.get("city")  or "").strip()
    phone       = (data.get("phone") or "").strip()
    plan        = data.get("plan", "pro")

    if not firm_name or not owner_name or not owner_email or not password:
        return jsonify({"error": "Firm name, your name, email, and password are required."}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400
    if get_user_by_email(owner_email):
        return jsonify({"error": "An account with this email already exists."}), 409

    try:
        org = register_org(
            firm_name=firm_name,
            owner_name=owner_name,
            owner_email=owner_email,
            password=password,
            city=city,
            phone=phone,
            plan=plan,
        )
    except Exception as exc:
        current_app.logger.exception("Registration error: %s", exc)
        return jsonify({"error": "Registration failed. Please try again."}), 500

    # Fire-and-forget confirmation email
    try:
        await _asyncio.to_thread(_send_pending, owner_email, firm_name)
    except Exception:
        pass  # email failure must not block signup

    return jsonify({"success": True, "org_id": org["org_id"]}), 201


@bp.route("/admin/registrations", methods=["GET"])
async def admin_get_registrations():
    """List all orgs with pending_payment status."""
    err = _require_platform_admin()
    if err:
        return err
    return jsonify({"registrations": get_pending_registrations()})


@bp.route("/admin/orgs/<org_id>/approve", methods=["PATCH"])
async def admin_approve_org(org_id: str):
    """Approve a pending registration → active + send email."""
    import asyncio as _asyncio
    from email_helper import send_registration_approved as _send_approved

    err = _require_platform_admin()
    if err:
        return err
    admin_session = _get_session()
    actor = (admin_session or {}).get("email") or SYSTEM

    org = approve_registration(org_id, actor=actor)
    if not org:
        return jsonify({"error": "Organization not found or already active."}), 404

    # Send approval email to org owner
    details = get_org_details(org_id)
    if details and details.get("users"):
        owner = next(
            (u for u in details["users"] if u.get("role") == "org_owner"), None
        )
        if owner:
            try:
                await _asyncio.to_thread(
                    _send_approved, owner["email"], org["name"]
                )
            except Exception:
                pass

    return jsonify({"success": True, "org": org})


@bp.route("/org/profile", methods=["PUT"])
async def update_org_profile_route():
    """Owner-only: update optional profile fields (phone, city, practice_areas, etc.)."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    updated = update_org_profile(
        session.get("org") or "",
        actor=actor,
        phone=data.get("phone"),
        city=data.get("city"),
        practice_areas=data.get("practice_areas"),
        bar_council_no=data.get("bar_council_no"),
        website=data.get("website"),
        team_size=data.get("team_size"),
    )
    return jsonify(updated or {})


@bp.route("/auth/change-password", methods=["POST"])
async def change_password():
    """Verify current password then set a new one."""
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    current_pw  = data.get("current_password", "")
    new_pw      = data.get("new_password", "")
    if not current_pw or not new_pw:
        return jsonify({"error": "Both current and new password are required"}), 400
    if len(new_pw) < 8:
        return jsonify({"error": "New password must be at least 8 characters"}), 400

    user = get_user_by_email(session.get("email") or "")
    if not user or not _check_pw(current_pw, user["password_hash"]):
        return jsonify({"error": "Current password is incorrect"}), 401

    from db import hash_password as _hash_pw
    actor = session.get("user_id") or SYSTEM
    with __import__("db").get_conn() as conn:
        conn.execute(
            """UPDATE users SET password_hash=?, must_change_password=0,
               modified_at=datetime('now'), modified_by=? WHERE user_id=?""",
            (_hash_pw(new_pw), actor, user["user_id"]),
        )
    _audit(session, "password_change", resource_type="user", resource_id=user["user_id"])
    return jsonify({"success": True})


@bp.route("/auth/forgot-password", methods=["POST"])
async def forgot_password():
    """Unauthenticated — generate a reset token and email a link.
    Always returns 200 (don't reveal whether email exists)."""
    import asyncio as _aio
    from email_helper import send_password_reset as _send_reset
    data  = await request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    if email:
        user = get_user_by_email(email)
        if user and user.get("is_active"):
            token    = create_reset_token(user["user_id"])
            base_url = (data.get("base_url") or "https://projectease.pk").rstrip("/")
            reset_url = f"{base_url}/#/?reset_token={token}"
            await _aio.to_thread(
                _send_reset, email, user.get("name", "there"), reset_url
            )
    return jsonify({"ok": True})


@bp.route("/auth/reset-password", methods=["POST"])
async def reset_password():
    """Unauthenticated — validate reset token and set new password."""
    from db import hash_password as _hash_pw
    data     = await request.get_json(silent=True) or {}
    token    = (data.get("token") or "").strip()
    new_pw   = (data.get("password") or "").strip()
    if not token or len(new_pw) < 8:
        return jsonify({"error": "Token and password (min 8 chars) are required."}), 400
    user_id = use_reset_token(token)
    if not user_id:
        return jsonify({"error": "This reset link is invalid or has expired."}), 400
    with __import__("db").get_conn() as conn:
        conn.execute(
            "UPDATE users SET password_hash=?, must_change_password=0, modified_at=datetime('now'), modified_by=? WHERE user_id=?",
            (_hash_pw(new_pw), "password_reset", user_id),
        )
    log_event("password_reset", ip_address=_get_ip(),
              details={"user_id": user_id, "method": "email_token"})
    return jsonify({"ok": True})


# ─── PROJECT EASE: Team API ───────────────────────────────────────────────────

@bp.route("/team", methods=["GET"])
async def list_team():
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    members = get_users_for_org(session.get("org") or "")
    return jsonify({"members": members})


@bp.route("/team", methods=["POST"])
async def invite_member():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    name  = (data.get("name") or "").strip()
    role  = data.get("role", "employee")
    if not email or not name:
        return jsonify({"error": "Name and email are required"}), 400

    # ── Plan limit check ──────────────────────────────────────────────────────
    org_id_check = session.get("org") or ""
    org_check    = get_org(org_id_check)
    if org_check:
        plan_key = org_check.get("plan", "trial")
        limits   = PLAN_CONFIG.get(plan_key, PLAN_CONFIG["trial"])
        current_users = len(get_users_for_org(org_id_check))
        if current_users >= limits["max_users"]:
            return jsonify({
                "error": f"User limit reached ({limits['max_users']} users on {limits['label']} plan). "
                         "Please upgrade your plan to add more team members.",
                "limit_reached": "users",
            }), 403
    # ─────────────────────────────────────────────────────────────────────────

    # Generate a temp password and email the invite
    import asyncio as _aio
    from email_helper import send_team_invite as _send_invite
    temp_pw = _secrets.token_urlsafe(8)
    try:
        user = create_user(
            org_id=session.get("org") or "",
            email=email, name=name, role=role,
            actor=session.get("user_id") or SYSTEM,
            password=temp_pw, must_change=True,
        )
        _audit(session, "member_invite",
               resource_type="user", resource_id=user["user_id"],
               resource_name=name, details={"email": email, "role": role})
        # Send invite email (fire-and-forget, non-blocking)
        org = get_org(session.get("org") or "")
        firm_name = org.get("name", "your firm") if org else "your firm"
        _aio.get_event_loop().run_in_executor(
            None, _send_invite, email, firm_name, temp_pw
        )
        return jsonify({**user, "temp_password": temp_pw}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@bp.route("/team/<user_id>", methods=["DELETE"])
async def remove_member(user_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    # Invalidate any active sessions for this user
    removed_user = get_user_by_id(user_id)
    for tok, s in list(_sessions.items()):
        if s.get("user_id") == user_id:
            _sessions.pop(tok, None)
    delete_user(user_id, actor=session.get("user_id") or SYSTEM)
    _audit(session, "member_remove",
           resource_type="user", resource_id=user_id,
           resource_name=(removed_user or {}).get("name"))
    return jsonify({"success": True})


@bp.route("/team/<user_id>/permissions", methods=["GET"])
async def get_member_permissions(user_id: str):
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"category_ids": get_permitted_categories(user_id)})


@bp.route("/team/<user_id>/permissions", methods=["PUT"])
async def set_member_permissions(user_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    category_ids = data.get("category_ids", [])
    set_permissions(user_id, category_ids, actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


@bp.route("/team/<user_id>/whatsapp", methods=["PATCH"])
async def update_member_whatsapp(user_id: str):
    """Set or clear the WhatsApp number for a team member (owner only)."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data    = await request.get_json(silent=True) or {}
    number  = (data.get("whatsapp_number") or "").strip()
    if number and not number.startswith("+"):
        number = "+" + number
    org_id  = session.get("org") or ""
    member  = get_user_by_id(user_id)
    if not member or member.get("org_id") != org_id:
        return jsonify({"error": "User not found"}), 404
    update_user_whatsapp(user_id, number or None, session.get("user_id") or SYSTEM)
    return jsonify({"success": True, "whatsapp_number": number or None})


# ─── PROJECT EASE: WhatsApp Webhook ──────────────────────────────────────────

# In-memory conversation history keyed by phone number (survives restarts poorly
# but is sufficient for MVP — a Redis cache can replace this later).
_wa_sessions: dict[str, list] = {}

# Multi-step self-registration state machine keyed by phone number.
# Each entry: { "step": int, "data": { "firm": str, "name": str, "email": str, "password": str } }
_wa_reg_sessions: dict[str, dict] = {}

_WA_HELP = (
    "📋 *Project Ease — Commands*\n"
    "• Send any legal question to search your firm's documents\n"
    "• *MATTER: <name>* — scope search to a specific matter\n"
    "  e.g. _MATTER: Khan vs State_ then your question\n"
    "• *LIST MATTERS* — see your active matters\n"
    "• *HELP* — show this message\n"
    "• *RESET* — clear conversation history"
)


def _twiml_reply(text: str):
    """Return a Twilio TwiML XML response containing a WhatsApp message."""
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    xml  = f'<?xml version="1.0" encoding="UTF-8"?><Response><Message>{safe}</Message></Response>'
    return current_app.response_class(xml, mimetype="text/xml")


async def _whisper_transcribe_bytes(audio_bytes: bytes, filename: str, content_type: str) -> str | None:
    """Transcribe raw audio bytes with Azure OpenAI Whisper.
    Returns None if no Whisper deployment is configured or if transcription fails.
    Whisper auto-detects language (Urdu, English, Roman Urdu speech, etc.) —
    we deliberately don't pin `language=` so mixed-language court dictation
    (very common: Urdu sentence, English case citation) still transcribes."""
    whisper_dep = os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT", "").strip()
    if not whisper_dep:
        return None
    try:
        openai_client = current_app.config[CONFIG_OPENAI_CLIENT]
        result = await openai_client.audio.transcriptions.create(
            model=whisper_dep,
            file=(filename, audio_bytes, content_type),
        )
        return result.text.strip() or None
    except Exception as exc:
        logging.warning("Whisper transcription failed: %s", exc)
        return None


async def _transcribe_voice(media_url: str) -> str | None:
    """Download a Twilio voice note and transcribe it with Azure OpenAI Whisper.
    Returns None if no Whisper deployment is configured or if transcription fails."""
    whisper_dep = os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT", "").strip()
    if not whisper_dep:
        return None
    try:
        import httpx
        sid   = os.getenv("TWILIO_ACCOUNT_SID", "")
        token = os.getenv("TWILIO_AUTH_TOKEN", "")
        auth  = (sid, token) if sid and token else None
        async with httpx.AsyncClient() as client:
            r = await client.get(media_url, auth=auth, timeout=30.0)
            r.raise_for_status()
            audio_bytes = r.content
            content_type = r.headers.get("content-type", "audio/ogg")
        return await _whisper_transcribe_bytes(audio_bytes, "voice_note.ogg", content_type)
    except Exception as exc:
        logging.warning("Whisper transcription failed: %s", exc)
        return None


_WA_CLIENT_TRIGGERS = {"MERA CASE", "MY CASE", "CASE STATUS", "STATUS", "MERA CASE?", "MY CASE?"}


async def _wa_client_self_service(phone: str, body: str):
    """Client self-service bot: a client (not a team member) texts a phone
    number that's on file for them, e.g. "mera case", and gets their next
    hearing date + last outcome — zero involvement from the firm.

    Deliberately a pure DB lookup, never routed through the AI/RAG chat
    pipeline — a client's own number must never be able to pull firm
    documents or other clients' data, so this only ever reads the handful
    of fields matched to *their own* client_id.

    Returns None if this number doesn't match any client on file (falls
    through to the existing team-member-onboarding flow), so the caller can
    tell "not a client" apart from "client asked something we handle here".
    """
    try:
        from whatsapp_helper import normalize_pk_number
    except ImportError:
        return None
    normalized = normalize_pk_number(phone)
    if not normalized:
        return None
    last10 = normalized[-10:]
    candidate = find_client_by_phone_suffix(last10)
    if not candidate or normalize_pk_number(candidate.get("phone") or "") != normalized:
        return None

    if body.strip().upper() not in _WA_CLIENT_TRIGGERS:
        return _twiml_reply(
            f"👋 Hi {candidate.get('name') or 'there'}, this is *Project Ease* on behalf of your lawyer.\n\n"
            "Reply *MERA CASE* (or *STATUS*) any time to get your next hearing date and last update."
        )

    matters = get_client_matters_status(candidate["client_id"], candidate["org_id"])
    if not matters:
        return _twiml_reply(f"Hi {candidate.get('name') or 'there'}, no active matters found on file for you right now.")

    lines = [f"📋 *Case Update for {candidate.get('name') or 'you'}*\n"]
    for m in matters:
        lines.append(f"*{m['title']}* ({m['status']})")
        if m.get("last_outcome"):
            lines.append(f"  Last hearing ({m.get('last_hearing_date', '?')}): {m['last_outcome']}")
        next_date = m.get("next_hearing_date") or m.get("next_date_from_order")
        lines.append(f"  Next date: {next_date}" if next_date else "  Next date: not yet fixed")
        lines.append("")
    lines.append("_Sent automatically by Project Ease. For details, contact your lawyer's office._")
    return _twiml_reply("\n".join(lines))


@bp.route("/webhook/whatsapp", methods=["POST"])
async def whatsapp_webhook():
    """Twilio WhatsApp webhook — receives messages, runs them through the chat
    pipeline, and returns a TwiML reply.

    Task #40 additions:
    - Self-registration flow for unregistered numbers
    - Matter-scoped queries via "MATTER: <name>" prefix
    - HELP / LIST MATTERS / RESET commands
    """
    form        = await request.form
    from_raw    = form.get("From", "")                  # "whatsapp:+923001234567"
    body        = (form.get("Body") or "").strip()
    media_type  = form.get("MediaContentType0", "")
    media_url   = form.get("MediaUrl0", "")

    # Strip the "whatsapp:" prefix Twilio prepends
    phone = from_raw.removeprefix("whatsapp:").strip()
    if not phone:
        return _twiml_reply("Could not identify your number. Please contact support.")

    # ── Self-registration flow (Task #40) ─────────────────────────────────────
    if phone in _wa_reg_sessions:
        return await _wa_handle_registration(phone, body)

    user = get_user_by_whatsapp(phone)
    if not user:
        # ── Client self-service bot: is this number a client's, not a team member's? ──
        client_reply = await _wa_client_self_service(phone, body)
        if client_reply is not None:
            return client_reply

        # Unregistered number — start onboarding
        cmd = body.upper()
        if cmd in ("YES", "Y", "REGISTER", "SIGN UP", "SIGNUP"):
            _wa_reg_sessions[phone] = {"step": 1, "data": {}}
            return _twiml_reply(
                "Great! Let's register your law firm on Project Ease.\n\n"
                "Step 1 of 4: What is the *name of your law firm*?"
            )
        return _twiml_reply(
            "👋 Welcome to *Project Ease* — AI-powered legal research for Pakistani law firms.\n\n"
            "Your WhatsApp number is not yet registered. Would you like to register your firm?\n"
            "Reply *YES* to begin registration or contact us at projectease.pk"
        )

    # ── Registered user — determine query text ────────────────────────────────
    if media_url and "audio" in media_type:
        transcript = await _transcribe_voice(media_url)
        if transcript is None:
            return _twiml_reply(
                "Voice messages are not supported yet — please send your question as text."
            )
        query_text = transcript
    elif body:
        query_text = body
    else:
        return _twiml_reply("Please send a text or voice message with your question.")

    # ── Handle commands ───────────────────────────────────────────────────────
    cmd_upper = query_text.upper().strip()

    if cmd_upper == "HELP":
        return _twiml_reply(_WA_HELP)

    if cmd_upper == "RESET":
        _wa_sessions.pop(phone, None)
        return _twiml_reply("Conversation history cleared. Start a new question!")

    if cmd_upper == "LIST MATTERS":
        matters = get_matters(user["org_id"])[:10]
        if not matters:
            return _twiml_reply("No active matters found for your firm.")
        lines = ["📁 *Active Matters*\n"]
        for m in matters:
            line = f"• {m['title']} ({m['status']})"
            if m.get("case_number"):
                line += f" — Case {m['case_number']}"
            lines.append(line)
        return _twiml_reply("\n".join(lines))

    # ── Matter-scoped query: "MATTER: <name>\n<question>" ────────────────────
    matter_filter: list[str] | None = None
    matter_label  = ""
    for prefix in ("MATTER:", "M:"):
        if query_text.upper().startswith(prefix):
            rest        = query_text[len(prefix):].strip()
            # Support "MATTER: Khan vs State\nWhat is the filing date?"
            if "\n" in rest:
                matter_ref, query_text = rest.split("\n", 1)
                matter_ref = matter_ref.strip()
                query_text = query_text.strip()
            else:
                matter_ref = rest
                query_text = rest   # treat the rest as the query too (single-line)
            matches = search_matters_by_keyword(user["org_id"], matter_ref)
            if not matches:
                return _twiml_reply(
                    f"No matter matching *{matter_ref}* found. "
                    "Use LIST MATTERS to see available matters."
                )
            best = matches[0]
            matter_filter = get_filenames_for_matter(best["matter_id"], user["org_id"])
            matter_label  = best["title"]
            break

    # ── Build permitted file list ─────────────────────────────────────────────
    if matter_filter is not None:
        permitted = matter_filter
        if not permitted:
            return _twiml_reply(
                f"Matter *{matter_label}* has no documents yet. "
                "Upload documents in the owner portal to enable matter-scoped search."
            )
    else:
        cat_ids   = get_permitted_categories(user["user_id"])
        docs      = get_docs_for_categories(user["org_id"], cat_ids) if cat_ids else []
        permitted = [d["filename"] for d in docs]

    # ── Maintain per-number conversation history (last 10 turns) ─────────────
    history = _wa_sessions.setdefault(phone, [])
    history.append({"role": "user", "content": query_text})
    if len(history) > 10:
        _wa_sessions[phone] = history[-10:]
        history = _wa_sessions[phone]

    # Prefix matter context into the query when scoped
    chat_messages = history.copy()
    if matter_label:
        system_note = (
            f"[Context: This query is scoped to matter '{matter_label}'. "
            "Answer using only the documents provided for this matter.]"
        )
        chat_messages = [{"role": "user", "content": system_note}] + chat_messages

    # ── Run through the chat pipeline ─────────────────────────────────────────
    approach: Approach = cast(Approach, current_app.config[CONFIG_CHAT_APPROACH])
    try:
        result = await approach.run(
            messages=chat_messages,
            context={
                "overrides": {
                    "retrieval_mode":  "hybrid",
                    "semantic_ranker": True,
                    "top":             3,
                    "permitted_sourcefiles": permitted,
                },
                "auth_claims": {
                    "organization_id": user["org_id"],
                },
            },
        )
        answer = (result.get("output_text") or "").strip()
        if not answer:
            answer = "I could not find a relevant answer in your firm's documents. Please try rephrasing your question."
    except Exception as exc:
        logging.exception("WhatsApp chat pipeline error: %s", exc)
        answer = "Sorry, I ran into an error processing your request. Please try again in a moment."

    if matter_label:
        answer = f"📁 *{matter_label}*\n\n{answer}"

    # Add assistant reply to history
    history.append({"role": "assistant", "content": answer})

    # WhatsApp messages are capped at 1600 characters
    if len(answer) > 1500:
        answer = answer[:1497] + "…"

    return _twiml_reply(answer)


async def _wa_handle_registration(phone: str, body: str):
    """Handle one step of the WhatsApp self-registration flow (Task #40).

    Steps:
        1 → firm name
        2 → contact name
        3 → email
        4 → password  → create org + owner account
    """
    session  = _wa_reg_sessions[phone]
    step     = session["step"]
    data     = session["data"]
    text     = body.strip()
    cancel   = text.upper() in ("CANCEL", "QUIT", "EXIT", "NO", "N")

    if cancel:
        del _wa_reg_sessions[phone]
        return _twiml_reply("Registration cancelled. Reply REGISTER any time to try again.")

    if step == 1:
        if len(text) < 2:
            return _twiml_reply("Please enter a valid firm name (at least 2 characters).")
        data["firm"] = text
        session["step"] = 2
        return _twiml_reply(
            f"✅ Firm: *{text}*\n\nStep 2 of 4: What is *your full name*?"
        )

    if step == 2:
        if len(text) < 2:
            return _twiml_reply("Please enter your full name.")
        data["name"] = text
        session["step"] = 3
        return _twiml_reply(
            f"✅ Name: *{text}*\n\nStep 3 of 4: Enter your *email address*. "
            "This will be your login for the web portal."
        )

    if step == 3:
        import re as _re
        if not _re.match(r"[^@\s]+@[^@\s]+\.[^@\s]+", text):
            return _twiml_reply("That doesn't look like a valid email address. Please try again.")
        # Check for duplicate
        if get_user_by_email(text.lower()):
            return _twiml_reply(
                "An account with that email already exists. "
                "Please use a different email or log in at the web portal."
            )
        data["email"] = text.lower()
        session["step"] = 4
        return _twiml_reply(
            f"✅ Email: *{text}*\n\nStep 4 of 4: Create a *password* for your account.\n"
            "It must be at least 8 characters. "
            "⚠️ This message will not be stored — set a strong password."
        )

    if step == 4:
        if len(text) < 8:
            return _twiml_reply("Password must be at least 8 characters. Please try again.")
        data["password"] = text
        # Create the org and owner account
        try:
            new_org  = create_org(data["firm"], plan="trial", actor="whatsapp_onboarding")
            new_user = create_user(
                org_id   = new_org["org_id"],
                email    = data["email"],
                name     = data["name"],
                role     = "org_owner",
                password = data["password"],
                actor    = "whatsapp_onboarding",
            )
            # Link WhatsApp number to the new owner account
            update_user_whatsapp(new_user["user_id"], phone, actor="whatsapp_onboarding")
        except Exception as exc:
            logging.exception("WhatsApp registration error: %s", exc)
            del _wa_reg_sessions[phone]
            return _twiml_reply(
                "Registration failed due to an internal error. Please try again or contact support."
            )
        del _wa_reg_sessions[phone]
        return _twiml_reply(
            f"🎉 *Registration complete!*\n\n"
            f"Firm: *{data['firm']}*\n"
            f"Login: *{data['email']}*\n\n"
            f"Visit *projectease.pk* and sign in with your email and password.\n\n"
            f"You're on a 14-day free trial. Send *HELP* to see what I can do for you right now!"
        )


# ─── PROJECT EASE: Super Admin API ───────────────────────────────────────────

def _require_platform_admin():
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Forbidden"}), 403
    return None


@bp.route("/admin/stats", methods=["GET"])
async def admin_stats():
    err = _require_platform_admin()
    if err: return err
    return jsonify(get_platform_stats())


@bp.route("/admin/orgs", methods=["GET"])
async def admin_list_orgs():
    err = _require_platform_admin()
    if err: return err
    return jsonify({"orgs": get_all_orgs()})


@bp.route("/admin/orgs", methods=["POST"])
async def admin_create_org():
    err = _require_platform_admin()
    if err: return err
    data = await request.get_json(silent=True) or {}
    name     = (data.get("name") or "").strip()
    plan     = data.get("plan", "free")
    industry = data.get("industry", "Other")
    owner_name  = (data.get("owner_name") or "").strip()
    owner_email = (data.get("owner_email") or "").strip().lower()
    if not name or not owner_name or not owner_email:
        return jsonify({"error": "Org name, owner name, and owner email are required"}), 400
    try:
        admin_session = _get_session()
        admin_id = (admin_session or {}).get("email") or SYSTEM
        org = create_org(name=name, plan=plan, industry=industry, actor=admin_id)
        temp_pw = _secrets.token_urlsafe(10)
        owner = create_user(
            org_id=org["org_id"], email=owner_email,
            name=owner_name, role="org_owner",
            actor=admin_id,
            password=temp_pw, must_change=True,
        )
        return jsonify({"org": org, "owner": owner, "temp_password": temp_pw}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@bp.route("/admin/orgs/<org_id>", methods=["GET"])
async def admin_get_org(org_id: str):
    err = _require_platform_admin()
    if err: return err
    details = get_org_details(org_id)
    if not details:
        return jsonify({"error": "Not found"}), 404
    return jsonify(details)


@bp.route("/admin/orgs/<org_id>", methods=["PUT"])
async def admin_update_org(org_id: str):
    err = _require_platform_admin()
    if err: return err
    data = await request.get_json(silent=True) or {}
    admin_session = _get_session()
    admin_id = (admin_session or {}).get("email") or SYSTEM
    updated = update_org(org_id, actor=admin_id, **data)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    # If org was suspended, invalidate all its sessions
    if data.get("status") == "suspended":
        for tok, s in list(_sessions.items()):
            if s.get("org") == org_id:
                _sessions.pop(tok, None)
    return jsonify(updated)


@bp.route("/admin/orgs/<org_id>", methods=["DELETE"])
async def admin_delete_org(org_id: str):
    err = _require_platform_admin()
    if err: return err
    # Invalidate all sessions for this org
    for tok, s in list(_sessions.items()):
        if s.get("org") == org_id:
            _sessions.pop(tok, None)
    admin_session = _get_session()
    admin_id = (admin_session or {}).get("email") or SYSTEM
    delete_org(org_id, actor=admin_id)
    return jsonify({"success": True})


# PROJECT EASE: admin eval endpoints
# Protected by a shared secret — set ADMIN_EVAL_API_KEY in your environment.
# Callers must send:  Authorization: Bearer <ADMIN_EVAL_API_KEY>
# This keeps law-firm query data from leaking to unauthenticated callers.
def _require_admin_key():
    """Returns an error response if the bearer token is missing or wrong, else None."""
    expected = os.environ.get("ADMIN_EVAL_API_KEY", "")
    if not expected:
        # Key not configured → block all access (fail-closed, not fail-open)
        return jsonify({"error": "Admin API is disabled. Set ADMIN_EVAL_API_KEY to enable it."}), 403
    auth_header = request.headers.get("Authorization", "")
    provided = auth_header.removeprefix("Bearer ").strip()
    if not provided or provided != expected:
        return jsonify({"error": "Unauthorized"}), 401
    return None


@bp.route("/admin/evals", methods=["GET"])
async def admin_evals():
    """Return recent eval results. Optional ?org=<org_id> to filter by tenant."""
    if err := _require_admin_key():
        return err
    from evals.db import fetch_recent, fetch_summary
    org = request.args.get("org")
    results = fetch_recent(limit=100)
    if org:
        results = [r for r in results if r.get("organization_id") == org]
    summary = fetch_summary(organization_id=org)
    return jsonify({"summary": summary, "results": results})


@bp.route("/admin/evals/summary", methods=["GET"])
async def admin_evals_summary():
    """Aggregate stats across all orgs or a specific org."""
    if err := _require_admin_key():
        return err
    from evals.db import fetch_summary
    org = request.args.get("org")
    return jsonify(fetch_summary(organization_id=org))


@bp.route("/config", methods=["GET"])
def config():
    return jsonify(
        {
            "showMultimodalOptions": current_app.config[CONFIG_MULTIMODAL_ENABLED],
            "showSemanticRankerOption": current_app.config[CONFIG_SEMANTIC_RANKER_DEPLOYED],
            "showQueryRewritingOption": current_app.config[CONFIG_QUERY_REWRITING_ENABLED],
            "showReasoningEffortOption": current_app.config[CONFIG_REASONING_EFFORT_ENABLED],
            "streamingEnabled": current_app.config[CONFIG_STREAMING_ENABLED],
            "defaultReasoningEffort": current_app.config[CONFIG_DEFAULT_REASONING_EFFORT],
            "reasoningEffortOptions": current_app.config[CONFIG_REASONING_EFFORT_OPTIONS],
            "defaultRetrievalReasoningEffort": current_app.config[CONFIG_DEFAULT_RETRIEVAL_REASONING_EFFORT],
            "showVectorOption": current_app.config[CONFIG_VECTOR_SEARCH_ENABLED],
            "showUserUpload": current_app.config[CONFIG_USER_UPLOAD_ENABLED],
            "showLanguagePicker": current_app.config[CONFIG_LANGUAGE_PICKER_ENABLED],
            "showSpeechInput": current_app.config[CONFIG_SPEECH_INPUT_ENABLED],
            "showSpeechOutputBrowser": current_app.config[CONFIG_SPEECH_OUTPUT_BROWSER_ENABLED],
            "showSpeechOutputAzure": current_app.config[CONFIG_SPEECH_OUTPUT_AZURE_ENABLED],
            "showChatHistoryBrowser": current_app.config[CONFIG_CHAT_HISTORY_BROWSER_ENABLED],
            "showChatHistoryCosmos": current_app.config[CONFIG_CHAT_HISTORY_COSMOS_ENABLED],
            "showAgenticRetrievalOption": current_app.config[CONFIG_AGENTIC_KNOWLEDGEBASE_ENABLED],
            "ragSearchTextEmbeddings": current_app.config[CONFIG_RAG_SEARCH_TEXT_EMBEDDINGS],
            "ragSearchImageEmbeddings": current_app.config[CONFIG_RAG_SEARCH_IMAGE_EMBEDDINGS],
            "ragSendTextSources": current_app.config[CONFIG_RAG_SEND_TEXT_SOURCES],
            "ragSendImageSources": current_app.config[CONFIG_RAG_SEND_IMAGE_SOURCES],
            "webSourceEnabled": current_app.config[CONFIG_WEB_SOURCE_ENABLED],
            "sharepointSourceEnabled": current_app.config[CONFIG_SHAREPOINT_SOURCE_ENABLED],
        }
    )


@bp.route("/speech", methods=["POST"])
async def speech():
    if not request.is_json:
        return jsonify({"error": "request must be json"}), 415

    if current_app.config.get(CONFIG_CREDENTIAL) is None:
        return jsonify({"error": "Speech isn't configured — Azure credentials required."}), 503

    speech_token = current_app.config.get(CONFIG_SPEECH_SERVICE_TOKEN)
    if speech_token is None or speech_token.expires_on < time.time() + 60:
        speech_token = await current_app.config[CONFIG_CREDENTIAL].get_token(
            "https://cognitiveservices.azure.com/.default"
        )
        current_app.config[CONFIG_SPEECH_SERVICE_TOKEN] = speech_token

    request_json = await request.get_json()
    text = request_json["text"]
    try:
        # Construct a token as described in documentation:
        # https://learn.microsoft.com/azure/ai-services/speech-service/how-to-configure-azure-ad-auth?pivots=programming-language-python
        auth_token = (
            "aad#"
            + current_app.config[CONFIG_SPEECH_SERVICE_ID]
            + "#"
            + current_app.config[CONFIG_SPEECH_SERVICE_TOKEN].token
        )
        speech_config = SpeechConfig(auth_token=auth_token, region=current_app.config[CONFIG_SPEECH_SERVICE_LOCATION])
        speech_config.speech_synthesis_voice_name = current_app.config[CONFIG_SPEECH_SERVICE_VOICE]
        speech_config.set_speech_synthesis_output_format(SpeechSynthesisOutputFormat.Audio16Khz32KBitRateMonoMp3)
        synthesizer = SpeechSynthesizer(speech_config=speech_config, audio_config=None)
        result: SpeechSynthesisResult = synthesizer.speak_text_async(text).get()
        if result.reason == ResultReason.SynthesizingAudioCompleted:
            return result.audio_data, 200, {"Content-Type": "audio/mp3"}
        elif result.reason == ResultReason.Canceled:
            cancellation_details = result.cancellation_details
            current_app.logger.error(
                "Speech synthesis canceled: %s %s", cancellation_details.reason, cancellation_details.error_details
            )
            raise Exception("Speech synthesis canceled. Check logs for details.")
        else:
            current_app.logger.error("Unexpected result reason: %s", result.reason)
            raise Exception("Speech synthesis failed. Check logs for details.")
    except Exception as e:
        current_app.logger.exception("Exception in /speech")
        return jsonify({"error": str(e)}), 500


@bp.post("/upload")
@authenticated
async def upload(auth_claims: dict[str, Any]):
    request_files = await request.files
    if "file" not in request_files:
        return jsonify({"message": "No file part in the request", "status": "failed"}), 400

    try:
        user_oid = auth_claims["oid"]
        file = request_files.getlist("file")[0]
        adls_manager: AdlsBlobManager = current_app.config[CONFIG_USER_BLOB_MANAGER]
        file_url = await adls_manager.upload_blob(file, file.filename, user_oid)
        ingester: UploadUserFileStrategy = current_app.config[CONFIG_INGESTER]
        await ingester.add_file(File(content=file, url=file_url, acls={"oids": [user_oid]}), user_oid=user_oid)
        return jsonify({"message": "File uploaded successfully"}), 200
    except Exception as error:
        current_app.logger.error("Error uploading file: %s", error)
        return jsonify({"message": "Error uploading file, check server logs for details.", "status": "failed"}), 500


@bp.post("/delete_uploaded")
@authenticated
async def delete_uploaded(auth_claims: dict[str, Any]):
    request_json = await request.get_json()
    filename = request_json.get("filename")
    user_oid = auth_claims["oid"]
    adls_manager: AdlsBlobManager = current_app.config[CONFIG_USER_BLOB_MANAGER]
    await adls_manager.remove_blob(filename, user_oid)
    ingester: UploadUserFileStrategy = current_app.config[CONFIG_INGESTER]
    await ingester.remove_file(filename, user_oid)
    return jsonify({"message": f"File {filename} deleted successfully"}), 200


@bp.get("/list_uploaded")
@authenticated
async def list_uploaded(auth_claims: dict[str, Any]):
    """Lists the uploaded documents for the current user.
    Only returns files directly in the user's directory, not in subdirectories.
    Excludes image files and the images directory."""
    user_oid = auth_claims["oid"]
    adls_manager: AdlsBlobManager = current_app.config[CONFIG_USER_BLOB_MANAGER]
    files = await adls_manager.list_blobs(user_oid)
    return jsonify(files), 200


# ─── PROJECT EASE: Clients API ───────────────────────────────────────────────

@bp.route("/clients", methods=["GET"])
async def list_clients():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"clients": get_clients(session.get("org") or "")})


@bp.route("/clients", methods=["POST"])
async def add_client():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Client name is required"}), 400
    client = create_client(
        org_id=session.get("org") or "",
        name=name,
        client_type=data.get("client_type", "Individual"),
        email=data.get("email") or None,
        phone=data.get("phone") or None,
        address=data.get("address") or None,
        cnic_ntn=data.get("cnic_ntn") or None,
        notes=data.get("notes") or None,
        referral_source=data.get("referral_source") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    _audit(session, "client_create",
           resource_type="client", resource_id=client["client_id"], resource_name=name)
    return jsonify(client), 201


@bp.route("/clients/<client_id>", methods=["GET"])
async def get_client_detail(client_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    client = get_client_with_matters(client_id, session.get("org") or "")
    if not client:
        return jsonify({"error": "Not found"}), 404
    return jsonify(client)


@bp.route("/clients/<client_id>", methods=["PATCH"])
async def edit_client(client_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_client(
        client_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **data,
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    _audit(session, "client_update",
           resource_type="client", resource_id=client_id,
           resource_name=(updated or {}).get("name"))
    return jsonify(updated)


@bp.route("/clients/<client_id>", methods=["DELETE"])
async def remove_client(client_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_client(client_id, session.get("org") or "",
                  actor=session.get("user_id") or SYSTEM)
    _audit(session, "client_delete", resource_type="client", resource_id=client_id)
    return jsonify({"success": True})


# ─── PROJECT EASE: Matter Teams API ──────────────────────────────────────────

@bp.route("/matter-teams", methods=["GET"])
async def list_matter_teams():
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"teams": get_matter_teams(session.get("org") or "")})


@bp.route("/matter-teams", methods=["POST"])
async def add_matter_team():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Team name is required"}), 400
    team = create_matter_team(session.get("org") or "", name,
                              actor=session.get("user_id") or SYSTEM)
    return jsonify(team), 201


@bp.route("/matter-teams/<team_id>", methods=["PATCH"])
async def edit_matter_team(team_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Team name is required"}), 400
    updated = update_matter_team(team_id, session.get("org") or "", name,
                                 actor=session.get("user_id") or SYSTEM)
    return jsonify(updated or {})


@bp.route("/matter-teams/<team_id>", methods=["DELETE"])
async def remove_matter_team(team_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_team(team_id, session.get("org") or "",
                       actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


@bp.route("/matter-teams/<team_id>/members/<member_user_id>", methods=["POST"])
async def add_member_to_matter_team(team_id: str, member_user_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    add_matter_team_member(team_id, member_user_id,
                           actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


@bp.route("/matter-teams/<team_id>/members/<member_user_id>", methods=["DELETE"])
async def remove_member_from_matter_team(team_id: str, member_user_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    remove_matter_team_member(team_id, member_user_id,
                              actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── PROJECT EASE: Courts API ─────────────────────────────────────────────────

_DEFAULT_COURTS = [
    "Supreme Court of Pakistan", "Federal Shariat Court",
    "Lahore High Court", "Sindh High Court", "Islamabad High Court",
    "Peshawar High Court", "Balochistan High Court",
    "Gilgit-Baltistan Chief Court", "Azad Kashmir High Court",
    "District & Sessions Court", "Civil Judge Court", "Magistrate Court",
    "Banking Court", "Labour Court", "National Accountability Court",
    "Customs Appellate Tribunal", "Income Tax Appellate Tribunal",
    "Anti-Corruption Establishment Court", "Service Tribunal", "Family Court",
]


@bp.route("/courts", methods=["GET"])
async def list_courts():
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    custom = get_custom_courts(session.get("org") or "")
    return jsonify({"default": _DEFAULT_COURTS, "custom": custom})


@bp.route("/courts", methods=["POST"])
async def add_court():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Court name is required"}), 400
    try:
        court = add_custom_court(session.get("org") or "", name,
                                 actor=session.get("user_id") or SYSTEM)
        return jsonify(court), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@bp.route("/courts/<court_id>", methods=["DELETE"])
async def remove_court(court_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_custom_court(court_id, session.get("org") or "",
                        actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── PROJECT EASE: Matters API ────────────────────────────────────────────────

@bp.route("/matters", methods=["GET"])
async def list_matters():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    client_id = request.args.get("client_id") or None
    return jsonify({"matters": get_matters(session.get("org") or "", client_id=client_id),
                    "priorities": list(MATTER_PRIORITIES)})


@bp.route("/matters", methods=["POST"])
async def add_matter():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data        = await request.get_json(silent=True) or {}
    client_id   = (data.get("client_id")   or "").strip()
    title       = (data.get("title")        or "").strip()
    matter_type = (data.get("matter_type")  or "").strip()
    if not client_id or not title or not matter_type:
        return jsonify({"error": "Client, title, and matter type are required"}), 400
    lim_type = data.get("limitation_type") or None
    coa_date = data.get("cause_of_action_date") or None
    lim_date = data.get("limitation_date") or (
        compute_limitation_date(lim_type, coa_date) if lim_type and coa_date else None
    )
    matter = create_matter(
        org_id=session.get("org") or "",
        client_id=client_id,
        title=title,
        matter_type=matter_type,
        status=data.get("status", "Active"),
        court_name=data.get("court_name") or None,
        case_number=data.get("case_number") or None,
        filing_date=data.get("filing_date") or None,
        opposing_party=data.get("opposing_party") or None,
        team_id=data.get("team_id") or None,
        notes=data.get("notes") or None,
        limitation_type=lim_type,
        cause_of_action_date=coa_date,
        limitation_date=lim_date,
        vakalatnama_status=data.get("vakalatnama_status", "Pending"),
        priority=data.get("priority", "Normal"),
        actor=session.get("user_id") or SYSTEM,
    )
    _audit(session, "matter_create",
           resource_type="matter", resource_id=matter["matter_id"], resource_name=title,
           details={"matter_type": matter_type, "status": data.get("status", "Active")})
    return jsonify(matter), 201


@bp.route("/matters/<matter_id>", methods=["GET"])
async def get_matter_detail(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    matter = get_matter_with_docs(matter_id, session.get("org") or "")
    if not matter:
        return jsonify({"error": "Not found"}), 404
    return jsonify(matter)


@bp.route("/matters/<matter_id>", methods=["PATCH"])
async def edit_matter(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    # Normalize empty strings to None for nullable fields
    for k in ("court_name", "case_number", "filing_date", "opposing_party", "team_id", "notes",
              "limitation_type", "cause_of_action_date", "limitation_date"):
        if k in data and data[k] == "":
            data[k] = None
    # Auto-compute limitation_date if type + coa changed but date not explicitly provided
    if ("limitation_type" in data or "cause_of_action_date" in data) and "limitation_date" not in data:
        lim_type = data.get("limitation_type")
        coa_date = data.get("cause_of_action_date")
        if lim_type and coa_date:
            data["limitation_date"] = compute_limitation_date(lim_type, coa_date)
    updated = update_matter(
        matter_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **data,
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    _audit(session, "matter_update",
           resource_type="matter", resource_id=matter_id,
           resource_name=(updated or {}).get("title"),
           details={"status": data.get("status")})
    return jsonify(updated)


@bp.route("/matters/<matter_id>", methods=["DELETE"])
async def remove_matter(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter(matter_id, session.get("org") or "",
                  actor=session.get("user_id") or SYSTEM)
    _audit(session, "matter_delete", resource_type="matter", resource_id=matter_id)
    return jsonify({"success": True})


@bp.route("/matters/<matter_id>/documents/<doc_id>", methods=["POST"])
async def link_doc_to_matter(matter_id: str, doc_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    ok = link_document_to_matter(doc_id, matter_id, session.get("org") or "",
                                 actor=session.get("user_id") or SYSTEM)
    if not ok:
        return jsonify({"error": "Document not found"}), 404
    return jsonify({"success": True})


@bp.route("/matters/<matter_id>/documents/<doc_id>", methods=["DELETE"])
async def unlink_doc_from_matter(matter_id: str, doc_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    unlink_document_from_matter(doc_id, session.get("org") or "",
                                actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── Court Orders (Task #130) ────────────────────────────────────────────────

async def _notify_client_of_order_outcome(session, matter_id: str, org_id: str, order: dict) -> None:
    """Task: auto client WhatsApp after a hearing outcome is marked.

    Fire-and-forget best effort — never raises, so a WhatsApp/Twilio hiccup
    can't block saving the court order itself.
    """
    outcome    = (order.get("outcome") or "").strip()
    next_date  = (order.get("next_date") or "").strip()
    if not outcome:
        return
    try:
        from whatsapp_helper import send_whatsapp_text, normalize_pk_number
    except ImportError:
        logging.warning("whatsapp_helper not available — client notification skipped.")
        return

    contact = get_matter_client_contact(matter_id, org_id)
    if not contact or not contact.get("client_phone"):
        return
    phone = normalize_pk_number(contact["client_phone"])
    if not phone:
        logging.info("Client phone for matter %s not a valid PK number — notification skipped.", matter_id)
        return

    title = contact.get("title") or "your matter"
    if outcome == "Decided":
        message = f'Your case "{title}" was decided today. Please contact our office for details.'
    elif outcome == "Adjourned":
        next_line = f"Next date: {next_date}." if next_date else "Next date will be shared shortly."
        message = f'Your case "{title}" was heard today and adjourned. {next_line}'
    else:
        next_line = f" Next date: {next_date}." if next_date else ""
        message = f'Your case "{title}" was heard today.{next_line}'

    import asyncio as _aio
    result = await _aio.to_thread(send_whatsapp_text, phone, message)
    _audit(session, "client_whatsapp_notify",
           resource_type="matter", resource_id=matter_id, resource_name=title,
           details={"outcome": outcome, "sent": result.get("sent"), "error": result.get("error")})


@bp.route("/voice/log-outcome", methods=["POST"])
async def voice_log_outcome():
    """Voice input for hearing outcomes (Urdu or English) — a lawyer records
    a short voice note in court ("agli tareekh pandrah August, muqadma suna
    gaya, judge ne bail manzoor ki") and this transcribes + extracts
    structured fields to PRE-FILL the Add Court Order form. Nothing is saved
    here — the lawyer always reviews/edits the extracted fields before the
    normal POST /matters/<id>/orders save, since a misheard word could
    otherwise silently corrupt a case record.
    """
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401

    whisper_configured = bool(os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT", "").strip())
    if not whisper_configured:
        return jsonify({
            "error": "Voice input isn't configured yet — AZURE_OPENAI_WHISPER_DEPLOYMENT "
                     "is not set. Type the order summary instead, or ask your admin to configure it."
        }), 400

    files = await request.files
    upload = files.get("audio")
    if upload is None:
        return jsonify({"error": "No audio file provided"}), 400

    audio_bytes = upload.read()
    content_type = upload.content_type or "audio/webm"
    transcript = await _whisper_transcribe_bytes(audio_bytes, upload.filename or "voice_note.webm", content_type)
    if not transcript:
        return jsonify({"error": "Could not transcribe that recording — try again, closer to the mic."}), 400

    result = {"transcript": transcript, "outcome": None, "next_date": None, "order_brief": transcript}

    if current_app.config.get("AZURE_CONFIGURED"):
        try:
            import datetime as _dt
            import re as _re
            openai_client = current_app.config[CONFIG_OPENAI_CLIENT]
            today = _dt.date.today().strftime("%Y-%m-%d")
            system_msg = (
                "You extract structured hearing-outcome data from a Pakistani lawyer's spoken "
                "voice note (often Urdu, English, or a mix). Return ONLY a JSON object with keys: "
                '"outcome" (one of "Adjourned", "Heard", "Decided", "Partially Heard", or null if unclear), '
                '"next_date" (an ISO YYYY-MM-DD date if a next hearing date was mentioned, else null — '
                f"today's date is {today}, resolve relative/spoken dates like \"15 August\" against it), "
                '"order_brief" (a clean one/two-sentence English summary of what happened, suitable for '
                "a case file — translate from Urdu if needed, but keep any case-specific names/terms accurate). "
                "If you are not confident about a field, use null rather than guessing."
            )
            ai_resp = await openai_client.chat.completions.create(
                model=current_app.config.get("OPENAI_CHATGPT_MODEL", "gpt-5-mini-1"),
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": transcript},
                ],
                temperature=0.1,
                max_tokens=400,
            )
            raw = (ai_resp.choices[0].message.content or "").strip()
            raw = _re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=_re.MULTILINE).strip()
            extracted = json.loads(raw)
            if extracted.get("outcome") in ("Adjourned", "Heard", "Decided", "Partially Heard"):
                result["outcome"] = extracted["outcome"]
            if extracted.get("next_date"):
                result["next_date"] = extracted["next_date"]
            if extracted.get("order_brief"):
                result["order_brief"] = extracted["order_brief"]
        except Exception as exc:
            logging.warning("Voice outcome extraction failed, falling back to raw transcript: %s", exc)

    return jsonify(result)


@bp.route("/matters/<matter_id>/orders", methods=["GET"])
async def list_court_orders(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    orders = get_court_orders(matter_id, org_id)
    return jsonify({"orders": orders})


@bp.route("/matters/<matter_id>/orders", methods=["POST"])
async def add_court_order(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    hearing_date = (data.get("hearing_date") or "").strip()
    order_brief  = (data.get("order_brief")  or "").strip()
    if not hearing_date or not order_brief:
        return jsonify({"error": "hearing_date and order_brief are required"}), 400
    org_id = session.get("org") or ""
    order = create_court_order(
        matter_id=matter_id,
        org_id=org_id,
        hearing_date=hearing_date,
        order_brief=order_brief,
        outcome=data.get("outcome", "Adjourned"),
        court_name=data.get("court_name") or None,
        next_date=data.get("next_date") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    if data.get("notify_client", True):
        await _notify_client_of_order_outcome(session, matter_id, org_id, order)
    return jsonify(order), 201


@bp.route("/matters/<matter_id>/orders/<order_id>", methods=["PATCH"])
async def edit_court_order(matter_id: str, order_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    org_id = session.get("org") or ""
    updated = update_court_order(
        order_id, org_id,
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"hearing_date", "court_name", "order_brief", "next_date", "outcome"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    if "outcome" in data and data.get("notify_client", True):
        await _notify_client_of_order_outcome(session, matter_id, org_id, updated)
    return jsonify(updated)


@bp.route("/matters/<matter_id>/orders/<order_id>", methods=["DELETE"])
async def remove_court_order(matter_id: str, order_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_court_order(order_id, session.get("org") or "",
                       actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── Limitation Alerts (Task #132) ──────────────────────────────────────────

@bp.route("/matters/limitation-alerts", methods=["GET"])
async def limitation_alerts():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    within = int(request.args.get("within_days", 60))
    alerts = get_matters_with_approaching_limitation(session.get("org") or "", within_days=within)
    return jsonify({"alerts": alerts, "limitation_types": list(LIMITATION_PERIODS.keys())})


# ─── Adverse Parties (Task #131) ─────────────────────────────────────────────

@bp.route("/matters/<matter_id>/adverse-parties", methods=["GET"])
async def list_adverse_parties(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    parties = get_adverse_parties(matter_id, session.get("org") or "")
    return jsonify({"parties": parties})


@bp.route("/matters/<matter_id>/adverse-parties", methods=["POST"])
async def add_adverse_party(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    party_name = (data.get("party_name") or "").strip()
    if not party_name:
        return jsonify({"error": "party_name is required"}), 400
    party = create_adverse_party(
        matter_id=matter_id,
        org_id=session.get("org") or "",
        party_name=party_name,
        party_type=data.get("party_type", "Individual"),
        counsel_name=data.get("counsel_name") or None,
        counsel_phone=data.get("counsel_phone") or None,
        counsel_firm=data.get("counsel_firm") or None,
        notes=data.get("notes") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(party), 201


@bp.route("/matters/<matter_id>/adverse-parties/<party_id>", methods=["PATCH"])
async def edit_adverse_party(matter_id: str, party_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    updated = update_adverse_party(
        party_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"party_name", "party_type", "counsel_name", "counsel_phone", "counsel_firm", "notes"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/adverse-parties/<party_id>", methods=["DELETE"])
async def remove_adverse_party(matter_id: str, party_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_adverse_party(party_id, session.get("org") or "",
                         actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


# ─── Time Tracking (Task #133) ───────────────────────────────────────────────

@bp.route("/matters/<matter_id>/time-entries", methods=["GET"])
async def list_time_entries(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    entries = get_time_entries(matter_id, session.get("org") or "")
    return jsonify({"entries": entries})


@bp.route("/matters/<matter_id>/time-entries", methods=["POST"])
async def add_time_entry(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    duration_minutes = int(data.get("duration_minutes") or 0)
    entry_date = (data.get("entry_date") or "").strip()
    if duration_minutes <= 0 or not entry_date:
        return jsonify({"error": "duration_minutes and entry_date are required"}), 400
    entry = create_time_entry(
        matter_id=matter_id,
        org_id=session.get("org") or "",
        duration_minutes=duration_minutes,
        entry_date=entry_date,
        description=data.get("description") or None,
        hourly_rate=int(data.get("hourly_rate") or 0),
        billable=int(data.get("billable", 1)),
        user_id=session.get("user_id") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(entry), 201


@bp.route("/matters/<matter_id>/time-entries/<entry_id>", methods=["PATCH"])
async def edit_time_entry(matter_id: str, entry_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    updated = update_time_entry(
        entry_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"description", "entry_date", "duration_minutes", "hourly_rate", "billable"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/time-entries/<entry_id>", methods=["DELETE"])
async def remove_time_entry(matter_id: str, entry_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_time_entry(entry_id, session.get("org") or "",
                      actor=session.get("user_id") or SYSTEM)
    return jsonify({"success": True})


@bp.route("/matters/<matter_id>/time-entries/bill", methods=["POST"])
async def bill_time_entries_route(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json() or {}
    entry_ids = data.get("entry_ids") or []
    description = (data.get("description") or "Time charges").strip()
    if not entry_ids:
        return jsonify({"error": "entry_ids required"}), 400
    fee = bill_time_entries(
        entry_ids=entry_ids,
        matter_id=matter_id,
        org_id=session.get("org") or "",
        fee_description=description,
        actor=session.get("user_id") or SYSTEM,
    )
    if not fee:
        return jsonify({"error": "No billable unbilled entries found"}), 400
    return jsonify(fee), 201


# ─── Cause List Integration — Task #137 ─────────────────────────────────────

async def _ocr_extract_text(file_bytes: bytes, filename: str) -> str:
    """OCR a single uploaded PDF/image via Azure Document Intelligence and
    return its plain text. Returns "" if Document Intelligence isn't
    configured in this environment (caller should handle that explicitly,
    not treat it the same as "OCR ran and found nothing")."""
    doc_int_service = os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE")
    if not doc_int_service:
        return ""
    doc_int_key = os.getenv("AZURE_DOCUMENTINTELLIGENCE_KEY")

    from prepdocslib.pdfparser import DocumentAnalysisParser

    if doc_int_key and doc_int_key != "YOUR_DOC_INTELLIGENCE_KEY_HERE":
        credential = AzureKeyCredential(doc_int_key)
    else:
        credential = AzureDeveloperCliCredential(process_timeout=60)

    parser = DocumentAnalysisParser(
        endpoint=f"https://{doc_int_service}.cognitiveservices.azure.com/",
        credential=credential,
    )
    stream = io.BytesIO(file_bytes)
    stream.name = filename
    texts = []
    async for page in parser.parse(stream):
        texts.append(page.text)
    return "\n".join(texts)


@bp.route("/cause-list", methods=["GET"])
async def list_cause_list():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    list_date = request.args.get("date") or None
    entries   = get_cause_list_entries(session.get("org") or "", list_date)
    return jsonify({"entries": entries})


@bp.route("/cause-list/today-matches", methods=["GET"])
async def today_cause_list_matches():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    matches = get_today_cause_list_matches(session.get("org") or "")
    return jsonify({"matches": matches})


@bp.route("/cause-list/parse", methods=["POST"])
async def parse_and_store_cause_list():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401

    raw_text   = ""
    list_date  = ""
    court_name = ""
    is_multipart = (request.content_type or "").startswith("multipart/form-data")
    if is_multipart:
        form  = await request.form
        files = await request.files
        list_date  = (form.get("list_date") or "").strip()
        court_name = (form.get("court_name") or "").strip()
        upload = files.get("file")
        if upload is not None:
            doc_int_configured = bool(os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE"))
            if not doc_int_configured:
                return jsonify({
                    "error": "Document scanning isn't configured yet — AZURE_DOCUMENTINTELLIGENCE_SERVICE "
                             "is not set. Paste the cause list text instead, or ask your admin to configure OCR."
                }), 400
            try:
                raw_text = await _ocr_extract_text(upload.read(), upload.filename or "cause_list.pdf")
            except Exception as exc:
                logging.exception("Cause list OCR failed: %s", exc)
                return jsonify({"error": "Could not read that file — try a clearer scan or paste the text instead."}), 400
    else:
        data       = await request.get_json(silent=True) or {}
        raw_text   = (data.get("text") or "").strip()
        list_date  = (data.get("list_date") or "").strip()
        court_name = (data.get("court_name") or "").strip()

    if not raw_text.strip():
        return jsonify({"error": "Cause list text is required"}), 400
    if not list_date:
        import datetime as _dt
        list_date = _dt.datetime.utcnow().strftime("%Y-%m-%d")
    entries = parse_cause_list_text(raw_text, court_name=court_name, list_date=list_date)
    if not entries:
        return jsonify({"error": "No entries could be parsed from the provided text"}), 400
    stored  = store_cause_list(session.get("org") or "", entries,
                               actor=session.get("user_id") or SYSTEM)
    matched = [e for e in stored if e.get("matter_id")]
    return jsonify({"entries": stored, "matched_count": len(matched), "total_count": len(stored)}), 201


@bp.route("/cause-list/<entry_id>", methods=["PATCH"])
async def update_cause_list_entry(entry_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data      = await request.get_json(silent=True) or {}
    matter_id = data.get("matter_id") or None
    link_cause_list_entry(entry_id, session.get("org") or "", matter_id,
                          actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


@bp.route("/cause-list/<entry_id>", methods=["DELETE"])
async def remove_cause_list_entry(entry_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_cause_list_entry(entry_id, session.get("org") or "",
                            actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Notes — Task #138 ────────────────────────────────────────────────

@bp.route("/matters/<matter_id>/notes", methods=["GET"])
async def list_matter_notes(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    notes = get_matter_notes(matter_id, session.get("org") or "")
    return jsonify({"notes": notes, "note_types": NOTE_TYPES})


@bp.route("/matters/<matter_id>/notes", methods=["POST"])
async def add_matter_note(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    note_text = (data.get("note_text") or "").strip()
    if not note_text:
        return jsonify({"error": "note_text required"}), 400
    note = create_matter_note(
        org_id    = session.get("org") or "",
        matter_id = matter_id,
        note_type = data.get("note_type", "Note"),
        note_text = note_text,
        note_date = data.get("note_date") or __import__("datetime").datetime.utcnow().strftime("%Y-%m-%d"),
        actor     = session.get("user_id") or SYSTEM,
    )
    return jsonify(note), 201


@bp.route("/matters/<matter_id>/notes/<note_id>", methods=["PATCH"])
async def edit_matter_note(matter_id: str, note_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    note = update_matter_note(
        note_id   = note_id,
        org_id    = session.get("org") or "",
        note_type = data.get("note_type"),
        note_text = data.get("note_text"),
        note_date = data.get("note_date"),
        actor     = session.get("user_id") or SYSTEM,
    )
    return jsonify(note)


@bp.route("/matters/<matter_id>/notes/<note_id>", methods=["DELETE"])
async def remove_matter_note(matter_id: str, note_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_note(note_id, session.get("org") or "",
                       actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Document Requests — Task #140 ──────────────────────────────────────────

@bp.route("/matters/<matter_id>/doc-requests", methods=["GET"])
async def list_doc_requests(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    reqs = get_document_requests(matter_id, session.get("org") or "")
    return jsonify({"requests": reqs, "statuses": list(DOC_REQUEST_STATUSES)})


@bp.route("/matters/<matter_id>/doc-requests", methods=["POST"])
async def add_doc_request(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data     = await request.get_json(silent=True) or {}
    doc_name = (data.get("doc_name") or "").strip()
    if not doc_name:
        return jsonify({"error": "doc_name required"}), 400
    req = create_document_request(
        org_id         = session.get("org") or "",
        matter_id      = matter_id,
        doc_name       = doc_name,
        requested_date = data.get("requested_date") or __import__("datetime").datetime.utcnow().strftime("%Y-%m-%d"),
        due_date       = data.get("due_date") or None,
        notes          = data.get("notes") or None,
        actor          = session.get("user_id") or SYSTEM,
    )
    return jsonify(req), 201


@bp.route("/matters/<matter_id>/doc-requests/<request_id>", methods=["PATCH"])
async def edit_doc_request(matter_id: str, request_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    req  = update_document_request(
        request_id = request_id,
        org_id     = session.get("org") or "",
        actor      = session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"doc_name", "requested_date", "due_date", "status", "notes", "received_date"}},
    )
    return jsonify(req)


@bp.route("/matters/<matter_id>/doc-requests/<request_id>", methods=["DELETE"])
async def remove_doc_request(matter_id: str, request_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_document_request(request_id, session.get("org") or "",
                            actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Witnesses — Task #141 ───────────────────────────────────────────────────

@bp.route("/matters/<matter_id>/witnesses", methods=["GET"])
async def list_witnesses(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    ws = get_witnesses(matter_id, session.get("org") or "")
    return jsonify({"witnesses": ws, "witness_types": list(WITNESS_TYPES),
                    "statement_statuses": list(STATEMENT_STATUSES)})


@bp.route("/matters/<matter_id>/witnesses", methods=["POST"])
async def add_witness(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    name = (data.get("witness_name") or "").strip()
    if not name:
        return jsonify({"error": "witness_name required"}), 400
    w = create_witness(
        org_id           = session.get("org") or "",
        matter_id        = matter_id,
        witness_name     = name,
        witness_type     = data.get("witness_type", "Defence"),
        contact_number   = data.get("contact_number") or None,
        address          = data.get("address") or None,
        statement_status = data.get("statement_status", "Not Taken"),
        notes            = data.get("notes") or None,
        actor            = session.get("user_id") or SYSTEM,
    )
    return jsonify(w), 201


@bp.route("/matters/<matter_id>/witnesses/<witness_id>", methods=["PATCH"])
async def edit_witness(matter_id: str, witness_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    w = update_witness(
        witness_id = witness_id,
        org_id     = session.get("org") or "",
        actor      = session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"witness_name", "witness_type", "contact_number",
                    "address", "statement_status", "notes"}},
    )
    return jsonify(w)


@bp.route("/matters/<matter_id>/witnesses/<witness_id>", methods=["DELETE"])
async def remove_witness(matter_id: str, witness_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_witness(witness_id, session.get("org") or "",
                   actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Deadlines — Task #142 ───────────────────────────────────────────

@bp.route("/matters/<matter_id>/deadlines", methods=["GET"])
async def list_matter_deadlines(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    dl = get_matter_deadlines(matter_id, session.get("org") or "")
    return jsonify({"deadlines": dl, "priorities": list(DEADLINE_PRIORITIES)})


@bp.route("/matters/<matter_id>/deadlines", methods=["POST"])
async def add_matter_deadline(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "title required"}), 400
    due_date = (data.get("due_date") or "").strip()
    if not due_date:
        return jsonify({"error": "due_date required"}), 400
    dl = create_matter_deadline(
        org_id    = session.get("org") or "",
        matter_id = matter_id,
        title     = title,
        due_date  = due_date,
        priority  = data.get("priority", "Medium"),
        notes     = data.get("notes") or None,
        actor     = session.get("user_id") or SYSTEM,
    )
    return jsonify(dl), 201


@bp.route("/matters/<matter_id>/deadlines/<deadline_id>", methods=["PATCH"])
async def edit_matter_deadline(matter_id: str, deadline_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    dl = update_matter_deadline(
        deadline_id = deadline_id,
        org_id      = session.get("org") or "",
        actor       = session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"title", "due_date", "priority", "notes", "completed", "completed_at"}},
    )
    return jsonify(dl)


@bp.route("/matters/<matter_id>/deadlines/<deadline_id>", methods=["DELETE"])
async def remove_matter_deadline(matter_id: str, deadline_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_deadline(deadline_id, session.get("org") or "",
                           actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Expenses — Task #143 ─────────────────────────────────────────────

@bp.route("/matters/<matter_id>/expenses", methods=["GET"])
async def list_matter_expenses(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    expenses = get_matter_expenses(matter_id, session.get("org") or "")
    return jsonify({"expenses": expenses, "categories": list(EXPENSE_CATEGORIES)})


@bp.route("/matters/<matter_id>/expenses", methods=["POST"])
async def add_matter_expense(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("description") or not str(data.get("description", "")).strip():
        return jsonify({"error": "description required"}), 400
    if data.get("amount_pkr") is None:
        return jsonify({"error": "amount_pkr required"}), 400
    if not data.get("expense_date"):
        return jsonify({"error": "expense_date required"}), 400
    expense = create_matter_expense(
        org_id=session.get("org") or "",
        matter_id=matter_id,
        description=data["description"].strip(),
        amount_pkr=float(data["amount_pkr"]),
        expense_date=data["expense_date"],
        category=data.get("category", "Misc"),
        billable=int(bool(data.get("billable", True))),
        receipt_ref=data.get("receipt_ref"),
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(expense), 201


@bp.route("/matters/<matter_id>/expenses/<expense_id>", methods=["PATCH"])
async def edit_matter_expense(matter_id: str, expense_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_expense(
        expense_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"description", "amount_pkr", "expense_date", "category", "billable", "receipt_ref"}},
    )
    return jsonify(updated)


@bp.route("/matters/<matter_id>/expenses/<expense_id>", methods=["DELETE"])
async def remove_matter_expense(matter_id: str, expense_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_expense(expense_id, session.get("org") or "",
                          actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Correspondence — Task #144 ───────────────────────────────────────

@bp.route("/matters/<matter_id>/correspondence", methods=["GET"])
async def list_matter_correspondence(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    items = get_matter_correspondence(matter_id, session.get("org") or "")
    return jsonify({"correspondence": items, "directions": list(CORR_DIRECTIONS), "types": list(CORR_TYPES)})


@bp.route("/matters/<matter_id>/correspondence", methods=["POST"])
async def add_matter_correspondence(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("subject") or not str(data.get("subject", "")).strip():
        return jsonify({"error": "subject required"}), 400
    if not data.get("corr_date"):
        return jsonify({"error": "corr_date required"}), 400
    item = create_matter_correspondence(
        org_id=session.get("org") or "",
        matter_id=matter_id,
        corr_date=data["corr_date"],
        subject=data["subject"].strip(),
        direction=data.get("direction", "Sent"),
        corr_type=data.get("corr_type", "Letter"),
        party=data.get("party"),
        reference_no=data.get("reference_no"),
        notes=data.get("notes"),
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(item), 201


@bp.route("/matters/<matter_id>/correspondence/<corr_id>", methods=["PATCH"])
async def edit_matter_correspondence(matter_id: str, corr_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_correspondence(
        corr_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"corr_date", "direction", "corr_type", "subject", "party", "reference_no", "notes"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/correspondence/<corr_id>", methods=["DELETE"])
async def remove_matter_correspondence(matter_id: str, corr_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_correspondence(corr_id, session.get("org") or "",
                                 actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Bail & Interim Relief — Task #145 ───────────────────────────────────────

@bp.route("/matters/<matter_id>/relief", methods=["GET"])
async def list_matter_relief(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    items = get_matter_relief(matter_id, session.get("org") or "")
    return jsonify({"relief": items, "types": list(RELIEF_TYPES), "statuses": list(RELIEF_STATUSES)})


@bp.route("/matters/<matter_id>/relief", methods=["POST"])
async def add_matter_relief(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("application_date"):
        return jsonify({"error": "application_date required"}), 400
    item = create_matter_relief(
        org_id=session.get("org") or "",
        matter_id=matter_id,
        application_date=data["application_date"],
        relief_type=data.get("relief_type", "Bail"),
        court=data.get("court"),
        judge=data.get("judge"),
        status=data.get("status", "Pending"),
        conditions=data.get("conditions"),
        surety_amount_pkr=float(data["surety_amount_pkr"]) if data.get("surety_amount_pkr") is not None else None,
        surety_name=data.get("surety_name"),
        notes=data.get("notes"),
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(item), 201


@bp.route("/matters/<matter_id>/relief/<relief_id>", methods=["PATCH"])
async def edit_matter_relief(matter_id: str, relief_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_relief(
        relief_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"application_date", "relief_type", "court", "judge", "status",
                    "conditions", "surety_amount_pkr", "surety_name", "notes"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/relief/<relief_id>", methods=["DELETE"])
async def remove_matter_relief(matter_id: str, relief_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_relief(relief_id, session.get("org") or "",
                         actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Outcome — Task #146 ──────────────────────────────────────────────

@bp.route("/matters/<matter_id>/outcome", methods=["GET"])
async def get_outcome(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    outcome = get_matter_outcome(matter_id, session.get("org") or "")
    return jsonify({"outcome": outcome, "outcome_types": list(OUTCOME_TYPES)})


@bp.route("/matters/<matter_id>/outcome", methods=["PUT"])
async def save_outcome(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    outcome = upsert_matter_outcome(
        org_id=session.get("org") or "",
        matter_id=matter_id,
        outcome_type=data.get("outcome_type", "Pending"),
        disposal_date=data.get("disposal_date") or None,
        court=data.get("court") or None,
        judge=data.get("judge") or None,
        decree_amount_pkr=float(data["decree_amount_pkr"]) if data.get("decree_amount_pkr") is not None else None,
        appeal_filed=int(bool(data.get("appeal_filed", False))),
        appeal_deadline=data.get("appeal_deadline") or None,
        notes=data.get("notes") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(outcome)


# ─── Matter Charges — Task #147 ──────────────────────────────────────────────

@bp.route("/matters/<matter_id>/charges", methods=["GET"])
async def list_matter_charges(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    charges = get_matter_charges(matter_id, session.get("org") or "")
    return jsonify({"charges": charges, "plea_options": list(PLEA_OPTIONS)})


@bp.route("/matters/<matter_id>/charges", methods=["POST"])
async def add_matter_charge(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("section_no") or not str(data.get("section_no", "")).strip():
        return jsonify({"error": "section_no required"}), 400
    charge = create_matter_charge(
        org_id=session.get("org") or "",
        matter_id=matter_id,
        section_no=data["section_no"].strip(),
        description=data.get("description"),
        plea=data.get("plea", "No Plea"),
        charge_framed=int(bool(data.get("charge_framed", False))),
        charge_framed_date=data.get("charge_framed_date") or None,
        court=data.get("court") or None,
        notes=data.get("notes") or None,
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(charge), 201


@bp.route("/matters/<matter_id>/charges/<charge_id>", methods=["PATCH"])
async def edit_matter_charge(matter_id: str, charge_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_charge(
        charge_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        **{k: v for k, v in data.items()
           if k in {"section_no", "description", "plea", "charge_framed",
                    "charge_framed_date", "court", "notes"}},
    )
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/charges/<charge_id>", methods=["DELETE"])
async def remove_matter_charge(matter_id: str, charge_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_charge(charge_id, session.get("org") or "",
                         actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter FIR — Task #148 ───────────────────────────────────────────────────

@bp.route("/documents/extract-fir", methods=["POST"])
async def extract_fir():
    """OCR + AI field extraction for a photographed/scanned FIR or court
    order, to pre-fill the Add FIR form. Beta: accuracy depends heavily on
    document quality/layout, which varies by police station — the lawyer
    always reviews the extracted fields before saving, nothing is written
    to the matter here.
    """
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401

    if not os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE"):
        return jsonify({
            "error": "Document scanning isn't configured yet — AZURE_DOCUMENTINTELLIGENCE_SERVICE "
                     "is not set. Enter the FIR details manually, or ask your admin to configure OCR."
        }), 400

    files = await request.files
    upload = files.get("file")
    if upload is None:
        return jsonify({"error": "No file provided"}), 400

    try:
        raw_text = await _ocr_extract_text(upload.read(), upload.filename or "fir.pdf")
    except Exception as exc:
        logging.exception("FIR OCR failed: %s", exc)
        return jsonify({"error": "Could not read that document — try a clearer photo or scan."}), 400

    if not raw_text.strip():
        return jsonify({"error": "No readable text found in that document."}), 400

    result: dict = {
        "raw_text": raw_text[:4000],
        "fir_number": None, "police_station": None, "district": None,
        "io_name": None, "complainant": None, "arrest_date": None,
        "sections_at_fir": None, "fir_date": None, "accused_name": None,
    }

    if current_app.config.get("AZURE_CONFIGURED"):
        try:
            import re as _re
            openai_client = current_app.config[CONFIG_OPENAI_CLIENT]
            system_msg = (
                "You extract structured fields from OCR text of a Pakistani police FIR "
                "(First Information Report) or a related court order. Return ONLY a JSON object "
                "with keys: \"fir_number\", \"police_station\", \"district\", \"io_name\" "
                "(investigating officer), \"complainant\", \"arrest_date\" (ISO YYYY-MM-DD if found), "
                "\"sections_at_fir\" (e.g. \"302, 324, 34 PPC\"), \"fir_date\" (ISO YYYY-MM-DD if found), "
                "\"accused_name\". Use null for any field not clearly present in the text — "
                "never guess or infer a value that isn't actually written there."
            )
            ai_resp = await openai_client.chat.completions.create(
                model=current_app.config.get("OPENAI_CHATGPT_MODEL", "gpt-5-mini-1"),
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": raw_text[:6000]},
                ],
                temperature=0.0,
                max_tokens=500,
            )
            raw = (ai_resp.choices[0].message.content or "").strip()
            raw = _re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=_re.MULTILINE).strip()
            extracted = json.loads(raw)
            for key in ("fir_number", "police_station", "district", "io_name", "complainant",
                        "arrest_date", "sections_at_fir", "fir_date", "accused_name"):
                if extracted.get(key):
                    result[key] = extracted[key]
        except Exception as exc:
            logging.warning("FIR field extraction failed, returning raw OCR text only: %s", exc)

    return jsonify(result)


@bp.route("/matters/<matter_id>/fir", methods=["GET"])
async def list_matter_fir(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    records = get_matter_fir(matter_id, session.get("org") or "")
    return jsonify({"fir": records})


@bp.route("/matters/<matter_id>/fir", methods=["POST"])
async def add_matter_fir(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("fir_number") or not data.get("police_station"):
        return jsonify({"error": "fir_number and police_station are required"}), 400
    record = create_matter_fir(
        matter_id=matter_id,
        org_id=session.get("org") or "",
        data=data,
        actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(record), 201


@bp.route("/matters/<matter_id>/fir/<fir_id>", methods=["PATCH"])
async def edit_matter_fir(matter_id: str, fir_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_fir(fir_id, session.get("org") or "", data,
                                actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/fir/<fir_id>", methods=["DELETE"])
async def remove_matter_fir(matter_id: str, fir_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_fir(fir_id, session.get("org") or "",
                      actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Matter Challan — Task #149 ───────────────────────────────────────────────

@bp.route("/matters/<matter_id>/challan", methods=["GET"])
async def list_matter_challan(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    records = get_matter_challan(matter_id, session.get("org") or "")
    return jsonify({"challan": records, "challan_types": list(CHALLAN_TYPES), "challan_statuses": list(CHALLAN_STATUSES)})


@bp.route("/matters/<matter_id>/challan", methods=["POST"])
async def add_matter_challan(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    record = create_matter_challan(
        matter_id=matter_id, org_id=session.get("org") or "",
        data=data, actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(record), 201


@bp.route("/matters/<matter_id>/challan/<challan_id>", methods=["PATCH"])
async def edit_matter_challan(matter_id: str, challan_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_challan(challan_id, session.get("org") or "", data,
                                    actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/challan/<challan_id>", methods=["DELETE"])
async def remove_matter_challan(matter_id: str, challan_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_challan(challan_id, session.get("org") or "",
                          actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Conflict of Interest Check — Task #150 ───────────────────────────────────

@bp.route("/conflicts/check", methods=["POST"])
async def check_conflicts():
    """
    Given a new_client_name and opponent_name, return any existing matters where
    either name appears as the client or as the opposing_party.
    """
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    new_client = (data.get("new_client_name") or "").strip().lower()
    opponent   = (data.get("opponent_name") or "").strip().lower()
    org_id = session.get("org") or ""

    if not new_client and not opponent:
        return jsonify({"conflicts": [], "checked": False})

    all_clients = get_clients(org_id)
    all_matters = get_matters(org_id)

    # Build client_id → name map
    client_map = {c["client_id"]: c["name"].lower() for c in all_clients}

    conflicts = []
    for m in all_matters:
        client_name = client_map.get(m.get("client_id") or "", "")
        opp = (m.get("opposing_party") or "").lower()
        match_reason = []
        # New client name matches an existing client in any matter → they may already be on file
        if new_client and new_client in client_name:
            match_reason.append(f"'{data.get('new_client_name')}' is already a client in this matter")
        # New client name matches the opponent in an existing matter
        if new_client and opp and new_client in opp:
            match_reason.append(f"'{data.get('new_client_name')}' appears as opposing party in this matter")
        # Opponent name matches an existing client in any matter
        if opponent and client_name and opponent in client_name:
            match_reason.append(f"'{data.get('opponent_name')}' is an existing client of the firm")
        # Opponent name matches the opposing party in another matter (same counsel, fine — but flag for awareness)
        if opponent and opp and opponent in opp:
            match_reason.append(f"'{data.get('opponent_name')}' appears as opposing party in another matter — recurring opponent")
        if match_reason:
            conflicts.append({
                "matter_id":    m.get("matter_id"),
                "matter_title": m.get("title"),
                "client_name":  next((c["name"] for c in all_clients if c["client_id"] == m.get("client_id")), ""),
                "opposing_party": m.get("opposing_party"),
                "status":       m.get("status"),
                "reasons":      match_reason,
            })

    return jsonify({"conflicts": conflicts, "checked": True})


# ─── Court Fee Calculator — Task #152 ────────────────────────────────────────

@bp.route("/matters/<matter_id>/court-fees", methods=["GET"])
async def list_court_fees(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    records = get_court_fee_payments(matter_id, session.get("org") or "")
    return jsonify({"payments": records, "fee_types": list(COURT_FEE_TYPES)})


@bp.route("/court-fees/calculate", methods=["POST"])
async def calculate_court_fee():
    """Quick calculator — no DB write."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    claim  = float(data.get("claim_amount_pkr", 0) or 0)
    ftype  = data.get("fee_type", "Ad Valorem")
    result = compute_court_fee(claim, ftype)
    return jsonify({"calculated_fee": result, "claim_amount_pkr": claim, "fee_type": ftype})


@bp.route("/matters/<matter_id>/court-fees", methods=["POST"])
async def add_court_fee(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    record = create_court_fee_payment(
        matter_id=matter_id, org_id=session.get("org") or "",
        data=data, actor=session.get("user_id") or SYSTEM,
    )
    return jsonify(record), 201


@bp.route("/matters/<matter_id>/court-fees/<fp_id>", methods=["PATCH"])
async def edit_court_fee(matter_id: str, fp_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_court_fee_payment(fp_id, session.get("org") or "", data,
                                       actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/court-fees/<fp_id>", methods=["DELETE"])
async def remove_court_fee(matter_id: str, fp_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_court_fee_payment(fp_id, session.get("org") or "",
                             actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Associate Fees — Task #153 ──────────────────────────────────────────────

@bp.route("/matters/<matter_id>/associate-fees", methods=["GET"])
async def list_associate_fees(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"fees": get_associate_fees(matter_id, session.get("org") or "")})


@bp.route("/matters/<matter_id>/associate-fees", methods=["POST"])
async def add_associate_fee(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("advocate_name"):
        return jsonify({"error": "advocate_name is required"}), 400
    record = create_associate_fee(matter_id, session.get("org") or "", data,
                                  actor=session.get("user_id") or SYSTEM)
    return jsonify(record), 201


@bp.route("/matters/<matter_id>/associate-fees/<af_id>", methods=["PATCH"])
async def edit_associate_fee(matter_id: str, af_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_associate_fee(af_id, session.get("org") or "", data,
                                   actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/associate-fees/<af_id>", methods=["DELETE"])
async def remove_associate_fee(matter_id: str, af_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_associate_fee(af_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


@bp.route("/associate-fees/summary", methods=["GET"])
async def associate_fees_summary():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"summary": get_associate_fees_summary(session.get("org") or "")})


# ─── Client Trust Ledger — Task #154 ─────────────────────────────────────────

@bp.route("/clients/<client_id>/trust-ledger", methods=["GET"])
async def list_trust_ledger(client_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    entries = get_trust_ledger(client_id, session.get("org") or "")
    balance = entries[-1]["balance_pkr"] if entries else 0.0
    return jsonify({"entries": entries, "balance": balance, "txn_types": list(TRUST_TXN_TYPES)})


@bp.route("/clients/<client_id>/trust-ledger", methods=["POST"])
async def add_trust_entry(client_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("description"):
        return jsonify({"error": "description is required"}), 400
    if not data.get("txn_date"):
        return jsonify({"error": "txn_date is required"}), 400
    record = create_trust_entry(client_id, session.get("org") or "", data,
                                actor=session.get("user_id") or SYSTEM)
    return jsonify(record), 201


@bp.route("/clients/<client_id>/trust-ledger/<ledger_id>", methods=["PATCH"])
async def edit_trust_entry(client_id: str, ledger_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_trust_entry(ledger_id, session.get("org") or "", data,
                                 actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/clients/<client_id>/trust-ledger/<ledger_id>", methods=["DELETE"])
async def remove_trust_entry(client_id: str, ledger_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_trust_entry(ledger_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Cheque Tracker — Task #155 ──────────────────────────────────────────────

@bp.route("/matters/<matter_id>/cheques", methods=["GET"])
async def list_matter_cheques(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({
        "cheques": get_matter_cheques(matter_id, session.get("org") or ""),
        "cheque_types": list(CHEQUE_TYPES),
        "cheque_statuses": list(CHEQUE_STATUSES),
    })


@bp.route("/matters/<matter_id>/cheques", methods=["POST"])
async def add_matter_cheque(matter_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("cheque_no"):
        return jsonify({"error": "cheque_no is required"}), 400
    record = create_matter_cheque(matter_id, session.get("org") or "", data,
                                  actor=session.get("user_id") or SYSTEM)
    return jsonify(record), 201


@bp.route("/matters/<matter_id>/cheques/<cheque_id>", methods=["PATCH"])
async def edit_matter_cheque(matter_id: str, cheque_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    updated = update_matter_cheque(cheque_id, session.get("org") or "", data,
                                   actor=session.get("user_id") or SYSTEM)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/cheques/<cheque_id>", methods=["DELETE"])
async def remove_matter_cheque(matter_id: str, cheque_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_matter_cheque(cheque_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── Vakalatnama Register — Task #156 ────────────────────────────────────────

@bp.route("/vakalatnama-register", methods=["GET"])
async def vakalatnama_register():
    """Cross-matter vakalatnama status register."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    matters = get_matters(org_id)
    # Attach client names
    clients = {c["client_id"]: c["name"] for c in get_clients(org_id)}
    register = []
    for m in matters:
        register.append({
            "matter_id":          m["matter_id"],
            "title":              m["title"],
            "matter_no":          m.get("matter_no"),
            "client_name":        clients.get(m.get("client_id", ""), "Unknown"),
            "court_name":         m.get("court_name"),
            "vakalatnama_status": m.get("vakalatnama_status", "Pending"),
            "status":             m.get("status"),
            "created_at":         m.get("created_at"),
        })
    return jsonify({"register": register})


@bp.route("/matters/<matter_id>/vakalatnama", methods=["PATCH"])
async def update_vakalatnama_status(matter_id: str):
    """Quick-update vakalatnama_status on a matter."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    new_status = data.get("vakalatnama_status")
    if new_status not in ("Pending", "Filed", "Rejected"):
        return jsonify({"error": "Invalid status"}), 400
    return jsonify(update_matter(
        matter_id, session.get("org") or "",
        actor=session.get("user_id") or SYSTEM,
        vakalatnama_status=new_status,
    ))


# ─── WHT Invoice Preview — Task #157 ─────────────────────────────────────────

@bp.route("/invoices/wht-preview", methods=["POST"])
async def wht_preview():
    """Calculate WHT for a gross amount + client_type without writing to DB."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    gross = float(data.get("gross", 0))
    client_type = data.get("client_type", "Individual")
    rate, wht_amount, net_payable = compute_wht(gross, client_type)
    return jsonify({
        "gross":       gross,
        "client_type": client_type,
        "wht_rate":    rate,
        "wht_amount":  wht_amount,
        "net_payable": net_payable,
        "note": "6% WHT applies to Corporate clients per Income Tax Ordinance 2001 §153" if rate > 0 else "No WHT — individual client",
    })


# ─── Intelligence Notes — Task #158 ─────────────────────────────────────────

@bp.route("/opposing-counsel", methods=["GET"])
async def list_opposing_counsel():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"counsel": get_opposing_counsel(session.get("org") or "")})


@bp.route("/opposing-counsel", methods=["POST"])
async def add_opposing_counsel():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("name"):
        return jsonify({"error": "name is required"}), 400
    record = create_opposing_counsel(session.get("org") or "", data, actor=session.get("user_id") or SYSTEM)
    return jsonify(record), 201


@bp.route("/opposing-counsel/<counsel_id>", methods=["PATCH"])
async def edit_opposing_counsel(counsel_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    return jsonify(update_opposing_counsel(counsel_id, session.get("org") or "", data, actor=session.get("user_id") or SYSTEM))


@bp.route("/opposing-counsel/<counsel_id>", methods=["DELETE"])
async def remove_opposing_counsel(counsel_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_opposing_counsel(counsel_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


@bp.route("/judge-notes", methods=["GET"])
async def list_judge_notes():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"judges": get_judge_notes(session.get("org") or "")})


@bp.route("/judge-notes", methods=["POST"])
async def add_judge_note():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    if not data.get("name"):
        return jsonify({"error": "name is required"}), 400
    record = create_judge_note(session.get("org") or "", data, actor=session.get("user_id") or SYSTEM)
    return jsonify(record), 201


@bp.route("/judge-notes/<judge_id>", methods=["PATCH"])
async def edit_judge_note(judge_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    return jsonify(update_judge_note(judge_id, session.get("org") or "", data, actor=session.get("user_id") or SYSTEM))


@bp.route("/judge-notes/<judge_id>/stats", methods=["GET"])
async def judge_track_record(judge_id: str):
    """Firm's own historical track record with this judge — computed live
    from this org's own hearings/bail_bonds. Not external/published data."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    judge = get_judge_note(judge_id, org_id)
    if not judge:
        return jsonify({"error": "Not found"}), 404
    return jsonify(get_judge_track_record(org_id, judge.get("name") or ""))


@bp.route("/judge-notes/<judge_id>", methods=["DELETE"])
async def remove_judge_note(judge_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_judge_note(judge_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ─── LHC Case Status Live Lookup — Task #159 ─────────────────────────────────
# Replace LHC_CASE_STATUS_URL with the actual Lahore High Court e-filing / case
# status endpoint when available.  The form field names (case_no_field etc.) must
# also be updated to match the live page's HTML form.
LHC_CASE_STATUS_URL = "PLACEHOLDER_LHC_CASE_STATUS_URL"
LHC_CASE_NO_FIELD   = "PLACEHOLDER_CASE_NO_FIELD"        # <input name="?"> in the LHC form

import aiohttp


@bp.route("/lhc/case-status", methods=["GET"])
async def lhc_case_status():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401

    case_no = request.args.get("case_no", "").strip()
    if not case_no:
        return jsonify({"error": "case_no is required"}), 400

    if LHC_CASE_STATUS_URL == "PLACEHOLDER_LHC_CASE_STATUS_URL":
        return jsonify({
            "status":  "unavailable",
            "message": "LHC_CASE_STATUS_URL placeholder not yet configured. "
                       "Replace LHC_CASE_STATUS_URL and LHC_CASE_NO_FIELD in app.py "
                       "with the actual Lahore High Court case-status portal URL and form field name.",
            "case_no": case_no,
        }), 503

    try:
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as http:
            async with http.post(
                LHC_CASE_STATUS_URL,
                data={LHC_CASE_NO_FIELD: case_no},
                headers={"User-Agent": "Mozilla/5.0 (compatible; ProjectEase/1.0)"},
            ) as resp:
                html = await resp.text()

        # ── Parse with BeautifulSoup ──────────────────────────────────────────
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")

        # Extract whatever the LHC portal returns — adjust selectors to match
        # the actual page structure once the URL is configured.
        result_text = soup.get_text(separator=" ", strip=True)[:2000]

        return jsonify({
            "status":  "ok",
            "case_no": case_no,
            "raw_text": result_text,
            "note": "Selectors may need adjustment once LHC_CASE_STATUS_URL is set.",
        })

    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc), "case_no": case_no}), 502


# ─── PROJECT EASE: Plan & Upgrade API ───────────────────────────────────────

# Bank transfer details — fill these in .env before going live (see PLACEHOLDERS.md)
_BANK_NAME    = __import__("os").getenv("BANK_NAME",    "Your Bank Name")
_BANK_ACCOUNT = __import__("os").getenv("BANK_ACCOUNT", "0000-0000-0000-0000")
_BANK_IBAN    = __import__("os").getenv("BANK_IBAN",    "PK00XXXX0000000000000000")
_BANK_TITLE   = __import__("os").getenv("BANK_TITLE",   "Project Ease Pvt Ltd")
_SUPPORT_WA   = __import__("os").getenv("SUPPORT_WHATSAPP", "+92-300-0000000")


@bp.route("/plan-config", methods=["GET"])
async def get_plan_config():
    """Return plan limits and pricing for the frontend."""
    session = _get_session()
    if not session:
        return jsonify({"error": "Unauthorized"}), 401
    org    = get_org(session.get("org") or "")
    plan   = (org or {}).get("plan", "trial")
    return jsonify({
        "plans":        PLAN_CONFIG,
        "current_plan": plan,
        "bank": {
            "name":    _BANK_NAME,
            "account": _BANK_ACCOUNT,
            "iban":    _BANK_IBAN,
            "title":   _BANK_TITLE,
        },
        "support_whatsapp": _SUPPORT_WA,
    })


@bp.route("/upgrade-request", methods=["POST"])
async def submit_upgrade_request():
    """Owner submits an upgrade request after making a bank transfer."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401

    data           = await request.get_json(silent=True) or {}
    requested_plan = (data.get("requested_plan") or "").strip().lower()
    payment_ref    = (data.get("payment_ref") or "").strip() or None
    notes          = (data.get("notes") or "").strip() or None

    if requested_plan not in PLAN_CONFIG:
        return jsonify({"error": "Invalid plan"}), 400

    org = get_org(session.get("org") or "")
    if not org:
        return jsonify({"error": "Organization not found"}), 404

    current_plan = org.get("plan", "trial")
    if requested_plan == current_plan:
        return jsonify({"error": "You are already on this plan"}), 400

    req = create_upgrade_request(
        org_id        = org["org_id"],
        current_plan  = current_plan,
        requested_plan= requested_plan,
        payment_ref   = payment_ref,
        notes         = notes,
    )
    _audit(session, "upgrade_request",
           resource_type="plan", resource_name=requested_plan,
           details={"from": current_plan, "to": requested_plan, "payment_ref": payment_ref})

    # Fire-and-forget acknowledgement email
    try:
        import asyncio as _asyncio
        from email_helper import send_upgrade_request_received as _send_ack
        owner = next(
            (u for u in get_users_for_org(org["org_id"]) if u["role"] == "org_owner"),
            None
        )
        if owner:
            await _asyncio.to_thread(_send_ack, owner["email"], org["name"], requested_plan)
    except Exception:
        pass

    return jsonify(req), 201


@bp.route("/admin/upgrade-requests", methods=["GET"])
async def admin_list_upgrade_requests():
    err = _require_platform_admin()
    if err:
        return err
    status = request.args.get("status") or None
    return jsonify({"requests": get_upgrade_requests(status=status)})


@bp.route("/admin/upgrade-requests/<request_id>/approve", methods=["PATCH"])
async def admin_approve_upgrade(request_id: str):
    import asyncio as _asyncio
    err = _require_platform_admin()
    if err:
        return err
    admin_session = _get_session()
    resolver = (admin_session or {}).get("email") or SYSTEM

    result = resolve_upgrade_request(request_id, "approved", resolver)
    if not result:
        return jsonify({"error": "Request not found"}), 404

    # Fire-and-forget approval email
    try:
        from email_helper import send_upgrade_approved as _send_upgrade
        org = get_org(result["org_id"])
        if org:
            owner = next(
                (u for u in get_users_for_org(result["org_id"]) if u["role"] == "org_owner"),
                None
            )
            if owner:
                await _asyncio.to_thread(
                    _send_upgrade, owner["email"], org["name"], result["requested_plan"]
                )
    except Exception:
        pass

    log_event("upgrade_approved", org_id=result["org_id"],
              actor_name=resolver, details={"plan": result["requested_plan"]})
    return jsonify(result)


@bp.route("/admin/upgrade-requests/<request_id>/reject", methods=["PATCH"])
async def admin_reject_upgrade(request_id: str):
    err = _require_platform_admin()
    if err:
        return err
    admin_session = _get_session()
    resolver = (admin_session or {}).get("email") or SYSTEM

    result = resolve_upgrade_request(request_id, "rejected", resolver)
    if not result:
        return jsonify({"error": "Request not found"}), 404

    log_event("upgrade_rejected", org_id=result["org_id"],
              actor_name=resolver, details={"plan": result["requested_plan"]})
    return jsonify(result)


# ─── PROJECT EASE: Audit Log API ─────────────────────────────────────────────

@bp.route("/audit-logs", methods=["GET"])
async def list_audit_logs():
    """Return paginated audit logs scoped to the caller's org (owner) or all orgs (admin)."""
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "platform_admin"):
        return jsonify({"error": "Unauthorized"}), 401

    role = session.get("role")
    if role not in ("org_owner", "platform_admin"):
        # Log the access attempt and reject
        _audit(session, "access_denied",
               details={"route": "/audit-logs", "reason": "insufficient_role"})
        return jsonify({"error": "Forbidden"}), 403

    # Platform admins can query across all orgs; org_owners are always scoped to their own
    org_id = None if role == "platform_admin" else session.get("org")
    # Admin may optionally filter by org
    if role == "platform_admin":
        org_id = request.args.get("org_id") or None

    event_type = request.args.get("event_type") or None
    user_id    = request.args.get("user_id")    or None
    date_from  = request.args.get("date_from")  or None
    date_to    = request.args.get("date_to")    or None
    limit      = min(int(request.args.get("limit",  "200")), 500)
    offset     = int(request.args.get("offset", "0"))

    logs  = get_audit_logs(org_id=org_id, event_type=event_type, user_id=user_id,
                           date_from=date_from, date_to=date_to, limit=limit, offset=offset)
    total = count_audit_logs(org_id=org_id, event_type=event_type, user_id=user_id,
                             date_from=date_from, date_to=date_to)
    return jsonify({"logs": logs, "total": total, "limit": limit, "offset": offset})


# ─── PROJECT EASE: Fees & Invoices API ───────────────────────────────────────


@bp.route("/fees", methods=["GET"])
async def list_fees():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id    = session.get("org") or ""
    matter_id = request.args.get("matter_id") or None
    return jsonify(get_fees(org_id, matter_id=matter_id))


@bp.route("/fees", methods=["POST"])
async def add_fee():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    data        = await request.get_json(silent=True) or {}
    description = (data.get("description") or "").strip()
    fee_date    = (data.get("fee_date") or "").strip()
    amount      = int(data.get("amount") or 0)
    if not description or not fee_date:
        return jsonify({"error": "description and fee_date are required"}), 400
    if amount < 0:
        return jsonify({"error": "amount must be non-negative"}), 400
    org_id = session.get("org") or ""
    actor  = session.get("user_id") or SYSTEM
    f = create_fee(
        org_id=org_id, description=description, fee_date=fee_date, amount=amount,
        matter_id = data.get("matter_id") or None,
        fee_type  = data.get("fee_type")  or "Consultation",
        notes     = data.get("notes")     or None,
        actor=actor,
    )
    _audit(session, "fee_create",
           resource_type="fee", resource_name=description,
           details={"amount": amount, "matter_id": data.get("matter_id")})
    return jsonify(f), 201


@bp.route("/fees/<fee_id>", methods=["PATCH"])
async def edit_fee(fee_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    f = get_fee(fee_id)
    if not f or f.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    # Coerce amount to int if present
    if "amount" in data:
        data["amount"] = int(data["amount"] or 0)
    # Auto-set paid_at when marking paid
    if data.get("is_paid") and not f.get("paid_at"):
        import datetime as _dt
        data["paid_at"] = _dt.datetime.utcnow().strftime("%Y-%m-%d")
    updated = update_fee(fee_id, session.get("org") or "", actor=actor, **data)
    return jsonify(updated)


@bp.route("/fees/<fee_id>", methods=["DELETE"])
async def remove_fee(fee_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    f = get_fee(fee_id)
    if not f or f.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    actor = session.get("user_id") or SYSTEM
    delete_fee(fee_id, session.get("org") or "", actor=actor)
    return jsonify({"ok": True})


@bp.route("/invoices", methods=["GET"])
async def list_invoices():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id    = session.get("org") or ""
    matter_id = request.args.get("matter_id") or None
    return jsonify(get_invoices(org_id, matter_id=matter_id))


@bp.route("/invoices", methods=["POST"])
async def add_invoice():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    data      = await request.get_json(silent=True) or {}
    matter_id = (data.get("matter_id") or "").strip()
    title     = (data.get("title") or "").strip()
    if not matter_id or not title:
        return jsonify({"error": "matter_id and title are required"}), 400
    import datetime as _dt
    org_id      = session.get("org") or ""
    actor       = session.get("user_id") or SYSTEM
    issued_date = data.get("issued_date") or _dt.datetime.utcnow().strftime("%Y-%m-%d")
    inv = create_invoice(
        org_id=org_id, matter_id=matter_id, title=title, issued_date=issued_date,
        client_id = data.get("client_id") or None,
        due_date  = data.get("due_date")  or None,
        notes     = data.get("notes")     or None,
        actor=actor,
    )
    _audit(session, "invoice_create",
           resource_type="invoice", resource_name=title,
           details={"matter_id": matter_id, "total": inv.get("total_amount")})
    return jsonify(inv), 201


@bp.route("/invoices/<invoice_id>", methods=["GET"])
async def get_invoice(invoice_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    inv = get_invoice_with_fees(invoice_id)
    if not inv or inv.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    return jsonify(inv)


@bp.route("/invoices/<invoice_id>", methods=["PATCH"])
async def edit_invoice(invoice_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    inv = get_invoice_with_fees(invoice_id)
    if not inv or inv.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    updated = update_invoice(invoice_id, session.get("org") or "", actor=actor, **data)
    return jsonify(updated)


# ─── PROJECT EASE: Court Calendar API ───────────────────────────────────────


@bp.route("/hearings", methods=["GET"])
async def list_hearings():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id    = session.get("org") or ""
    from_date = request.args.get("from_date") or None
    to_date   = request.args.get("to_date")   or None
    return jsonify(get_hearings(org_id, from_date=from_date, to_date=to_date))


@bp.route("/hearings", methods=["POST"])
async def add_hearing():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    title        = (data.get("title") or "").strip()
    hearing_date = (data.get("hearing_date") or "").strip()
    if not title or not hearing_date:
        return jsonify({"error": "title and hearing_date are required"}), 400
    org_id = session.get("org") or ""
    actor  = session.get("user_id") or SYSTEM
    h = create_hearing(
        org_id=org_id, title=title, hearing_date=hearing_date,
        matter_id   = data.get("matter_id")    or None,
        hearing_time= data.get("hearing_time") or None,
        court_name  = data.get("court_name")   or None,
        judge_name  = data.get("judge_name")   or None,
        notes       = data.get("notes")        or None,
        wa_reminder = bool(data.get("wa_reminder", False)),
        assigned_to = data.get("assigned_to")   or None,
        actor=actor,
    )
    _audit(session, "hearing_create",
           resource_type="hearing", resource_name=title,
           details={"date": hearing_date, "matter_id": data.get("matter_id")})
    return jsonify(h), 201


@bp.route("/hearings/<hearing_id>", methods=["GET"])
async def get_hearing_detail(hearing_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    h = get_hearing(hearing_id)
    if not h or h.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    return jsonify(h)


async def _notify_on_hearing_outcome(session, org_id: str, hearing: dict) -> None:
    """Associate dispatch: when a hearing's outcome is marked, WhatsApp the
    org owner (so they see it without calling the associate) and, if the
    hearing is matter-linked, the client too — same message pipeline as
    Court Orders. Fire-and-forget, never raises."""
    outcome = (hearing.get("hearing_outcome") or "").strip()
    if not outcome:
        return
    try:
        from whatsapp_helper import send_whatsapp_text, normalize_pk_number
    except ImportError:
        return
    import asyncio as _aio

    title = hearing.get("matter_title") or hearing.get("title") or "the hearing"
    next_date = (hearing.get("hearing_date") or "")  # informational only here

    # Notify the org owner (skip if the owner themself made the edit)
    owner = get_org_owner_contact(org_id)
    if owner and owner.get("whatsapp_number") and owner.get("user_id") != session.get("user_id"):
        owner_phone = normalize_pk_number(owner["whatsapp_number"])
        if owner_phone:
            assignee_name = session.get("name") or "Your associate"
            msg = f'📋 {assignee_name} marked "{title}" as *{outcome}* today.'
            if hearing.get("adj_reason"):
                msg += f"\nReason: {hearing['adj_reason']}"
            result = await _aio.to_thread(send_whatsapp_text, owner_phone, msg)
            _audit(session, "owner_whatsapp_notify", resource_type="hearing",
                   resource_id=hearing.get("hearing_id"), resource_name=title,
                   details={"outcome": outcome, "sent": result.get("sent")})

    # Notify the client, if this hearing is linked to a matter
    matter_id = hearing.get("matter_id")
    if matter_id:
        contact = get_matter_client_contact(matter_id, org_id)
        if contact and contact.get("client_phone"):
            phone = normalize_pk_number(contact["client_phone"])
            if phone:
                msg = f'Your case "{contact.get("title") or title}" was heard today.'
                if outcome == "Adjourned":
                    msg += " Adjourned."
                elif outcome:
                    msg += f" {outcome}."
                result = await _aio.to_thread(send_whatsapp_text, phone, msg)
                _audit(session, "client_whatsapp_notify", resource_type="matter",
                       resource_id=matter_id, resource_name=contact.get("title"),
                       details={"outcome": outcome, "sent": result.get("sent"), "via": "hearing"})


@bp.route("/hearings/<hearing_id>", methods=["PATCH"])
async def edit_hearing(hearing_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    h = get_hearing(hearing_id)
    if not h or h.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    org_id = session.get("org") or ""
    updated = update_hearing(hearing_id, org_id, actor=actor, **data)
    if updated and "hearing_outcome" in data and data.get("notify", True):
        await _notify_on_hearing_outcome(session, org_id, updated)
    return jsonify(updated)


@bp.route("/hearings/<hearing_id>", methods=["DELETE"])
async def remove_hearing(hearing_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    h = get_hearing(hearing_id)
    if not h or h.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    actor = session.get("user_id") or SYSTEM
    delete_hearing(hearing_id, session.get("org") or "", actor=actor)
    _audit(session, "hearing_delete", resource_type="hearing", resource_name=h.get("title", ""))
    return jsonify({"ok": True})


@bp.route("/calendar/notify-holiday/preview", methods=["GET"])
async def preview_holiday_notify():
    """Which clients have a hearing/deadline in this date range — shown
    before sending the bulk court-holiday WhatsApp blast, so the lawyer
    sees exactly who's about to get messaged."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    from_date = request.args.get("from_date") or ""
    to_date   = request.args.get("to_date") or from_date
    if not from_date:
        return jsonify({"error": "from_date is required"}), 400
    clients = get_clients_with_hearings_in_range(session.get("org") or "", from_date, to_date)
    return jsonify({"clients": clients})


@bp.route("/calendar/notify-holiday", methods=["POST"])
async def send_holiday_notify():
    """Bulk WhatsApp: one click notifies every client with a hearing/deadline
    in the given date range that the court is closed. Sends individually
    (not a broadcast list) and reports per-send success/failure — never
    silent about partial failures."""
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data      = await request.get_json(silent=True) or {}
    from_date = (data.get("from_date") or "").strip()
    to_date   = (data.get("to_date") or from_date).strip()
    message   = (data.get("message") or "").strip()
    if not from_date:
        return jsonify({"error": "from_date is required"}), 400
    org_id = session.get("org") or ""
    clients = get_clients_with_hearings_in_range(org_id, from_date, to_date)
    if not clients:
        return jsonify({"notified": 0, "failed": 0, "skipped_no_phone": 0, "details": []})

    try:
        from whatsapp_helper import send_whatsapp_text, normalize_pk_number
    except ImportError:
        return jsonify({"error": "whatsapp_helper not available"}), 500

    import asyncio as _aio
    date_label = from_date if from_date == to_date else f"{from_date} to {to_date}"
    default_msg = (
        f"📢 *Court Holiday Notice*\n\nPlease note: the court is closed {date_label}. "
        "Any hearings scheduled in this period will be rescheduled — your lawyer's office "
        "will confirm the new date shortly."
    )

    notified, failed, skipped = 0, 0, 0
    details = []
    for c in clients:
        phone = normalize_pk_number(c.get("client_phone") or "")
        if not phone:
            skipped += 1
            details.append({"client_name": c.get("client_name"), "sent": False, "reason": "no_valid_phone"})
            continue
        msg = message or default_msg
        result = await _aio.to_thread(send_whatsapp_text, phone, msg)
        if result.get("sent"):
            notified += 1
        else:
            failed += 1
        details.append({"client_name": c.get("client_name"), "sent": result.get("sent"), "reason": result.get("error")})

    _audit(session, "bulk_holiday_notify", details={"from_date": from_date, "to_date": to_date,
                                                     "notified": notified, "failed": failed, "skipped": skipped})
    return jsonify({"notified": notified, "failed": failed, "skipped_no_phone": skipped, "details": details})


@bp.route("/diary/<date_str>", methods=["GET"])
async def get_diary(date_str: str):
    """Task #161: Return all hearings + deadlines for a given date (YYYY-MM-DD)."""
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    hearings = get_hearings(org_id, from_date=date_str, to_date=date_str)
    deadlines = get_deadlines(org_id, from_date=date_str, to_date=date_str)
    return jsonify({"date": date_str, "hearings": hearings, "deadlines": deadlines})


@bp.route("/diary/send-brief", methods=["POST"])
async def send_diary_morning_brief():
    """Task #172: Send daily court diary as a WhatsApp morning brief.

    Body: { to_number: "+923001234567", date: "YYYY-MM-DD" }
    The number must include the country code.  Twilio credentials are resolved
    from TWILIO_ACCOUNT_SID / TWILIO_AUTH_TOKEN / TWILIO_WA_FROM environment
    variables (see PLACEHOLDERS.md).
    """
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401

    body       = await request.get_json() or {}
    to_number  = (body.get("to_number") or "").strip()
    date_str   = (body.get("date") or "").strip()

    if not to_number:
        return jsonify({"error": "to_number is required"}), 400
    if not date_str:
        import datetime as _dt
        date_str = _dt.date.today().isoformat()

    org_id    = session.get("org") or ""
    hearings  = get_hearings(org_id,  from_date=date_str, to_date=date_str)
    deadlines = get_deadlines(org_id, from_date=date_str, to_date=date_str)

    # Format brief text
    try:
        import datetime as _dt
        d_obj = _dt.date.fromisoformat(date_str)
        day_label = d_obj.strftime("%A, %d %B %Y")
    except Exception:
        day_label = date_str

    lines: list[str] = [f"📅 *Morning Brief — {day_label}*\n"]

    if hearings:
        lines.append(f"⚖️ *Court Hearings ({len(hearings)})*")
        for h in hearings:
            time_part = f"{h['hearing_time']} — " if h.get("hearing_time") else ""
            court     = f" @ {h['court_name']}"   if h.get("court_name") else ""
            matter    = f" [{h['matter_title']}]" if h.get("matter_title") else ""
            lines.append(f"• {time_part}{h['title']}{matter}{court}")
        lines.append("")

    if deadlines:
        lines.append(f"⏰ *Deadlines ({len(deadlines)})*")
        for dl in deadlines:
            matter = f" [{dl['matter_title']}]" if dl.get("matter_title") else ""
            lines.append(f"• {dl['title']}{matter}")
        lines.append("")

    if not hearings and not deadlines:
        lines.append("✅ No hearings or deadlines today — clear day!")

    lines.append("_Sent via Project Ease_")
    message = "\n".join(lines)

    # Send via Twilio WhatsApp (placeholder if helper not yet configured)
    try:
        from whatsapp_helper import send_whatsapp_text  # type: ignore[import]
        # Normalise number — Twilio expects "whatsapp:+92..." format internally
        normalised = to_number if to_number.startswith("+") else f"+{to_number}"
        import asyncio as _aio
        result = await _aio.to_thread(send_whatsapp_text, normalised, message)
        sent = bool(result.get("sent")) if isinstance(result, dict) else bool(result)
        if sent:
            _audit(session, "diary_brief_sent", resource_type="diary", resource_name=date_str)
            return jsonify({"sent": True, "to": normalised, "date": date_str})
        reason = result.get("error") if isinstance(result, dict) else "send_failed"
        return jsonify({"sent": False, "reason": reason or "send_failed", "message": message})
    except ImportError:
        # Twilio not yet configured — return the formatted text so the
        # caller can fall back to the WhatsApp share link.
        return jsonify({"sent": False, "reason": "whatsapp_not_configured", "message": message})
    except Exception as exc:
        logging.getLogger(__name__).error("send_diary_brief error: %s", exc)
        return jsonify({"error": str(exc)}), 500


# ── Feature Flags — Task #162 ─────────────────────────────────────────────────

@bp.route("/org-flags", methods=["GET"])
async def read_org_flags():
    """Owner/employee reads their own org's feature flags."""
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    flags = get_org_flags(session.get("org") or "")
    return jsonify({"flags": flags, "feature_labels": FEATURE_LABELS})


@bp.route("/admin/org-flags", methods=["GET"])
async def admin_list_org_flags_route():
    """Admin reads feature flags for every active org."""
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"orgs": get_all_org_flags(), "feature_keys": list(FEATURE_KEYS), "feature_labels": FEATURE_LABELS})


@bp.route("/admin/org-flags/<org_id>", methods=["PUT"])
async def admin_set_org_flags_route(org_id: str):
    """Admin sets feature flags for a specific org."""
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Unauthorized"}), 401
    body  = await request.get_json() or {}
    flags = body.get("flags", {})
    actor = session.get("user_id") or SYSTEM
    result = set_org_flags(org_id, flags, actor=actor)
    _audit(session, "feature_flags_update", resource_type="org", resource_name=org_id)
    return jsonify({"flags": result})


# ── Server-side feature-flag enforcement ──────────────────────────────────────
# get_org_flags()/is_feature_enabled() back the admin Feature Access panel, but
# until now nothing ever called is_feature_enabled() outside of that panel's own
# GET/PUT routes — a disabled flag only hid a sidebar item, the API stayed fully
# callable. This makes it a real boundary. Checked longest/most-specific pattern
# first so a route with its own dedicated flag (e.g. vakalatnama under /matters)
# is matched before the broader parent route group's flag.
_FEATURE_ROUTE_RULES = [
    (re.compile(r"^/matters/[^/]+/vakalatnama$"), "vakalat"),
    (re.compile(r"^/diary/send-brief$"),          "whatsapp"),
    (re.compile(r"^/calendar/notify-holiday"),    "whatsapp"),
    (re.compile(r"^/team/[^/]+/whatsapp$"),       "whatsapp"),
    (re.compile(r"^/court-fees/calculate$"),      "matters"),
    (re.compile(r"^/associate-fees/summary$"),    "matters"),
    (re.compile(r"^/bail-stages"),                "matters"),
    (re.compile(r"^/matters(/|$)"),               "matters"),
    (re.compile(r"^/clients(/|$)"),               "clients"),
    (re.compile(r"^/hearings(/|$)"),               "calendar"),
    (re.compile(r"^/calendar(/|$)"),               "calendar"),
    (re.compile(r"^/diary(/|$)"),                  "diary"),
    (re.compile(r"^/invoices(/|$)"),                "invoices"),
    (re.compile(r"^/fees(/|$)"),                    "invoices"),
    (re.compile(r"^/outstanding-dues$"),            "invoices"),
    (re.compile(r"^/team(/|$)"),                    "team"),
    (re.compile(r"^/templates(/|$)"),               "drafting"),
    (re.compile(r"^/draft$"),                       "drafting"),
    (re.compile(r"^/cause-list(/|$)"),              "causelist"),
    (re.compile(r"^/vakalatnama-register$"),        "vakalat"),
    (re.compile(r"^/judge-notes(/|$)"),             "intelligence"),
    (re.compile(r"^/opposing-counsel(/|$)"),        "intelligence"),
    (re.compile(r"^/audit-logs$"),                  "audit"),
    (re.compile(r"^/client-tokens(/|$)"),           "client_portal"),
    (re.compile(r"^/documents(/|$)"),               "documents"),
    (re.compile(r"^/categories(/|$)"),              "documents"),
]


@bp.before_request
async def _enforce_feature_flags():
    """Block a route at the API layer when its feature is disabled for the org.

    Skips platform_admin (cross-org by design) and any request with no valid
    session — those are either public routes or get their own 401 from the
    route body. The public /portal/* routes are token-scoped, not session-scoped,
    so they check is_feature_enabled() directly instead of going through here.
    """
    session = _get_session()
    if not session or session.get("role") == "platform_admin":
        return None
    org_id = session.get("org")
    if not org_id:
        return None
    path = request.path
    for pattern, feature_key in _FEATURE_ROUTE_RULES:
        if pattern.match(path):
            if not is_feature_enabled(org_id, feature_key):
                return jsonify({
                    "error": f"'{FEATURE_LABELS.get(feature_key, feature_key)}' is not enabled for your organization.",
                    "feature": feature_key,
                }), 403
            break
    return None


@bp.route("/deadlines", methods=["GET"])
async def list_deadlines():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id    = session.get("org") or ""
    from_date = request.args.get("from_date") or None
    to_date   = request.args.get("to_date")   or None
    return jsonify(get_deadlines(org_id, from_date=from_date, to_date=to_date))


@bp.route("/deadlines", methods=["POST"])
async def add_deadline():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    title    = (data.get("title") or "").strip()
    due_date = (data.get("due_date") or "").strip()
    if not title or not due_date:
        return jsonify({"error": "title and due_date are required"}), 400
    org_id = session.get("org") or ""
    actor  = session.get("user_id") or SYSTEM
    d = create_deadline(
        org_id=org_id, title=title, due_date=due_date,
        matter_id     = data.get("matter_id")      or None,
        deadline_type = data.get("deadline_type")  or "Filing",
        notes         = data.get("notes")          or None,
        wa_reminder   = bool(data.get("wa_reminder", False)),
        actor=actor,
    )
    _audit(session, "deadline_create",
           resource_type="deadline", resource_name=title,
           details={"due_date": due_date, "matter_id": data.get("matter_id")})
    return jsonify(d), 201


@bp.route("/deadlines/<deadline_id>", methods=["PATCH"])
async def edit_deadline(deadline_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    d = get_deadline(deadline_id)
    if not d or d.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    data  = await request.get_json(silent=True) or {}
    actor = session.get("user_id") or SYSTEM
    updated = update_deadline(deadline_id, session.get("org") or "", actor=actor, **data)
    return jsonify(updated)


@bp.route("/deadlines/<deadline_id>", methods=["DELETE"])
async def remove_deadline(deadline_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    d = get_deadline(deadline_id)
    if not d or d.get("org_id") != session.get("org"):
        return jsonify({"error": "Not found"}), 404
    actor = session.get("user_id") or SYSTEM
    delete_deadline(deadline_id, session.get("org") or "", actor=actor)
    _audit(session, "deadline_delete", resource_type="deadline", resource_name=d.get("title", ""))
    return jsonify({"ok": True})


# ─── Case Law (PLD / SCMR / MLD / CLC) — Task #33 ──────────────────────────
# Documents are indexed into Azure Search with category="__case_law__" so they
# are visible to every authenticated user alongside their own org's documents.
# Only the platform admin can upload / delete case law documents.

CASE_LAW_CATEGORY = "__case_law__"

VALID_PUBLISHERS = {"PLD", "SCMR", "MLD", "CLC", "OTHER"}


@bp.route("/admin/case-law", methods=["GET"])
async def admin_list_case_law():
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Unauthorized"}), 401
    publisher = request.args.get("publisher")
    docs = list_case_law_docs(publisher or None)
    return jsonify({"docs": docs})


@bp.route("/admin/case-law/upload", methods=["POST"])
async def admin_upload_case_law():
    """Admin uploads a PLD/SCMR/MLD/CLC PDF to the shared case law index."""
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Unauthorized"}), 401

    files = await request.files
    form  = await request.form
    if "file" not in files:
        return jsonify({"error": "No file provided"}), 400

    uploaded_file = files["file"]
    filename: str = uploaded_file.filename or "upload"

    import os as _os
    ext = _os.path.splitext(filename)[1].lower()
    if ext != ".pdf":
        return jsonify({"error": "Only PDF files are supported for case law."}), 400

    publisher = (form.get("publisher") or "OTHER").upper()
    if publisher not in VALID_PUBLISHERS:
        publisher = "OTHER"
    title  = (form.get("title") or filename).strip()
    year_s = form.get("year") or ""
    volume = (form.get("volume") or "").strip() or None
    court  = (form.get("court")  or "").strip() or None

    try:
        year = int(year_s) if year_s.strip().isdigit() else None
    except ValueError:
        year = None

    actor = session.get("user_id") or SYSTEM

    try:
        import tempfile, asyncio as _asyncio
        from dotenv import load_dotenv as _ld
        _ld(dotenv_path=_os.path.join(_os.path.dirname(__file__), ".env"), override=True)

        from azure.core.credentials import AzureKeyCredential as _AKC
        from azure.identity.aio import AzureDeveloperCliCredential as _AzCred
        from prepdocslib.servicesetup import (
            OpenAIHost, setup_search_info, setup_blob_manager,
            setup_embeddings_service, setup_openai_client,
            build_file_processors,
        )
        from prepdocslib.listfilestrategy import LocalListFileStrategy
        from prepdocslib.filestrategy import FileStrategy
        from prepdocslib.strategy import DocumentAction

        azure_credential = _AzCred(process_timeout=60)
        OPENAI_HOST = OpenAIHost(_os.environ["OPENAI_HOST"])

        search_key = _os.getenv("AZURE_SEARCH_KEY")
        search_info = setup_search_info(
            search_service=_os.environ["AZURE_SEARCH_SERVICE"],
            index_name=_os.environ["AZURE_SEARCH_INDEX"],
            azure_credential=azure_credential,
            search_key=search_key,
        )

        blob_manager = setup_blob_manager(
            azure_credential=azure_credential,
            storage_account=_os.environ["AZURE_STORAGE_ACCOUNT"],
            storage_container=_os.environ["AZURE_STORAGE_CONTAINER"],
            storage_resource_group=_os.environ.get("AZURE_STORAGE_RESOURCE_GROUP"),
            subscription_id=_os.environ.get("AZURE_SUBSCRIPTION_ID"),
            storage_key=_os.getenv("AZURE_STORAGE_KEY"),
        )

        openai_client, azure_openai_endpoint = setup_openai_client(
            openai_host=OPENAI_HOST,
            azure_credential=azure_credential,
            azure_openai_service=_os.getenv("AZURE_OPENAI_SERVICE"),
            azure_openai_api_key=_os.getenv("AZURE_OPENAI_API_KEY_OVERRIDE"),
        )

        doc_int_service = _os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE")
        doc_int_key = _os.getenv("AZURE_DOCUMENTINTELLIGENCE_KEY")
        file_processors, figure_processor = build_file_processors(
            azure_credential=azure_credential,
            document_intelligence_service=doc_int_service,
            document_intelligence_key=doc_int_key if doc_int_key and doc_int_key != "YOUR_DOC_INTELLIGENCE_KEY_HERE" else None,
            use_local_pdf_parser=not doc_int_service,
            use_local_html_parser=not doc_int_service,
        ), None
        if isinstance(file_processors, tuple):
            file_processors, figure_processor = file_processors

        embeddings_service = setup_embeddings_service(
            OPENAI_HOST, openai_client,
            emb_model_name=_os.environ["AZURE_OPENAI_EMB_MODEL_NAME"],
            emb_model_dimensions=int(_os.getenv("AZURE_OPENAI_EMB_DIMENSIONS", "1536")),
            azure_openai_deployment=_os.getenv("AZURE_OPENAI_EMB_DEPLOYMENT"),
            azure_openai_endpoint=azure_openai_endpoint,
        )

        file_bytes = uploaded_file.read()
        size_bytes = len(file_bytes)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        # Create DB record
        doc_record = create_case_law_doc(
            publisher=publisher, title=title,
            filename=filename, size_bytes=size_bytes,
            year=year, volume=volume, court=court,
            actor=actor,
        )
        doc_id = doc_record["doc_id"]

        async def _index_case_law():
            try:
                list_strategy = LocalListFileStrategy(path_pattern=tmp_path)
                strategy = FileStrategy(
                    search_info=search_info,
                    list_file_strategy=list_strategy,
                    blob_manager=blob_manager,
                    file_processors=file_processors,
                    document_action=DocumentAction.Add,
                    embeddings=embeddings_service,
                    image_embeddings=None,
                    search_field_name_embedding=_os.getenv("AZURE_SEARCH_FIELD_NAME_EMBEDDING", "embedding"),
                    use_acls=False,
                    category=CASE_LAW_CATEGORY,   # ← shared pool, NOT org_id
                    figure_processor=figure_processor,
                )
                await strategy.setup()
                await strategy.run()
                set_case_law_doc_status(doc_id, "ready", actor=actor)
            except Exception as exc:
                set_case_law_doc_status(doc_id, "error", error_msg=str(exc), actor=actor)
                logging.getLogger(__name__).error("Case law indexing failed: %s", exc)
            finally:
                try:
                    _os.remove(tmp_path)
                    md5 = tmp_path + ".md5"
                    if _os.path.exists(md5):
                        _os.remove(md5)
                except Exception:
                    pass

        _asyncio.ensure_future(_index_case_law())
        _audit(session, "case_law_upload", resource_type="case_law_doc",
               resource_name=title, resource_id=doc_id)
        return jsonify({"doc": doc_record, "message": "Indexing started. Status will update to 'ready' when complete."})

    except Exception as exc:
        logging.getLogger(__name__).error("Case law upload error: %s", exc)
        return jsonify({"error": f"Upload failed: {exc}"}), 500


@bp.route("/admin/case-law/<doc_id>", methods=["DELETE"])
async def admin_delete_case_law(doc_id: str):
    session = _get_session()
    if not session or session.get("role") != "platform_admin":
        return jsonify({"error": "Unauthorized"}), 401
    doc = get_case_law_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    actor = session.get("user_id") or SYSTEM
    delete_case_law_doc(doc_id, actor=actor)
    _audit(session, "case_law_delete", resource_type="case_law_doc",
           resource_name=doc.get("title", ""), resource_id=doc_id)
    # Note: blob + search index cleanup left to admin manually for now
    # (or add a separate cleanup job — see PLACEHOLDERS.md)
    return jsonify({"ok": True})


# ─── WhatsApp reminder scheduler (runs hourly in background) ─────────────────

_scheduler_started = False


def _start_reminder_scheduler(app_ref):
    """Start APScheduler to send WhatsApp reminders once per hour."""
    global _scheduler_started
    if _scheduler_started:
        return
    try:
        from apscheduler.schedulers.background import BackgroundScheduler  # type: ignore
        scheduler = BackgroundScheduler(daemon=True)
        scheduler.add_job(_send_due_reminders, "interval", hours=1, id="wa_reminders")
        scheduler.add_job(_send_cause_list_nudge, "cron", hour=7, minute=0, id="cause_list_nudge")
        scheduler.add_job(_send_cause_list_digest, "cron", hour=8, minute=0, id="cause_list_digest")
        scheduler.start()
        _scheduler_started = True
        logging.getLogger(__name__).info("WhatsApp reminder scheduler started (hourly reminders, 7am cause-list nudge, 8am digest).")
    except ImportError:
        logging.getLogger(__name__).warning(
            "apscheduler not installed — WhatsApp reminders disabled. "
            "Run: pip install apscheduler"
        )


def _send_due_reminders():
    """Fire-and-forget: send WhatsApp reminders for hearings/deadlines due tomorrow."""
    import asyncio as _aio
    try:
        loop = _aio.new_event_loop()
        loop.run_until_complete(_send_due_reminders_async())
    except Exception as exc:
        logging.getLogger(__name__).error("Reminder scheduler error: %s", exc)


async def _send_due_reminders_async():
    import asyncio as _aio
    from whatsapp_helper import send_whatsapp_text  # type: ignore[import]
    from email_helper import (
        send_hearing_reminder_email as _hrmail,
        send_deadline_reminder_email as _dlmail,
    )
    _log = logging.getLogger(__name__)

    # ── Hearings ──────────────────────────────────────────────────────────────
    for h in get_hearings_needing_reminder():
        time_str = f" at {h['hearing_time']}" if h.get("hearing_time") else ""
        sent = False

        # WhatsApp
        wa = h.get("owner_wa")
        if wa:
            msg = (
                f"🏛️ *Project Ease Reminder*\n\n"
                f"*Hearing tomorrow{time_str}*\n"
                f"📌 {h['title']}\n"
                f"🗓️ {h['hearing_date']}\n"
                + (f"🏛️ {h['court_name']}\n" if h.get("court_name") else "")
                + (f"⚖️ {h['matter_title']}\n" if h.get("matter_title") else "")
                + "\n_Sent by Project Ease_"
            )
            try:
                await _aio.to_thread(send_whatsapp_text, wa, msg)
                sent = True
            except Exception as exc:
                _log.error("Hearing WA reminder error: %s", exc)

        # Email
        email = h.get("owner_email")
        if email:
            try:
                await _aio.to_thread(
                    _hrmail, email,
                    h.get("owner_name", "there"),
                    h.get("matter_title") or h.get("title", ""),
                    h["hearing_date"],
                    h.get("hearing_time") or "",
                    h.get("court_name") or "",
                )
                sent = True
            except Exception as exc:
                _log.error("Hearing email reminder error: %s", exc)

        if sent:
            mark_hearing_reminder_sent(h["hearing_id"])

    # ── Deadlines ─────────────────────────────────────────────────────────────
    for d in get_deadlines_needing_reminder():
        sent = False

        # WhatsApp
        wa = d.get("owner_wa")
        if wa:
            msg = (
                f"⚠️ *Project Ease Reminder*\n\n"
                f"*Deadline tomorrow*\n"
                f"📌 {d['title']}\n"
                f"🗓️ {d['due_date']} · {d.get('deadline_type', 'Filing')}\n"
                + (f"⚖️ {d['matter_title']}\n" if d.get("matter_title") else "")
                + "\n_Sent by Project Ease_"
            )
            try:
                await _aio.to_thread(send_whatsapp_text, wa, msg)
                sent = True
            except Exception as exc:
                _log.error("Deadline WA reminder error: %s", exc)

        # Email
        email = d.get("owner_email")
        if email:
            try:
                await _aio.to_thread(
                    _dlmail, email,
                    d.get("owner_name", "there"),
                    d.get("matter_title") or "",
                    d["due_date"],
                    d["title"],
                    d.get("deadline_type") or "Filing",
                )
                sent = True
            except Exception as exc:
                _log.error("Deadline email reminder error: %s", exc)

        if sent:
            mark_deadline_reminder_sent(d["deadline_id"])


def _send_cause_list_nudge():
    """07:00 — if an org owner hasn't pasted/uploaded today's cause list yet,
    send a one-line WhatsApp nudge. Runs as a sync APScheduler job."""
    import asyncio as _aio
    try:
        loop = _aio.new_event_loop()
        loop.run_until_complete(_send_cause_list_nudge_async())
    except Exception as exc:
        logging.getLogger(__name__).error("Cause-list nudge scheduler error: %s", exc)


async def _send_cause_list_nudge_async():
    import datetime as _dt
    from whatsapp_helper import send_whatsapp_text  # type: ignore[import]
    import asyncio as _aio
    _log = logging.getLogger(__name__)
    today = _dt.datetime.utcnow().strftime("%Y-%m-%d")
    for org in get_orgs_with_owner_wa():
        entries = get_cause_list_entries(org["org_id"], today)
        if entries:
            continue  # already pasted/uploaded — no nudge needed
        msg = (
            "🏛️ *Project Ease*\n\n"
            "Today's LHC cause list hasn't been added yet. Paste or upload it in "
            "*Cause List* so your 8am matter digest goes out on time."
        )
        try:
            await _aio.to_thread(send_whatsapp_text, org["owner_wa"], msg)
        except Exception as exc:
            _log.error("Cause-list nudge send error for org %s: %s", org["org_id"], exc)


def _send_cause_list_digest():
    """08:00 — send each org owner today's matched cause-list entries."""
    import asyncio as _aio
    try:
        loop = _aio.new_event_loop()
        loop.run_until_complete(_send_cause_list_digest_async())
    except Exception as exc:
        logging.getLogger(__name__).error("Cause-list digest scheduler error: %s", exc)


async def _send_cause_list_digest_async():
    from whatsapp_helper import send_whatsapp_text  # type: ignore[import]
    import asyncio as _aio
    _log = logging.getLogger(__name__)
    for org in get_orgs_with_owner_wa():
        matches = get_today_cause_list_matches(org["org_id"])
        if not matches:
            continue  # nothing matched today — nothing to send
        lines = [f"🏛️ *Morning Brief — {len(matches)} of your matters are listed today*\n"]
        for m in matches:
            label = m.get("bench") or m.get("court_name") or m.get("matter_court") or "Court TBD"
            lines.append(f"• {m.get('matter_title', 'Matter')} — {label}")
        lines.append("\n_Sent automatically by Project Ease at 8am_")
        try:
            await _aio.to_thread(send_whatsapp_text, org["owner_wa"], "\n".join(lines))
        except Exception as exc:
            _log.error("Cause-list digest send error for org %s: %s", org["org_id"], exc)


# ─── Answer Export (PDF via browser print / Word via python-docx) ─────────────

@bp.route("/export/answer", methods=["POST"])
async def export_answer():
    """Generate a .docx file from an AI answer and return it for download.

    Request JSON:
        question  : str   – the user's question
        answer    : str   – the AI's answer text
        citations : list  – list of citation strings (filenames / doc titles)
        org_name  : str   – firm name for header branding
    """
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        from docx import Document as _DocxDocument  # type: ignore[import]
        from docx.shared import Pt, RGBColor, Inches  # type: ignore[import]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 503

    body = await request.get_json(force=True) or {}
    question  = (body.get("question")  or "").strip()
    answer    = (body.get("answer")    or "").strip()
    citations = [str(c) for c in (body.get("citations") or [])]
    org_name  = (body.get("org_name")  or "Project Ease").strip()

    if not answer:
        return jsonify({"error": "answer is required"}), 400

    doc = _DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Brand header ──────────────────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(org_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xB8, 0x96, 0x4C)  # --gold

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("AI Research Export — Project Ease")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    import datetime as _dt
    date_str = _dt.datetime.now().strftime("%-d %B %Y")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Question ──────────────────────────────────────────────────────────────
    if question:
        q_label = doc.add_paragraph()
        q_label_run = q_label.add_run("Question")
        q_label_run.bold = True
        q_label_run.font.size = Pt(10)
        q_label_run.font.color.rgb = RGBColor(0xB8, 0x96, 0x4C)

        q_para = doc.add_paragraph()
        q_para.paragraph_format.left_indent = Inches(0.25)
        q_run = q_para.add_run(question)
        q_run.font.size = Pt(11)
        q_run.italic = True

        doc.add_paragraph()  # spacer

    # ── Answer ────────────────────────────────────────────────────────────────
    a_label = doc.add_paragraph()
    a_label_run = a_label.add_run("Answer")
    a_label_run.bold = True
    a_label_run.font.size = Pt(10)
    a_label_run.font.color.rgb = RGBColor(0xB8, 0x96, 0x4C)

    # Split answer into paragraphs so line breaks render correctly
    for para_text in answer.split("\n"):
        if para_text.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            run_a = p.add_run(para_text)
            run_a.font.size = Pt(11)

    # ── Citations ─────────────────────────────────────────────────────────────
    if citations:
        doc.add_paragraph()  # spacer
        c_label = doc.add_paragraph()
        c_label_run = c_label.add_run("Sources")
        c_label_run.bold = True
        c_label_run.font.size = Pt(10)
        c_label_run.font.color.rgb = RGBColor(0xB8, 0x96, 0x4C)

        for i, citation in enumerate(citations, 1):
            c_para = doc.add_paragraph()
            c_para.paragraph_format.left_indent = Inches(0.25)
            c_run = c_para.add_run(f"{i}. {citation}")
            c_run.font.size = Pt(10)
            c_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Generated by Project Ease · AI answers should be verified by a qualified lawyer.")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.italic = True

    # ── Stream file ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import re as _re
    safe_q = _re.sub(r"[^\w\s-]", "", question[:40]).strip().replace(" ", "_") if question else "answer"
    filename = f"ProjectEase_Export_{safe_q}.docx"

    response = await make_response(buf.read())
    response.headers["Content-Type"]        = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ─── Templates & Document Drafting ─────────────────────────────────────────

@bp.route("/templates", methods=["GET"])
async def templates_list():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    ttype  = request.args.get("type")
    rows   = list_templates(org_id, ttype if ttype else None)
    return jsonify(rows)


@bp.route("/templates", methods=["POST"])
async def templates_create():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]
    body   = await request.get_json(force=True) or {}

    title         = (body.get("title") or "").strip()
    template_type = (body.get("template_type") or "general").strip()
    content       = (body.get("content") or "").strip()
    description   = (body.get("description") or "").strip()

    if not title:
        return jsonify({"error": "title is required"}), 400
    if template_type not in TEMPLATE_TYPES:
        return jsonify({"error": f"template_type must be one of {TEMPLATE_TYPES}"}), 400

    row = create_template(org_id, title, template_type, content, description, actor)
    log_event(org_id=org_id, user_id=actor, event_type="template_created",
              resource_type="template", resource_id=row["template_id"], resource_name=title)
    return jsonify(row), 201


@bp.route("/templates/<template_id>", methods=["PATCH"])
async def templates_update(template_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]

    existing = get_template(template_id)
    if not existing or existing["org_id"] != org_id:
        return jsonify({"error": "Not found"}), 404

    body          = await request.get_json(force=True) or {}
    title         = body.get("title")
    template_type = body.get("template_type")
    content       = body.get("content")
    description   = body.get("description")

    if template_type is not None and template_type not in TEMPLATE_TYPES:
        return jsonify({"error": f"template_type must be one of {TEMPLATE_TYPES}"}), 400

    row = update_template(template_id, actor,
                          title=title, template_type=template_type,
                          content=content, description=description)
    log_event(org_id=org_id, user_id=actor, event_type="template_updated",
              resource_type="template", resource_id=template_id,
              resource_name=row["title"] if row else template_id)
    return jsonify(row)


@bp.route("/templates/<template_id>", methods=["DELETE"])
async def templates_delete(template_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]

    existing = get_template(template_id)
    if not existing or existing["org_id"] != org_id:
        return jsonify({"error": "Not found"}), 404

    delete_template(template_id, actor)
    log_event(org_id=org_id, user_id=actor, event_type="template_deleted",
              resource_type="template", resource_id=template_id, resource_name=existing["title"])
    return jsonify({"ok": True})


@bp.route("/draft", methods=["POST"])
async def draft_document():
    """Fill a template with matter/client context and AI, return .docx.

    Request JSON:
        template_id : str  – which template to use
        matter_id   : str  – matter to pull context from (optional)
        extra_vars  : dict – caller-supplied {{variable}} overrides (optional)
    """
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]

    try:
        from docx import Document as _DocxDocument  # type: ignore[import]
        from docx.shared import Pt, RGBColor, Inches  # type: ignore[import]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 503

    body        = await request.get_json(force=True) or {}
    template_id = (body.get("template_id") or "").strip()
    matter_id   = (body.get("matter_id")   or "").strip()
    extra_vars  = body.get("extra_vars") or {}

    if not template_id:
        return jsonify({"error": "template_id is required"}), 400

    # ── Load template ──────────────────────────────────────────────────────────
    tmpl = get_template(template_id)
    if not tmpl or tmpl["org_id"] != org_id:
        return jsonify({"error": "Template not found"}), 404

    # ── Load org name ──────────────────────────────────────────────────────────
    org_row = get_org(org_id)
    org_name = org_row["name"] if org_row else "Your Firm"

    # ── Load matter + client context ───────────────────────────────────────────
    matter_ctx: dict = {}
    if matter_id:
        m = get_matter_with_docs(matter_id)
        if m and m["org_id"] == org_id:
            client_row = None
            if m.get("client_id"):
                clients = get_clients(org_id)
                client_row = next((c for c in clients if c["client_id"] == m["client_id"]), None)
            matter_ctx = {
                "matter_title":       m.get("title", ""),
                "matter_type":        m.get("matter_type", ""),
                "matter_description": m.get("description", ""),
                "case_number":        m.get("case_number", ""),
                "court_name":         m.get("court", ""),
                "client_name":        (client_row or {}).get("name", ""),
                "client_cnic":        (client_row or {}).get("cnic", ""),
                "client_email":       (client_row or {}).get("email", ""),
                "client_phone":       (client_row or {}).get("phone", ""),
            }

    import datetime as _dt
    known_vars: dict = {
        "org_name":     org_name,
        "date":         _dt.date.today().strftime("%d-%m-%Y"),
        "date_long":    _dt.date.today().strftime("%-d %B %Y"),
        "advocate_name": org_row.get("contact_name", "") if org_row else "",
        **matter_ctx,
        **{str(k): str(v) for k, v in extra_vars.items()},
    }

    # ── Fill known variables ───────────────────────────────────────────────────
    import re as _re
    content = tmpl["content"]
    for var, val in known_vars.items():
        content = content.replace("{{" + var + "}}", val)

    # ── Find remaining unfilled placeholders ───────────────────────────────────
    remaining = list(dict.fromkeys(_re.findall(r"\{\{(\w+)\}\}", content)))

    # ── AI-fill remaining placeholders via Azure OpenAI ───────────────────────
    if remaining and current_app.config.get("AZURE_CONFIGURED"):
        try:
            openai_client = current_app.config[CONFIG_OPENAI_CLIENT]
            context_summary = "\n".join(f"  {k}: {v}" for k, v in known_vars.items() if v)
            vars_list       = ", ".join("{{" + v + "}}" for v in remaining)
            system_msg = (
                "You are a Pakistani legal document drafting assistant. "
                "Fill in the placeholder variables in a legal document based on the context provided. "
                "Return ONLY a JSON object mapping each variable name to its filled value. "
                "Use formal legal language appropriate for Pakistani courts. "
                "If a value cannot be determined, use a sensible placeholder like '[To be provided]'."
            )
            user_msg = (
                f"Document type: {tmpl['template_type']}\n"
                f"Known context:\n{context_summary}\n\n"
                f"Fill these variables: {vars_list}\n\n"
                "Return JSON only, e.g.: {\"variable_name\": \"filled value\"}"
            )
            ai_resp = await openai_client.chat.completions.create(
                model=current_app.config.get("OPENAI_CHATGPT_MODEL", "gpt-5-mini-1"),
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user",   "content": user_msg},
                ],
                temperature=0.2,
                max_tokens=800,
            )
            raw = (ai_resp.choices[0].message.content or "").strip()
            # Strip markdown fences if present
            raw = _re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=_re.MULTILINE).strip()
            ai_vars = json.loads(raw)
            for var, val in ai_vars.items():
                content = content.replace("{{" + var + "}}", str(val))
        except Exception as exc:
            logging.getLogger(__name__).warning("AI draft fill failed: %s", exc)
            # Leave remaining placeholders as-is — document still downloads

    # ── Build .docx ───────────────────────────────────────────────────────────
    doc = _DocxDocument()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Brand header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = h.add_run(org_name)
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = RGBColor(0xB8, 0x96, 0x4C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subr = sub.add_run(tmpl["title"])
    subr.font.size = Pt(11)
    subr.bold = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_r = date_p.add_run(_dt.date.today().strftime("%-d %B %Y"))
    date_r.font.size = Pt(9)
    date_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # Body — split by line breaks, preserve empty lines as paragraph breaks
    for line in content.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(11)

    # Footer disclaimer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Generated by Project Ease · Verify all details before filing.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = _re.sub(r"[^\w\s-]", "", tmpl["title"][:40]).strip().replace(" ", "_")
    filename = f"ProjectEase_Draft_{safe_title}.docx"

    log_event(org_id=org_id, user_id=actor, event_type="document_drafted",
              resource_type="template", resource_id=template_id, resource_name=tmpl["title"])

    resp = await make_response(buf.read())
    resp.headers["Content-Type"]        = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


# ─── Client Portal ────────────────────────────────────────────────────────────

@bp.route("/client-tokens", methods=["GET"])
async def client_tokens_list():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id    = session.get("org") or ""
    client_id = request.args.get("client_id")
    rows      = list_client_tokens(org_id, client_id if client_id else None)
    return jsonify(rows)


@bp.route("/client-tokens", methods=["POST"])
async def client_tokens_create():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]
    body   = await request.get_json(force=True) or {}

    client_id  = (body.get("client_id") or "").strip()
    matter_id  = (body.get("matter_id") or "").strip() or None
    label      = (body.get("label") or "").strip() or None
    expires_at = (body.get("expires_at") or "").strip() or None

    if not client_id:
        return jsonify({"error": "client_id is required"}), 400

    # Verify client belongs to org
    clients = get_clients(org_id)
    if not any(c["client_id"] == client_id for c in clients):
        return jsonify({"error": "Client not found"}), 404

    row = create_client_token(org_id, client_id, actor,
                              matter_id=matter_id, label=label, expires_at=expires_at)
    log_event(org_id=org_id, user_id=actor, event_type="client_token_created",
              resource_type="client_token", resource_id=row["token_id"],
              resource_name=row.get("client_name", ""))
    return jsonify(row), 201


@bp.route("/client-tokens/<token_id>", methods=["DELETE"])
async def client_tokens_revoke(token_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    org_id = session.get("org") or ""
    actor  = session["user_id"]

    existing = get_client_token_by_id(token_id)
    if not existing or existing["org_id"] != org_id:
        return jsonify({"error": "Not found"}), 404

    revoke_client_token(token_id, actor)
    log_event(org_id=org_id, user_id=actor, event_type="client_token_revoked",
              resource_type="client_token", resource_id=token_id,
              resource_name=existing.get("client_name", ""))
    return jsonify({"ok": True})


# ── Unauthenticated portal endpoints (token in query param) ──────────────────

@bp.route("/portal/me", methods=["GET"])
async def portal_me():
    """Client fetches their own info using a secret token (no login required)."""
    token = request.args.get("token", "").strip()
    if not token:
        return jsonify({"error": "token is required"}), 400
    row = get_client_token_by_value(token)
    if not row:
        return jsonify({"error": "Invalid or expired token"}), 401
    if not is_feature_enabled(row["org_id"], "client_portal"):
        return jsonify({"error": "The client portal is not enabled for this organization."}), 403
    return jsonify({
        "client_name":   row["client_name"],
        "client_email":  row.get("client_email"),
        "client_phone":  row.get("client_phone"),
        "matter_title":  row.get("matter_title"),
        "matter_type":   row.get("matter_type"),
        "case_number":   row.get("case_number"),
        "court_name":    row.get("court_name"),
        "matter_status": row.get("matter_status"),
        "org_name":      row["org_name"],
        "label":         row.get("label"),
    })


@bp.route("/portal/documents", methods=["GET"])
async def portal_documents():
    """Client fetches documents shared with them via a token."""
    token = request.args.get("token", "").strip()
    if not token:
        return jsonify({"error": "token is required"}), 400
    row = get_client_token_by_value(token)
    if not row:
        return jsonify({"error": "Invalid or expired token"}), 401
    if not is_feature_enabled(row["org_id"], "client_portal"):
        return jsonify({"error": "The client portal is not enabled for this organization."}), 403

    org_id    = row["org_id"]
    matter_id = row.get("matter_id")

    # Return documents linked to this matter (or all org docs if no matter scoped)
    if matter_id:
        matter = get_matter_with_docs(matter_id)
        docs   = matter.get("documents", []) if matter else []
    else:
        # No matter scope — return nothing (security: don't expose all org docs)
        docs = []

    return jsonify([{
        "doc_id":        d["doc_id"],
        "name":          d.get("filename", d.get("name", "")),
        "size_bytes":    d.get("size_bytes", 0),
        "status":        d.get("status", "ready"),
        "uploaded_at":   d.get("uploaded_at", d.get("created_at", "")),
        "category_name": d.get("category_name", ""),
    } for d in docs if d.get("status") == "ready"])


@bp.route("/portal/documents/<doc_id>/download", methods=["GET"])
async def portal_document_download(doc_id: str):
    """Client downloads a specific document (token-authenticated)."""
    token = request.args.get("token", "").strip()
    if not token:
        return jsonify({"error": "token is required"}), 400
    row = get_client_token_by_value(token)
    if not row:
        return jsonify({"error": "Invalid or expired token"}), 401
    if not is_feature_enabled(row["org_id"], "client_portal"):
        return jsonify({"error": "The client portal is not enabled for this organization."}), 403

    org_id    = row["org_id"]
    matter_id = row.get("matter_id")

    if not matter_id:
        return jsonify({"error": "No documents available"}), 403

    # Verify the doc belongs to this matter
    matter = get_matter_with_docs(matter_id)
    if not matter:
        return jsonify({"error": "Matter not found"}), 404
    doc_ids = [d["doc_id"] for d in (matter.get("documents") or [])]
    if doc_id not in doc_ids:
        return jsonify({"error": "Document not found or not shared"}), 403

    # Stream from Azure Blob Storage
    if not current_app.config.get("AZURE_CONFIGURED"):
        return jsonify({"error": "Azure not configured"}), 503

    try:
        blob_manager: BlobManager = current_app.config[CONFIG_GLOBAL_BLOB_MANAGER]
        # Find doc name from matter docs
        doc_entry = next((d for d in matter["documents"] if d["doc_id"] == doc_id), None)
        filename  = doc_entry.get("filename", doc_id) if doc_entry else doc_id
        stream    = await blob_manager.download_blob(filename)
        content   = await stream.readall()
        mime, _   = mimetypes.guess_type(filename)
        resp = await make_response(content)
        resp.headers["Content-Type"]        = mime or "application/octet-stream"
        resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp
    except Exception as exc:
        logging.getLogger(__name__).error("Portal download error: %s", exc)
        return jsonify({"error": "Download failed"}), 500


@bp.before_app_serving
async def setup_clients():
    # Load .env — works whether entry point is main.py or quart --app app:create_app
    from dotenv import load_dotenv as _load_dotenv
    _load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"), override=True)

    # ── PROJECT EASE: graceful local-dev mode ──────────────────────────────────
    # When Azure credentials are not configured the app starts in auth-only mode.
    # /auth/* endpoints work; Azure-backed routes (/chat, /config, etc.) will
    # return 503 until real credentials are provided.
    if not os.environ.get("AZURE_STORAGE_ACCOUNT"):
        current_app.logger.warning(
            "AZURE_STORAGE_ACCOUNT not set — starting in auth-only mode. "
            "Chat and search features require Azure credentials."
        )
        current_app.config["AZURE_CONFIGURED"] = False
        return
    current_app.config["AZURE_CONFIGURED"] = True

    # Replace these with your own values, either in environment variables or directly here
    AZURE_STORAGE_ACCOUNT = os.environ["AZURE_STORAGE_ACCOUNT"]
    AZURE_STORAGE_CONTAINER = os.environ["AZURE_STORAGE_CONTAINER"]
    AZURE_IMAGESTORAGE_CONTAINER = os.environ.get("AZURE_IMAGESTORAGE_CONTAINER")
    AZURE_USERSTORAGE_ACCOUNT = os.environ.get("AZURE_USERSTORAGE_ACCOUNT")
    AZURE_USERSTORAGE_CONTAINER = os.environ.get("AZURE_USERSTORAGE_CONTAINER")
    AZURE_SEARCH_SERVICE = os.environ["AZURE_SEARCH_SERVICE"]
    AZURE_SEARCH_ENDPOINT = f"https://{AZURE_SEARCH_SERVICE}.search.windows.net"
    AZURE_SEARCH_INDEX = os.environ["AZURE_SEARCH_INDEX"]
    AZURE_SEARCH_KNOWLEDGEBASE_NAME = os.getenv("AZURE_SEARCH_KNOWLEDGEBASE_NAME", "")
    # Shared by all OpenAI deployments
    OPENAI_HOST = OpenAIHost(os.getenv("OPENAI_HOST", "azure"))
    OPENAI_CHATGPT_MODEL = os.environ["AZURE_OPENAI_CHATGPT_MODEL"]
    AZURE_OPENAI_KNOWLEDGEBASE_MODEL = os.getenv("AZURE_OPENAI_KNOWLEDGEBASE_MODEL")
    AZURE_OPENAI_KNOWLEDGEBASE_DEPLOYMENT = os.getenv("AZURE_OPENAI_KNOWLEDGEBASE_DEPLOYMENT")
    OPENAI_EMB_MODEL = os.getenv("AZURE_OPENAI_EMB_MODEL_NAME", "text-embedding-ada-002")
    OPENAI_EMB_DIMENSIONS = int(os.getenv("AZURE_OPENAI_EMB_DIMENSIONS") or 1536)
    OPENAI_REASONING_EFFORT = os.getenv("AZURE_OPENAI_REASONING_EFFORT")
    # Used with Azure OpenAI deployments
    AZURE_OPENAI_SERVICE = os.getenv("AZURE_OPENAI_SERVICE")
    AZURE_OPENAI_CHATGPT_DEPLOYMENT = (
        os.getenv("AZURE_OPENAI_CHATGPT_DEPLOYMENT")
        if OPENAI_HOST in [OpenAIHost.AZURE, OpenAIHost.AZURE_CUSTOM]
        else None
    )
    AZURE_OPENAI_EMB_DEPLOYMENT = (
        os.getenv("AZURE_OPENAI_EMB_DEPLOYMENT") if OPENAI_HOST in [OpenAIHost.AZURE, OpenAIHost.AZURE_CUSTOM] else None
    )
    AZURE_OPENAI_CUSTOM_URL = os.getenv("AZURE_OPENAI_CUSTOM_URL")
    AZURE_VISION_ENDPOINT = os.getenv("AZURE_VISION_ENDPOINT", "")
    AZURE_OPENAI_API_KEY_OVERRIDE = os.getenv("AZURE_OPENAI_API_KEY_OVERRIDE")
    # Used only with non-Azure OpenAI deployments
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    OPENAI_ORGANIZATION = os.getenv("OPENAI_ORGANIZATION")

    AZURE_TENANT_ID = os.getenv("AZURE_TENANT_ID")
    AZURE_USE_AUTHENTICATION = os.getenv("AZURE_USE_AUTHENTICATION", "").lower() == "true"
    AZURE_ENFORCE_ACCESS_CONTROL = os.getenv("AZURE_ENFORCE_ACCESS_CONTROL", "").lower() == "true"
    AZURE_ENABLE_UNAUTHENTICATED_ACCESS = os.getenv("AZURE_ENABLE_UNAUTHENTICATED_ACCESS", "").lower() == "true"
    AZURE_SERVER_APP_ID = os.getenv("AZURE_SERVER_APP_ID")
    AZURE_SERVER_APP_SECRET = os.getenv("AZURE_SERVER_APP_SECRET")
    AZURE_CLIENT_APP_ID = os.getenv("AZURE_CLIENT_APP_ID")
    AZURE_AUTH_TENANT_ID = os.getenv("AZURE_AUTH_TENANT_ID", AZURE_TENANT_ID)

    KB_FIELDS_CONTENT = os.getenv("KB_FIELDS_CONTENT", "content")
    KB_FIELDS_SOURCEPAGE = os.getenv("KB_FIELDS_SOURCEPAGE", "sourcepage")

    AZURE_SEARCH_QUERY_LANGUAGE = os.getenv("AZURE_SEARCH_QUERY_LANGUAGE") or "en-us"
    AZURE_SEARCH_QUERY_SPELLER = os.getenv("AZURE_SEARCH_QUERY_SPELLER") or "lexicon"
    AZURE_SEARCH_SEMANTIC_RANKER = os.getenv("AZURE_SEARCH_SEMANTIC_RANKER", "free").lower()
    AZURE_SEARCH_QUERY_REWRITING = os.getenv("AZURE_SEARCH_QUERY_REWRITING", "false").lower()
    # This defaults to the previous field name "embedding", for backwards compatibility
    AZURE_SEARCH_FIELD_NAME_EMBEDDING = os.getenv("AZURE_SEARCH_FIELD_NAME_EMBEDDING", "embedding")

    AZURE_SPEECH_SERVICE_ID = os.getenv("AZURE_SPEECH_SERVICE_ID")
    AZURE_SPEECH_SERVICE_LOCATION = os.getenv("AZURE_SPEECH_SERVICE_LOCATION")
    AZURE_SPEECH_SERVICE_VOICE = os.getenv("AZURE_SPEECH_SERVICE_VOICE") or "en-US-AndrewMultilingualNeural"

    USE_MULTIMODAL = os.getenv("USE_MULTIMODAL", "").lower() == "true"
    RAG_SEARCH_TEXT_EMBEDDINGS = os.getenv("RAG_SEARCH_TEXT_EMBEDDINGS", "true").lower() == "true"
    RAG_SEARCH_IMAGE_EMBEDDINGS = os.getenv("RAG_SEARCH_IMAGE_EMBEDDINGS", "true").lower() == "true"
    RAG_SEND_TEXT_SOURCES = os.getenv("RAG_SEND_TEXT_SOURCES", "true").lower() == "true"
    RAG_SEND_IMAGE_SOURCES = os.getenv("RAG_SEND_IMAGE_SOURCES", "true").lower() == "true"
    USE_USER_UPLOAD = os.getenv("USE_USER_UPLOAD", "").lower() == "true"
    ENABLE_LANGUAGE_PICKER = os.getenv("ENABLE_LANGUAGE_PICKER", "").lower() == "true"
    USE_SPEECH_INPUT_BROWSER = os.getenv("USE_SPEECH_INPUT_BROWSER", "").lower() == "true"
    USE_SPEECH_OUTPUT_BROWSER = os.getenv("USE_SPEECH_OUTPUT_BROWSER", "").lower() == "true"
    USE_SPEECH_OUTPUT_AZURE = os.getenv("USE_SPEECH_OUTPUT_AZURE", "").lower() == "true"
    USE_CHAT_HISTORY_BROWSER = os.getenv("USE_CHAT_HISTORY_BROWSER", "").lower() == "true"
    USE_CHAT_HISTORY_COSMOS = os.getenv("USE_CHAT_HISTORY_COSMOS", "").lower() == "true"
    USE_AGENTIC_KNOWLEDGEBASE = os.getenv("USE_AGENTIC_KNOWLEDGEBASE", "").lower() == "true"
    USE_WEB_SOURCE = os.getenv("USE_WEB_SOURCE", "").lower() == "true"
    USE_SHAREPOINT_SOURCE = os.getenv("USE_SHAREPOINT_SOURCE", "").lower() == "true"
    AGENTIC_KNOWLEDGEBASE_REASONING_EFFORT = os.getenv("AGENTIC_KNOWLEDGEBASE_REASONING_EFFORT", "minimal")
    USE_VECTORS = os.getenv("USE_VECTORS", "").lower() != "false"

    if USE_SHAREPOINT_SOURCE and not AZURE_ENFORCE_ACCESS_CONTROL:
        current_app.logger.warning("AZURE_ENFORCE_ACCESS_CONTROL must be true when USE_SHAREPOINT_SOURCE is true")

    # WEBSITE_HOSTNAME is always set by App Service, RUNNING_IN_PRODUCTION is set in main.bicep
    RUNNING_ON_AZURE = os.getenv("WEBSITE_HOSTNAME") is not None or os.getenv("RUNNING_IN_PRODUCTION") is not None

    # Use the current user identity for keyless authentication to Azure services.
    # This assumes you use 'azd auth login' locally, and managed identity when deployed on Azure.
    # The managed identity is setup in the infra/ folder.
    azure_credential: AzureDeveloperCliCredential | ManagedIdentityCredential
    azure_ai_token_provider: Callable[[], Awaitable[str]]
    if RUNNING_ON_AZURE:
        current_app.logger.info("Setting up Azure credential using ManagedIdentityCredential")
        if AZURE_CLIENT_ID := os.getenv("AZURE_CLIENT_ID"):
            # ManagedIdentityCredential should use AZURE_CLIENT_ID if set in env, but its not working for some reason,
            # so we explicitly pass it in as the client ID here. This is necessary for user-assigned managed identities.
            current_app.logger.info(
                "Setting up Azure credential using ManagedIdentityCredential with client_id %s", AZURE_CLIENT_ID
            )
            azure_credential = ManagedIdentityCredential(client_id=AZURE_CLIENT_ID)
        else:
            current_app.logger.info("Setting up Azure credential using ManagedIdentityCredential")
            azure_credential = ManagedIdentityCredential()
    elif AZURE_TENANT_ID:
        current_app.logger.info(
            "Setting up Azure credential using AzureDeveloperCliCredential with tenant_id %s", AZURE_TENANT_ID
        )
        azure_credential = AzureDeveloperCliCredential(tenant_id=AZURE_TENANT_ID, process_timeout=60)
    else:
        current_app.logger.info("Setting up Azure credential using AzureDeveloperCliCredential for home tenant")
        azure_credential = AzureDeveloperCliCredential(process_timeout=60)
    azure_ai_token_provider = get_bearer_token_provider(
        azure_credential, "https://cognitiveservices.azure.com/.default"
    )

    # Set the Azure credential in the app config for use in other parts of the app
    current_app.config[CONFIG_CREDENTIAL] = azure_credential

    # Set up clients for AI Search and Storage
    AZURE_SEARCH_KEY = os.getenv("AZURE_SEARCH_KEY")
    search_credential = AzureKeyCredential(AZURE_SEARCH_KEY) if AZURE_SEARCH_KEY else azure_credential
    search_client = SearchClient(
        endpoint=AZURE_SEARCH_ENDPOINT,
        index_name=AZURE_SEARCH_INDEX,
        credential=search_credential,
    )

    knowledgebase_client = KnowledgeBaseRetrievalClient(
        endpoint=AZURE_SEARCH_ENDPOINT,
        knowledge_base_name=AZURE_SEARCH_KNOWLEDGEBASE_NAME,
        credential=azure_credential,
    )
    knowledgebase_client_with_web = None
    knowledgebase_client_with_sharepoint = None
    knowledgebase_client_with_web_and_sharepoint = None

    if AZURE_SEARCH_KNOWLEDGEBASE_NAME:
        if USE_WEB_SOURCE:
            knowledgebase_client_with_web = KnowledgeBaseRetrievalClient(
                endpoint=AZURE_SEARCH_ENDPOINT,
                knowledge_base_name=f"{AZURE_SEARCH_KNOWLEDGEBASE_NAME}-with-web",
                credential=azure_credential,
            )
        if USE_SHAREPOINT_SOURCE:
            knowledgebase_client_with_sharepoint = KnowledgeBaseRetrievalClient(
                endpoint=AZURE_SEARCH_ENDPOINT,
                knowledge_base_name=f"{AZURE_SEARCH_KNOWLEDGEBASE_NAME}-with-sp",
                credential=azure_credential,
            )
        if USE_WEB_SOURCE and USE_SHAREPOINT_SOURCE:
            knowledgebase_client_with_web_and_sharepoint = KnowledgeBaseRetrievalClient(
                endpoint=AZURE_SEARCH_ENDPOINT,
                knowledge_base_name=f"{AZURE_SEARCH_KNOWLEDGEBASE_NAME}-with-web-and-sp",
                credential=azure_credential,
            )

    # Set up the global blob storage manager (used for global content/images, but not user uploads)
    global_blob_manager = BlobManager(
        endpoint=f"https://{AZURE_STORAGE_ACCOUNT}.blob.core.windows.net",
        credential=azure_credential,
        container=AZURE_STORAGE_CONTAINER,
        image_container=AZURE_IMAGESTORAGE_CONTAINER,
    )
    current_app.config[CONFIG_GLOBAL_BLOB_MANAGER] = global_blob_manager

    # Set up authentication helper
    search_index = None
    if AZURE_USE_AUTHENTICATION:
        current_app.logger.info("AZURE_USE_AUTHENTICATION is true, setting up search index client")
        search_index_client = SearchIndexClient(
            endpoint=AZURE_SEARCH_ENDPOINT,
            credential=azure_credential,
        )
        search_index = await search_index_client.get_index(AZURE_SEARCH_INDEX)
        await search_index_client.close()
    auth_helper = AuthenticationHelper(
        search_index=search_index,
        use_authentication=AZURE_USE_AUTHENTICATION,
        server_app_id=AZURE_SERVER_APP_ID,
        server_app_secret=AZURE_SERVER_APP_SECRET,
        client_app_id=AZURE_CLIENT_APP_ID,
        tenant_id=AZURE_AUTH_TENANT_ID,
        enforce_access_control=AZURE_ENFORCE_ACCESS_CONTROL,
        enable_unauthenticated_access=AZURE_ENABLE_UNAUTHENTICATED_ACCESS,
    )

    if USE_SPEECH_OUTPUT_AZURE:
        current_app.logger.info("USE_SPEECH_OUTPUT_AZURE is true, setting up Azure speech service")
        if not AZURE_SPEECH_SERVICE_ID or AZURE_SPEECH_SERVICE_ID == "":
            raise ValueError("Azure speech resource not configured correctly, missing AZURE_SPEECH_SERVICE_ID")
        if not AZURE_SPEECH_SERVICE_LOCATION or AZURE_SPEECH_SERVICE_LOCATION == "":
            raise ValueError("Azure speech resource not configured correctly, missing AZURE_SPEECH_SERVICE_LOCATION")
        current_app.config[CONFIG_SPEECH_SERVICE_ID] = AZURE_SPEECH_SERVICE_ID
        current_app.config[CONFIG_SPEECH_SERVICE_LOCATION] = AZURE_SPEECH_SERVICE_LOCATION
        current_app.config[CONFIG_SPEECH_SERVICE_VOICE] = AZURE_SPEECH_SERVICE_VOICE
        # Wait until token is needed to fetch for the first time
        current_app.config[CONFIG_SPEECH_SERVICE_TOKEN] = None

    openai_client, azure_openai_endpoint = setup_openai_client(
        openai_host=OPENAI_HOST,
        azure_credential=azure_credential,
        azure_openai_service=AZURE_OPENAI_SERVICE,
        azure_openai_custom_url=AZURE_OPENAI_CUSTOM_URL,
        azure_openai_api_key=AZURE_OPENAI_API_KEY_OVERRIDE,
        openai_api_key=OPENAI_API_KEY,
        openai_organization=OPENAI_ORGANIZATION,
    )

    user_blob_manager = None
    if USE_USER_UPLOAD:
        current_app.logger.info("USE_USER_UPLOAD is true, setting up user upload feature")
        if not AZURE_USERSTORAGE_ACCOUNT or not AZURE_USERSTORAGE_CONTAINER:
            raise ValueError(
                "AZURE_USERSTORAGE_ACCOUNT and AZURE_USERSTORAGE_CONTAINER must be set when USE_USER_UPLOAD is true"
            )
        if not AZURE_ENFORCE_ACCESS_CONTROL:
            raise ValueError("AZURE_ENFORCE_ACCESS_CONTROL must be true when USE_USER_UPLOAD is true")
        user_blob_manager = AdlsBlobManager(
            endpoint=f"https://{AZURE_USERSTORAGE_ACCOUNT}.dfs.core.windows.net",
            container=AZURE_USERSTORAGE_CONTAINER,
            credential=azure_credential,
        )
        current_app.config[CONFIG_USER_BLOB_MANAGER] = user_blob_manager

        # Set up ingester
        file_processors, figure_processor = setup_file_processors(
            azure_credential=azure_credential,
            document_intelligence_service=os.getenv("AZURE_DOCUMENTINTELLIGENCE_SERVICE"),
            local_pdf_parser=os.getenv("USE_LOCAL_PDF_PARSER", "").lower() == "true",
            local_html_parser=os.getenv("USE_LOCAL_HTML_PARSER", "").lower() == "true",
            use_content_understanding=os.getenv("USE_CONTENT_UNDERSTANDING", "").lower() == "true",
            content_understanding_endpoint=os.getenv("AZURE_CONTENTUNDERSTANDING_ENDPOINT"),
            use_multimodal=USE_MULTIMODAL,
            openai_client=openai_client,
            openai_model=OPENAI_CHATGPT_MODEL,
            openai_deployment=AZURE_OPENAI_CHATGPT_DEPLOYMENT if OPENAI_HOST == OpenAIHost.AZURE else None,
        )
        search_info = setup_search_info(
            search_service=AZURE_SEARCH_SERVICE,
            index_name=AZURE_SEARCH_INDEX,
            azure_credential=azure_credential,
            use_agentic_knowledgebase=USE_AGENTIC_KNOWLEDGEBASE,
            azure_openai_endpoint=azure_openai_endpoint,
            knowledgebase_name=AZURE_SEARCH_KNOWLEDGEBASE_NAME,
            azure_openai_knowledgebase_deployment=AZURE_OPENAI_KNOWLEDGEBASE_DEPLOYMENT,
            azure_openai_knowledgebase_model=AZURE_OPENAI_KNOWLEDGEBASE_MODEL,
        )

        text_embeddings_service = None
        if USE_VECTORS:
            text_embeddings_service = setup_embeddings_service(
                open_ai_client=openai_client,
                openai_host=OPENAI_HOST,
                emb_model_name=OPENAI_EMB_MODEL,
                emb_model_dimensions=OPENAI_EMB_DIMENSIONS,
                azure_openai_deployment=AZURE_OPENAI_EMB_DEPLOYMENT,
                azure_openai_endpoint=azure_openai_endpoint,
            )

        image_embeddings_service = setup_image_embeddings_service(
            azure_credential=azure_credential,
            vision_endpoint=AZURE_VISION_ENDPOINT,
            use_multimodal=USE_MULTIMODAL,
        )
        ingester = UploadUserFileStrategy(
            search_info=search_info,
            file_processors=file_processors,
            embeddings=text_embeddings_service,
            image_embeddings=image_embeddings_service,
            search_field_name_embedding=AZURE_SEARCH_FIELD_NAME_EMBEDDING,
            blob_manager=user_blob_manager,
            figure_processor=figure_processor,
        )
        current_app.config[CONFIG_INGESTER] = ingester

    image_embeddings_client = None
    if USE_MULTIMODAL:
        image_embeddings_client = ImageEmbeddings(AZURE_VISION_ENDPOINT, azure_ai_token_provider)

    current_app.config[CONFIG_OPENAI_CLIENT] = openai_client
    current_app.config[CONFIG_SEARCH_CLIENT] = search_client
    current_app.config[CONFIG_KNOWLEDGEBASE_CLIENT] = knowledgebase_client
    current_app.config[CONFIG_KNOWLEDGEBASE_CLIENT_WITH_WEB] = knowledgebase_client_with_web
    current_app.config[CONFIG_KNOWLEDGEBASE_CLIENT_WITH_SHAREPOINT] = knowledgebase_client_with_sharepoint
    current_app.config[CONFIG_KNOWLEDGEBASE_CLIENT_WITH_WEB_AND_SHAREPOINT] = (
        knowledgebase_client_with_web_and_sharepoint
    )
    current_app.config[CONFIG_AUTH_CLIENT] = auth_helper

    current_app.config[CONFIG_SEMANTIC_RANKER_DEPLOYED] = AZURE_SEARCH_SEMANTIC_RANKER != "disabled"
    current_app.config[CONFIG_QUERY_REWRITING_ENABLED] = (
        AZURE_SEARCH_QUERY_REWRITING == "true" and AZURE_SEARCH_SEMANTIC_RANKER != "disabled"
    )
    current_app.config[CONFIG_DEFAULT_REASONING_EFFORT] = OPENAI_REASONING_EFFORT
    current_app.config[CONFIG_DEFAULT_RETRIEVAL_REASONING_EFFORT] = AGENTIC_KNOWLEDGEBASE_REASONING_EFFORT
    current_app.config[CONFIG_REASONING_EFFORT_ENABLED] = Approach.is_reasoning_model(OPENAI_CHATGPT_MODEL)
    current_app.config[CONFIG_REASONING_EFFORT_OPTIONS] = Approach.get_reasoning_effort_options(OPENAI_CHATGPT_MODEL)
    current_app.config[CONFIG_STREAMING_ENABLED] = True
    current_app.config[CONFIG_VECTOR_SEARCH_ENABLED] = bool(USE_VECTORS)
    current_app.config[CONFIG_USER_UPLOAD_ENABLED] = bool(USE_USER_UPLOAD)
    current_app.config[CONFIG_LANGUAGE_PICKER_ENABLED] = ENABLE_LANGUAGE_PICKER
    current_app.config[CONFIG_SPEECH_INPUT_ENABLED] = USE_SPEECH_INPUT_BROWSER
    current_app.config[CONFIG_SPEECH_OUTPUT_BROWSER_ENABLED] = USE_SPEECH_OUTPUT_BROWSER
    current_app.config[CONFIG_SPEECH_OUTPUT_AZURE_ENABLED] = USE_SPEECH_OUTPUT_AZURE
    current_app.config[CONFIG_CHAT_HISTORY_BROWSER_ENABLED] = USE_CHAT_HISTORY_BROWSER
    current_app.config[CONFIG_CHAT_HISTORY_COSMOS_ENABLED] = USE_CHAT_HISTORY_COSMOS
    current_app.config[CONFIG_AGENTIC_KNOWLEDGEBASE_ENABLED] = USE_AGENTIC_KNOWLEDGEBASE
    current_app.config[CONFIG_MULTIMODAL_ENABLED] = USE_MULTIMODAL
    current_app.config[CONFIG_RAG_SEARCH_TEXT_EMBEDDINGS] = RAG_SEARCH_TEXT_EMBEDDINGS
    current_app.config[CONFIG_RAG_SEARCH_IMAGE_EMBEDDINGS] = RAG_SEARCH_IMAGE_EMBEDDINGS
    current_app.config[CONFIG_RAG_SEND_TEXT_SOURCES] = RAG_SEND_TEXT_SOURCES
    current_app.config[CONFIG_RAG_SEND_IMAGE_SOURCES] = RAG_SEND_IMAGE_SOURCES
    current_app.config[CONFIG_WEB_SOURCE_ENABLED] = USE_WEB_SOURCE
    if AGENTIC_KNOWLEDGEBASE_REASONING_EFFORT == "minimal" and current_app.config[CONFIG_WEB_SOURCE_ENABLED]:
        raise ValueError("Web source cannot be used with minimal retrieval reasoning effort")
    current_app.config[CONFIG_SHAREPOINT_SOURCE_ENABLED] = USE_SHAREPOINT_SOURCE

    prompt_manager = PromptManager()

    # ChatReadRetrieveReadApproach is used by /chat for multi-turn conversation
    current_app.config[CONFIG_CHAT_APPROACH] = ChatReadRetrieveReadApproach(
        search_client=search_client,
        search_index_name=AZURE_SEARCH_INDEX,
        knowledgebase_model=AZURE_OPENAI_KNOWLEDGEBASE_MODEL,
        knowledgebase_deployment=AZURE_OPENAI_KNOWLEDGEBASE_DEPLOYMENT,
        knowledgebase_client=knowledgebase_client,
        knowledgebase_client_with_web=knowledgebase_client_with_web,
        knowledgebase_client_with_sharepoint=knowledgebase_client_with_sharepoint,
        knowledgebase_client_with_web_and_sharepoint=knowledgebase_client_with_web_and_sharepoint,
        openai_client=openai_client,
        chatgpt_model=OPENAI_CHATGPT_MODEL,
        chatgpt_deployment=AZURE_OPENAI_CHATGPT_DEPLOYMENT,
        embedding_model=OPENAI_EMB_MODEL,
        embedding_deployment=AZURE_OPENAI_EMB_DEPLOYMENT,
        embedding_dimensions=OPENAI_EMB_DIMENSIONS,
        embedding_field=AZURE_SEARCH_FIELD_NAME_EMBEDDING,
        sourcepage_field=KB_FIELDS_SOURCEPAGE,
        content_field=KB_FIELDS_CONTENT,
        query_language=AZURE_SEARCH_QUERY_LANGUAGE,
        query_speller=AZURE_SEARCH_QUERY_SPELLER,
        prompt_manager=prompt_manager,
        reasoning_effort=OPENAI_REASONING_EFFORT,
        multimodal_enabled=USE_MULTIMODAL,
        image_embeddings_client=image_embeddings_client,
        global_blob_manager=global_blob_manager,
        user_blob_manager=user_blob_manager,
        use_web_source=current_app.config[CONFIG_WEB_SOURCE_ENABLED],
        use_sharepoint_source=current_app.config[CONFIG_SHAREPOINT_SOURCE_ENABLED],
        retrieval_reasoning_effort=AGENTIC_KNOWLEDGEBASE_REASONING_EFFORT,
    )


@bp.after_app_serving
async def close_clients():
    await current_app.config[CONFIG_SEARCH_CLIENT].close()
    await current_app.config[CONFIG_GLOBAL_BLOB_MANAGER].close_clients()
    if user_blob_manager := current_app.config.get(CONFIG_USER_BLOB_MANAGER):
        await user_blob_manager.close_clients()
    await current_app.config[CONFIG_CREDENTIAL].close()


# ── Legal Notices — Task #165 ─────────────────────────────────────────────────

@bp.route("/legal-notices", methods=["GET"])
async def list_legal_notices():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    matter_id = request.args.get("matter_id") or None
    notices = get_legal_notices(session.get("org") or "", matter_id=matter_id)
    return jsonify({"notices": notices, "notice_types": list(NOTICE_TYPES),
                    "notice_via": list(NOTICE_VIA), "statuses": list(NOTICE_STATUSES)})


@bp.route("/legal-notices", methods=["POST"])
async def add_legal_notice():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    notice = create_legal_notice(
        org_id=session.get("org") or "", actor=actor,
        sent_to=b.get("sent_to", ""), sent_date=b.get("sent_date", ""),
        subject=b.get("subject", ""), notice_type=b.get("notice_type", "General"),
        sent_via=b.get("sent_via", "Registered Post"), matter_id=b.get("matter_id"),
        client_id=b.get("client_id"), response_due=b.get("response_due"),
        content=b.get("content"), tracking_no=b.get("tracking_no"), notes=b.get("notes"),
    )
    _audit(session, "legal_notice_create", resource_type="legal_notice", resource_name=b.get("subject",""))
    return jsonify(notice), 201


@bp.route("/legal-notices/<notice_id>", methods=["PATCH"])
async def edit_legal_notice(notice_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    return jsonify(update_legal_notice(notice_id, session.get("org") or "", actor=actor, **b))


@bp.route("/legal-notices/<notice_id>", methods=["DELETE"])
async def remove_legal_notice(notice_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    delete_legal_notice(notice_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ── Court Transfers — Task #170 ───────────────────────────────────────────────

@bp.route("/matters/<matter_id>/transfers", methods=["GET"])
async def list_court_transfers(matter_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"transfers": get_court_transfers(matter_id, session.get("org") or "")})


@bp.route("/matters/<matter_id>/transfers", methods=["POST"])
async def add_court_transfer(matter_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    t = create_court_transfer(
        matter_id=matter_id, org_id=session.get("org") or "", actor=actor,
        transfer_date=b.get("transfer_date", ""), to_court=b.get("to_court", ""),
        from_court=b.get("from_court"), from_judge=b.get("from_judge"),
        to_judge=b.get("to_judge"), reason=b.get("reason"),
        order_ref=b.get("order_ref"), notes=b.get("notes"),
    )
    return jsonify(t), 201


@bp.route("/matters/<matter_id>/transfers/<transfer_id>", methods=["PATCH"])
async def edit_court_transfer(matter_id: str, transfer_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    updated = update_court_transfer(transfer_id, session.get("org") or "",
                                    actor=session.get("user_id") or SYSTEM, **b)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/transfers/<transfer_id>", methods=["DELETE"])
async def remove_court_transfer(matter_id: str, transfer_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    delete_court_transfer(transfer_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ── Bail Bonds — Task #167 ────────────────────────────────────────────────────

@bp.route("/matters/<matter_id>/bail-bonds", methods=["GET"])
async def list_bail_bonds(matter_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"bonds": get_bail_bonds(matter_id, session.get("org") or ""),
                    "bail_types": list(BAIL_TYPES), "statuses": list(BAIL_STATUSES)})


@bp.route("/matters/<matter_id>/bail-bonds", methods=["POST"])
async def add_bail_bond(matter_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    bond = create_bail_bond(
        matter_id=matter_id, org_id=session.get("org") or "", actor=actor,
        accused_name=b.get("accused_name", ""), surety_name=b.get("surety_name", ""),
        bail_amount_pkr=float(b.get("bail_amount_pkr") or 0),
        bail_type=b.get("bail_type", "Post-Arrest"),
        surety_cnic=b.get("surety_cnic"), surety_address=b.get("surety_address"),
        surety_property=b.get("surety_property"),
        property_value=float(b["property_value"]) if b.get("property_value") else None,
        court=b.get("court"), judge=b.get("judge"),
        granted_date=b.get("granted_date"), expiry_date=b.get("expiry_date"),
        status=b.get("status", "Active"), bail_order_ref=b.get("bail_order_ref"),
        notes=b.get("notes"),
    )
    return jsonify(bond), 201


@bp.route("/matters/<matter_id>/bail-bonds/<bond_id>", methods=["PATCH"])
async def edit_bail_bond(matter_id: str, bond_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    updated = update_bail_bond(bond_id, session.get("org") or "",
                               actor=session.get("user_id") or SYSTEM, **b)
    if not updated:
        return jsonify({"error": "Not found"}), 404
    return jsonify(updated)


@bp.route("/matters/<matter_id>/bail-bonds/<bond_id>", methods=["DELETE"])
async def remove_bail_bond(matter_id: str, bond_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    delete_bail_bond(bond_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ── Bail Workflow Checklist (configurable, default = 6-stage flow) ──────────

@bp.route("/bail-stages", methods=["GET"])
async def list_bail_stages():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    include_inactive = request.args.get("all") == "1" and session.get("role") == "org_owner"
    return jsonify({"stages": get_org_bail_stages(session.get("org") or "", include_inactive=include_inactive)})


@bp.route("/bail-stages", methods=["POST"])
async def add_bail_stage():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data  = await request.get_json(silent=True) or {}
    label = (data.get("label") or "").strip()
    if not label:
        return jsonify({"error": "label is required"}), 400
    stage = add_org_bail_stage(session.get("org") or "", label, actor=session.get("user_id") or SYSTEM)
    return jsonify(stage), 201


@bp.route("/bail-stages/<stage_key>", methods=["PATCH"])
async def edit_bail_stage(stage_key: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    update_org_bail_stage(
        session.get("org") or "", stage_key,
        **{k: v for k, v in data.items() if k in {"label", "sort_order", "is_active"}},
    )
    return jsonify({"ok": True})


@bp.route("/matters/<matter_id>/bail-bonds/<bond_id>/stages", methods=["GET"])
async def get_bail_bond_stages(matter_id: str, bond_id: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    stages      = get_org_bail_stages(session.get("org") or "")
    completions = get_bail_stage_completions(bond_id)
    return jsonify({
        "stages": [
            {**s, "completed_at": (completions.get(s["stage_key"]) or {}).get("completed_at"),
             "completed_by": (completions.get(s["stage_key"]) or {}).get("completed_by")}
            for s in stages
        ]
    })


@bp.route("/matters/<matter_id>/bail-bonds/<bond_id>/stages/<stage_key>", methods=["PATCH"])
async def toggle_bail_bond_stage(matter_id: str, bond_id: str, stage_key: str):
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    data = await request.get_json(silent=True) or {}
    result = set_bail_stage_completion(
        bond_id, stage_key, completed=bool(data.get("completed", True)),
        actor=session.get("user_id") or SYSTEM, notes=data.get("notes"),
    )
    return jsonify(result)


# ── Staff & Attendance — Task #171 ────────────────────────────────────────────

@bp.route("/staff", methods=["GET"])
async def list_staff():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"staff": get_staff(session.get("org") or ""),
                    "roles": list(STAFF_ROLES)})


@bp.route("/staff", methods=["POST"])
async def add_staff():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    s = create_staff(
        org_id=session.get("org") or "", actor=actor,
        name=b.get("name", ""), role=b.get("role", "Munshi"),
        monthly_salary_pkr=float(b.get("monthly_salary_pkr") or 0),
        join_date=b.get("join_date"), cnic=b.get("cnic"),
        phone=b.get("phone"), notes=b.get("notes"),
    )
    return jsonify(s), 201


@bp.route("/staff/<staff_id>", methods=["PATCH"])
async def edit_staff(staff_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    return jsonify(update_staff(staff_id, session.get("org") or "",
                                actor=session.get("user_id") or SYSTEM, **b))


@bp.route("/staff/<staff_id>", methods=["DELETE"])
async def remove_staff(staff_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_staff(staff_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


@bp.route("/staff/attendance", methods=["GET"])
async def list_attendance():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    staff_id = request.args.get("staff_id") or None
    month    = request.args.get("month") or None
    return jsonify({"attendance": get_attendance(session.get("org") or "", staff_id=staff_id, month=month),
                    "statuses": list(ATT_STATUSES)})


@bp.route("/staff/attendance", methods=["POST"])
async def mark_attendance():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    rec = upsert_attendance(
        org_id=session.get("org") or "", actor=actor,
        staff_id=b.get("staff_id", ""), att_date=b.get("att_date", ""),
        status=b.get("status", "Present"), time_in=b.get("time_in"),
        time_out=b.get("time_out"), notes=b.get("notes"),
    )
    return jsonify(rec)


@bp.route("/staff/salary", methods=["GET"])
async def list_salary():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    staff_id = request.args.get("staff_id") or None
    return jsonify({"payments": get_salary_payments(session.get("org") or "", staff_id=staff_id),
                    "pay_modes": list(SALARY_PAY_MODES)})


@bp.route("/staff/salary", methods=["POST"])
async def add_salary_payment():
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    b = await request.get_json() or {}
    actor = session.get("user_id") or SYSTEM
    p = create_salary_payment(
        org_id=session.get("org") or "", actor=actor,
        staff_id=b.get("staff_id", ""), month=b.get("month", ""),
        gross_pkr=float(b.get("gross_pkr") or 0),
        advance_deduction=float(b.get("advance_deduction") or 0),
        absence_deduction=float(b.get("absence_deduction") or 0),
        paid_date=b.get("paid_date"), payment_mode=b.get("payment_mode", "Cash"),
        notes=b.get("notes"),
    )
    return jsonify(p), 201


@bp.route("/staff/salary/<payment_id>", methods=["DELETE"])
async def remove_salary_payment(payment_id: str):
    session = _get_session()
    if not session or session.get("role") != "org_owner":
        return jsonify({"error": "Unauthorized"}), 401
    delete_salary_payment(payment_id, session.get("org") or "", actor=session.get("user_id") or SYSTEM)
    return jsonify({"ok": True})


# ── Outstanding Dues — Task #169 ──────────────────────────────────────────────

@bp.route("/outstanding-dues", methods=["GET"])
async def list_outstanding_dues():
    session = _get_session()
    if not session or session.get("role") not in ("org_owner", "employee"):
        return jsonify({"error": "Unauthorized"}), 401
    dues = get_outstanding_invoices(session.get("org") or "")
    return jsonify({"invoices": dues, "total_outstanding": sum(float(d.get("total_amount") or 0) for d in dues)})


def create_app():
    app = Quart(__name__)
    app.register_blueprint(bp)
    app.register_blueprint(chat_history_cosmosdb_bp)

    if os.getenv("APPLICATIONINSIGHTS_CONNECTION_STRING"):
        app.logger.info("APPLICATIONINSIGHTS_CONNECTION_STRING is set, enabling Azure Monitor")
        configure_azure_monitor(
            instrumentation_options={
                "django": {"enabled": False},
                "psycopg2": {"enabled": False},
                "fastapi": {"enabled": False},
            }
        )
        # This tracks HTTP requests made by aiohttp:
        AioHttpClientInstrumentor().instrument()
        # This tracks HTTP requests made by httpx:
        HTTPXClientInstrumentor().instrument()
        # This tracks OpenAI SDK requests:
        OpenAIInstrumentor().instrument()
        # This middleware tracks app route requests:
        app.asgi_app = OpenTelemetryMiddleware(app.asgi_app)  # type: ignore[assignment]

    # Log levels should be one of https://docs.python.org/3/library/logging.html#logging-levels
    # Set root level to WARNING to avoid seeing overly verbose logs from SDKS
    logging.basicConfig(level=logging.WARNING)
    # Set our own logger levels to INFO by default
    app_level = os.getenv("APP_LOG_LEVEL", "INFO")
    app.logger.setLevel(os.getenv("APP_LOG_LEVEL", app_level))
    logging.getLogger("scripts").setLevel(app_level)

    if allowed_origin := os.getenv("ALLOWED_ORIGIN"):
        allowed_origins = allowed_origin.split(";")
        if len(allowed_origins) > 0:
            app.logger.info("CORS enabled for %s", allowed_origins)
            cors(app, allow_origin=allowed_origins, allow_methods=["GET", "POST"])

    # Start WhatsApp reminder scheduler (Task #32)
    _start_reminder_scheduler(app)

    return app
